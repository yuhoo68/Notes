# vers 1.05
import base64
import email
import html
import io
import os
import re
import urllib.parse
import zipfile
# from typing import Callable

import pandas as pd
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document


# --- SIMPLE HTML SANITIZER (no bleach) ---

_DANGEROUS_TAGS = {
    "script", "iframe", "object", "embed", "link", "meta", "base",
    "form", "input", "button", "textarea", "select", "option",
    "svg", "math",
}

_ALLOWED_TAGS = {
    "div", "p", "span", "br",
    "strong", "b", "em", "i", "u", "s",
    "a", "img",
    "ul", "ol", "li",
    "h1", "h2", "h3", "h4", "h5", "h6",
    "blockquote", "pre", "code",
    "table", "thead", "tbody", "tfoot", "tr", "td", "th", "colgroup", "col",
    "mark",
}

_ALLOWED_ATTRS_COMMON = {"style", "class"}

_ALLOWED_ATTRS = {
    "a": {"href", "title", "target", "rel"} | _ALLOWED_ATTRS_COMMON,
    "img": {"src", "alt", "title", "width", "height", "style"} | _ALLOWED_ATTRS_COMMON,
    "td": {"colspan", "rowspan"} | _ALLOWED_ATTRS_COMMON,
    "th": {"colspan", "rowspan"} | _ALLOWED_ATTRS_COMMON,
    "col": {"span", "width"} | _ALLOWED_ATTRS_COMMON,
    "mark": _ALLOWED_ATTRS_COMMON,
    "*": _ALLOWED_ATTRS_COMMON,
    "li": {"value"} | _ALLOWED_ATTRS_COMMON,
    "ol": {"start"} | _ALLOWED_ATTRS_COMMON,
}

_ALLOWED_CSS_PROPS = {
    "font-family", "font-size", "font-weight", "font-style",
    "text-decoration", "color", "line-height",
    "background", "background-color",
    "text-align",
    "margin", "margin-top", "margin-right", "margin-bottom", "margin-left",
    "padding", "padding-top", "padding-right", "padding-bottom", "padding-left",
    "border", "border-width", "border-style", "border-color",
    "border-collapse",
    "vertical-align",
    "width",
    "white-space",
}

_RE_STYLE_DECL = re.compile(r"\s*([-\w]+)\s*:\s*([^;]+)\s*;?", re.UNICODE)


def _is_bad_url(url: str) -> bool:
    u = (url or "").strip().lower()
    return (
        u.startswith("javascript:")
        or u.startswith("vbscript:")
        or u.startswith("data:text/html")
        or u.startswith("data:application/xhtml+xml")
    )


def _sanitize_style(style: str) -> str:
    if not style:
        return ""
    s = style
    s = re.sub(r"expression\s*\([^)]*\)", "", s, flags=re.I)
    s = re.sub(r"javascript\s*:", "", s, flags=re.I)
    s = re.sub(r"url\s*\(\s*['\"]?\s*javascript:[^)]*\)", "", s, flags=re.I)
    s = re.sub(r"[\x00-\x1f\x7f]", "", s)

    out: list[str] = []
    for prop, val in _RE_STYLE_DECL.findall(s):
        prop_l = prop.strip().lower()
        val = val.strip()

        # OneNote
        if prop_l == "mso-highlight":
            prop_l = "background-color"

        if prop_l not in _ALLOWED_CSS_PROPS:
            continue

        # запретим любые url(...) в значениях
        if re.search(r"url\s*\(", val, flags=re.I):
            continue

        out.append(f"{prop_l}: {val}")

    return "; ".join(out)


def sanitize_html_safe(html_in: str) -> str:
    """
    Универсальная санитизация без bleach:
    - удаляет опасные теги
    - оставляет только whitelist тегов/атрибутов
    - режет on* обработчики
    - фильтрует href/src
    - чистит style по whitelist CSS свойств
    """
    if not html_in:
        return ""

    soup = BeautifulSoup(html_in, "html.parser")
    root = soup.body or soup

    # 1) удалить опасные теги целиком
    for bad in list(root.find_all(list(_DANGEROUS_TAGS))):
        bad.decompose()

    # 2) нормализация и чистка
    for tag in list(root.find_all(True)):
        name = (tag.name or "").lower()

        # неизвестные теги — unwrap (сохраняем текст)
        if name not in _ALLOWED_TAGS:
            tag.unwrap()
            continue

        allowed_attrs = _ALLOWED_ATTRS.get(name, _ALLOWED_ATTRS["*"])

        for attr in list(tag.attrs.keys()):
            al = str(attr).lower()

            # убрать обработчики событий
            if al.startswith("on"):
                del tag.attrs[attr]
                continue

            # whitelist атрибутов
            if attr not in allowed_attrs:
                del tag.attrs[attr]
                continue

        # href/src safety
        if tag.has_attr("href") and _is_bad_url(tag.get("href", "")):
            del tag.attrs["href"]
        if tag.has_attr("src") and _is_bad_url(tag.get("src", "")):
            del tag.attrs["src"]

        # style
        if tag.has_attr("style"):
            safe_style = _sanitize_style(tag.get("style", ""))
            if safe_style:
                tag["style"] = safe_style
            else:
                del tag.attrs["style"]

        # ссылки
        if name == "a":
            tag["target"] = "_blank"
            tag["rel"] = "noopener noreferrer"

    return str(root)


def preserve_onenote_blank_paragraphs(html_in: str) -> str:
    """
    OneNote пустые строки между абзацами делает как <p>&nbsp;</p>.
    Мы превращаем такие абзацы в <p><br/></p>, чтобы они не удалялись
    нормализацией/санитизацией и корректно отображались в редакторе.
    """
    if not html_in:
        return ""

    soup = BeautifulSoup(html_in, "html.parser")
    root = soup.body or soup

    for p in root.find_all("p"):
        if not isinstance(p, Tag):
            continue

        # если внутри есть реальная структура — не трогаем
        if p.find(["img", "table", "a", "br"]):
            continue

        txt = p.get_text("", strip=False) or ""
        # только пробелы/переводы + NBSP
        if txt.replace("\xa0", "").strip() == "":
            p.clear()
            p.append(soup.new_tag("br"))

    return str(root)


def _onenote_keep_blank_lines_in_div(page_div: Tag) -> None:
    """
    OneNote пустые строки = <p>&nbsp;</p>.
    Превращаем их в <p><br/></p>, чтобы _normalize_onenote_page_div не удалил.
    """
    if not isinstance(page_div, Tag):
        return

    soup = page_div if hasattr(page_div, "new_tag") else None  # обычно это BeautifulSoup Tag

    for p in page_div.find_all("p"):
        if not isinstance(p, Tag):
            continue

        # если в <p> уже есть структура — не трогаем
        if p.find(["img", "table", "a", "br"]):
            continue

        # текст без пробелов, но NBSP считаем пустотой
        txt = p.get_text("", strip=False) or ""
        if txt.replace("\xa0", "").strip() == "":
            p.clear()
            # создаём <br>
            br = p._parent.new_tag("br") if getattr(p, "_parent", None) else None
            if br is None:
                # fallback
                br = BeautifulSoup("", "html.parser").new_tag("br")
            p.append(br)


def sanitize_html_simple_keep_formatting(html_in: str) -> str:
    if not html_in:
        return ""

    soup = BeautifulSoup(html_in, "html.parser")
    root = soup.body or soup

    # 1) убрать опасные теги целиком
    for bad in list(root.find_all(_DANGEROUS_TAGS)):
        bad.decompose()

    # 2) пройти по всем тегам
    for tag in list(root.find_all(True)):
        name = (tag.name or "").lower()

        # неразрешённые теги разворачиваем, сохраняя текст
        if name not in _ALLOWED_TAGS:
            tag.unwrap()
            continue

        allowed = _ALLOWED_ATTRS.get(name, _ALLOWED_ATTRS.get("*", set()))
        allowed_all = set(allowed) | set(_ALLOWED_ATTRS.get("*", set()))

        # attrs whitelist + remove on*
        for attr in list(tag.attrs.keys()):
            al = str(attr).lower()
            if al.startswith("on"):
                del tag.attrs[attr]
                continue
            if attr not in allowed_all:
                del tag.attrs[attr]
                continue

        # href/src safety
        if tag.has_attr("href") and _is_bad_url(tag.get("href", "")):
            del tag.attrs["href"]
        if tag.has_attr("src") and _is_bad_url(tag.get("src", "")):
            del tag.attrs["src"]

        # clean style
        if tag.has_attr("style"):
            safe_style = _sanitize_style(tag.get("style", ""))
            if safe_style:
                tag["style"] = safe_style
            else:
                del tag.attrs["style"]

        # normalize links
        if name == "a":
            tag["target"] = "_blank"
            tag["rel"] = "noopener noreferrer"

    return str(root)


_IMPORT_ALLOWED_PROTOCOLS = ["http", "https", "mailto", "data"]

_RE_STYLE_DECL2 = re.compile(r"\s*([-\w]+)\s*:\s*([^;]+)\s*;?", re.UNICODE)

_MONTHS_RU = r"(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)"
_RE_DATE_RU = re.compile(rf"^\s*\d{{1,2}}\s+{_MONTHS_RU}\s+\d{{4}}\s*г\.?\s*$", re.IGNORECASE)
_RE_TIME = re.compile(r"^\s*\d{1,2}:\d{2}\s*$")


def strip_onenote_datetime_block(html: str) -> str:
    if not html:
        return ""

    soup = BeautifulSoup(html, "html.parser")
    root = soup.body or soup

    # Сначала удаляем p со временем/датой
    for p in root.find_all("p"):
        txt = (p.get_text(" ", strip=True) or "").replace("\xa0", " ").strip()
        if not txt:
            continue

        if _RE_TIME.match(txt) or _RE_DATE_RU.match(txt):
            p.decompose()

    # Затем подчистим пустые div, которые могли остаться
    for div in root.find_all("div"):
        # если в div нет текста и нет картинок/таблиц/ссылок — убрать
        text = div.get_text(" ", strip=True).replace("\xa0", " ").strip()
        has_struct = div.find(["img", "table", "a"])
        if (not text) and (not has_struct):
            div.decompose()

    return str(root)


def _looks_like_imported_html(html: str) -> bool:
    if not html:
        return False
    h = html.lstrip().lower()
    if h.startswith("<div") and "direction:ltr" in h and "font-family" in h:
        return True
    if "mso-" in h or "margin-left:0in" in h:
        return True
    return False

def _unwrap_onenote_header_wrappers(page_div: Tag) -> None:
    """
    OneNote экспортирует заголовок/дату в узких wrapper-div с width:...in,
    из-за чего в редакторе/preview может быть плохой перенос.
    Здесь мы НЕ удаляем заголовок/дату, а превращаем эти div в нормальные <p>:
      <div style="...width:1.7in"><p>...</p></div>  ->  <p>...</p>
    Так сохраняется гиперссылка в заголовке (<a>...</a>) и даты/время.
    """
    if not isinstance(page_div, Tag):
        return

    # берем только прямых детей, т.к. header/date у OneNote обычно сверху
    for div in list(page_div.find_all("div", recursive=False)):
        if not isinstance(div, Tag):
            continue

        # если внутри нет вложенных div и есть только p (1..N) — это типичный wrapper
        inner_divs = div.find_all("div", recursive=False)
        if inner_divs:
            continue

        ps = div.find_all("p", recursive=False)
        if not ps:
            continue

        # wrapper должен быть "простым": кроме <p> и пробелов ничего
        ok = True
        for c in div.contents:
            if isinstance(c, NavigableString) and not str(c).strip():
                continue
            if isinstance(c, Tag) and c.name and c.name.lower() == "p":
                continue
            ok = False
            break
        if not ok:
            continue

        # разворачиваем wrapper: переносим <p> на уровень page_div
        insert_before = div
        for p in ps:
            # вынимаем p из div и вставляем перед div
            p.extract()
            insert_before.insert_before(p)

        # удаляем пустой wrapper
        div.decompose()


def html_to_body(text: str, fallback_title: str):
    soup = BeautifulSoup(text, "html.parser")
    title = soup.title.string.strip() if soup.title and soup.title.string else fallback_title
    body = str(soup.body or soup)
    return title, body


def _is_onenote_datetime_div(div: Tag) -> bool:
    if not isinstance(div, Tag):
        return False
    ps = div.find_all("p")
    if not ps:
        return False
    text = div.get_text(" ", strip=True)
    if len(text) > 40:
        return False
    for p in ps:
        style = (p.get("style") or "").lower()
        if "font-size:10" not in style and "font-size:9" not in style and "font-size:11" not in style:
            return False
        if "color:#767676" not in style and "color: #767676" not in style:
            return False
    return True

def _is_onenote_datetime_p(p: Tag) -> bool:
    if not isinstance(p, Tag):
        return False
    text = p.get_text(" ", strip=True)
    if not text or len(text) > 40:
        return False
    style = (p.get("style") or "").lower()
    if "font-size:10" not in style and "font-size:9" not in style and "font-size:11" not in style:
        return False
    if "color:#767676" not in style and "color: #767676" not in style and "color:gray" not in style:
        return False
    has_time = re.search(r"\b\d{1,2}:\d{2}\b", text) is not None
    has_date = re.search(r"\b\d{1,2}\s+[A-Za-zА-Яа-я]+\s+\d{4}\b", text) is not None
    return bool(has_time or has_date)



def _remove_onenote_datetime(page_div: Tag) -> None:
    if not isinstance(page_div, Tag):
        return
    for div in list(page_div.find_all("div")):
        if _is_onenote_datetime_div(div):
            div.decompose()
    for p in list(page_div.find_all("p")):
        if _is_onenote_datetime_p(p):
            p.decompose()



def _normalize_onenote_page_div(page_div: Tag) -> None:
    if not isinstance(page_div, Tag):
        return

    def _has_visible_content(tag: Tag) -> bool:
        if tag.find(["img", "br"]):
            return True
        text = tag.get_text("", strip=False)
        # OneNote often stores spaces inside language/formatting spans; keep them.
        if tag.name == "span" and text:
            return True
        return bool((text or "").strip())

    def _has_non_whitespace_text(tag: Tag) -> bool:
        for child in tag.contents:
            if isinstance(child, NavigableString) and child.strip():
                return True
        return False

    def _remove_whitespace_text_nodes(tag: Tag) -> None:
        for child in list(tag.contents):
            if isinstance(child, NavigableString) and not child.strip():
                # Preserve explicit spaces that OneNote splits into separate nodes.
                child.replace_with(" ")

    for tag in list(page_div.find_all(["div", "p", "span"])):
        if _has_visible_content(tag):
            continue
        tag.decompose()

    for tag in page_div.find_all(True):
        _remove_whitespace_text_nodes(tag)

    def _merge_wrapper_style(curr_style: str, child_style: str) -> str:
        if not child_style:
            return curr_style
        if not curr_style:
            return child_style
        if child_style in curr_style:
            return curr_style
        return f"{curr_style.rstrip(';')}; {child_style}"

# --- 2) Нормализация OneNote inline-style -> семантические теги (устойчиво к санитизации/рендеру) ---

def _parse_style_attr(style: str) -> dict[str, str]:
    out: dict[str, str] = {}
    s = (style or "").strip()
    if not s:
        return out
    # простейший разбор "k:v; k2:v2"
    parts = [p.strip() for p in s.split(";") if p.strip()]
    for p in parts:
        if ":" not in p:
            continue
        k, v = p.split(":", 1)
        k = k.strip().lower()
        v = v.strip()
        if k:
            out[k] = v
    return out


def _build_style_attr(style_map: dict[str, str]) -> str:
    # соберём обратно; порядок не критичен
    items = []
    for k, v in style_map.items():
        v2 = (v or "").strip()
        if v2:
            items.append(f"{k}:{v2}")
    return "; ".join(items)

def normalize_onenote_rich_html(html_in: str) -> str:
    """
    OneNote часто кладёт форматирование в <span style="...">:
      - font-weight:bold
      - text-decoration:underline
      - color:#000099
      - background:lime / mso-highlight:lime

    Чтобы форматирование НЕ терялось:
      - превращаем bold/italic/underline/strike в <strong>/<em>/<u>/<s>
      - подсветку переносим в <mark style="background-color:...">
      - mso-highlight / background -> background-color
    """
    if not html_in:
        return ""

    soup = BeautifulSoup(html_in, "html.parser")
    root = soup.body or soup

    def _wrap_children(node: Tag, wrapper_name: str) -> None:
        wrapper = soup.new_tag(wrapper_name)
        for ch in list(node.contents):
            wrapper.append(ch.extract())
        node.append(wrapper)

    # Идём по span (и иногда p), потому что OneNote активно кладёт стили туда
    for node in root.find_all(["span", "p"]):
        style_map = _parse_style_attr(node.get("style", ""))
        if not style_map:
            continue

        # --- подсветка: mso-highlight / background -> background-color ---
        # OneNote: background:lime; mso-highlight:lime
        if "mso-highlight" in style_map and "background-color" not in style_map:
            style_map["background-color"] = style_map.get("mso-highlight", "")
        if "background" in style_map and "background-color" not in style_map:
            style_map["background-color"] = style_map.get("background", "")

        # --- семантика ---
        make_bold = style_map.get("font-weight", "").strip().lower() in {"bold", "700", "800", "900"}
        make_italic = style_map.get("font-style", "").strip().lower() == "italic"
        tdec = style_map.get("text-decoration", "").strip().lower()

        make_underline = "underline" in tdec
        make_strike = ("line-through" in tdec) or ("strike" in tdec)

        # Удаляем преобразованные свойства, чтобы не зависеть от CSS в дальнейшем
        if "font-weight" in style_map and make_bold:
            style_map.pop("font-weight", None)
        if "font-style" in style_map and make_italic:
            style_map.pop("font-style", None)
        if "text-decoration" in style_map and (make_underline or make_strike):
            style_map.pop("text-decoration", None)

        # --- подсветку переносим в <mark> внутрь узла ---
        bg = (style_map.get("background-color") or "").strip()
        if bg:
            # уберём background/background-color/mso-highlight из style, а визуализацию сделаем mark-ом
            style_map.pop("background-color", None)
            style_map.pop("background", None)
            style_map.pop("mso-highlight", None)

            mark = soup.new_tag("mark")
            mark["style"] = f"background-color:{bg};"

            # перенесём текущие children внутрь mark
            for ch in list(node.contents):
                mark.append(ch.extract())
            node.append(mark)

        # применим обёртки (важно: оборачиваем содержимое node целиком)
        # порядок: strike -> underline -> italic -> bold (можно любой, визуально ок)
        if make_strike:
            _wrap_children(node, "s")
        if make_underline:
            _wrap_children(node, "u")
        if make_italic:
            _wrap_children(node, "em")
        if make_bold:
            _wrap_children(node, "strong")

        # восстановим style без “проблемных” свойств
        new_style = _build_style_attr(style_map)
        if new_style:
            node["style"] = new_style
        else:
            node.attrs.pop("style", None)

    return str(root)



def _split_onenote_html_into_pages(soup: BeautifulSoup, filename: str):
    base_title = filename.rsplit(".", 1)[0]
    page_divs = soup.find_all("div", style=lambda v: v and "border-width:100%" in v)

    pages: list[tuple[str, str]] = []

    if not page_divs:
        title, body_html = html_to_body(str(soup), base_title)
        pages.append((title, body_html))
        return pages

    for idx, div in enumerate(page_divs, start=1):
        # Заголовок страницы OneNote обычно в первом <p> (часто внутри <a>)
        title_p = div.find("p")
        title_text = title_p.get_text(" ", strip=True) if title_p else ""
        if not title_text:
            title_text = f"{base_title} {idx}"

        # ✅ ВАЖНО: НЕ удаляем заголовок/дату (как раньше делал _strip_onenote_header),
        # а нормализуем их wrapper-div, чтобы не было "узкого" переноса.
        _unwrap_onenote_header_wrappers(div)
        _remove_onenote_datetime(div)

        # ✅ СОХРАНЯЕМ ПУСТЫЕ СТРОКИ ДО НОРМАЛИЗАЦИИ
        _onenote_keep_blank_lines_in_div(div)

        # дальше как раньше: нормализуем страницу от мусорных обёрток/пустот
        _normalize_onenote_page_div(div)


        body_html = str(div)
        pages.append((title_text, body_html))

    return pages

# def parse_mht_to_pages(
#     data: bytes,
#     filename: str,
#     *,
#     split_onenote_html_into_pages: Callable[[BeautifulSoup, str], list[tuple[str, str]]],
#     preserve_onenote_blank_paragraphs: Callable[[str], str],
#     strip_onenote_datetime_block: Callable[[str], str],
#     normalize_onenote_rich_html: Callable[[str], str],
#     sanitize_html_safe: Callable[[str], str],
# ) -> list[tuple[str, str]]:


def parse_mht_to_pages(
    data: bytes,
    filename: str,
) -> list[tuple[str, str]]:
    """
    Импорт OneNote .mht:
    - используем email (ОК) для корректного decode quoted-printable/base64
    - НЕ используем quopri и bleach
    - встраиваем ресурсы (img и пр.) как data: URL
    - делим на страницы
    - сохраняем форматирование: normalize_onenote_rich_html()
    - простая своя санитизация: sanitize_html_safe()
    """
    msg = email.message_from_bytes(data)

    html_bytes: bytes | None = None
    html_charset: str = "utf-8"
    resources: list[tuple[str, bytes, str | None, str | None]] = []

    # 1) достаём HTML и ресурсы
    for part in msg.walk():
        ctype = (part.get_content_type() or "").lower()
        if ctype == "text/html" and html_bytes is None:
            html_charset = part.get_content_charset() or "utf-8"
            # decode=True -> email сам снимет quoted-printable (=3D, =\r\n) и base64
            html_bytes = part.get_payload(decode=True) or b""
        else:
            cid = part.get("Content-ID")
            loc = part.get("Content-Location")
            payload = part.get_payload(decode=True) or b""
            if (cid or loc) and payload:
                resources.append((ctype, payload, cid, loc))

    if not html_bytes:
        raise ValueError("В .mht не найден HTML-контент")

    try:
        html_part = html_bytes.decode(html_charset, errors="replace")
    except Exception:
        html_part = html_bytes.decode("utf-8", errors="replace")

    def norm(val: str) -> str:
        v = urllib.parse.unquote(val or "").strip()
        v = v.replace("\\", "/")
        if v.lower().startswith("cid:"):
            v = "cid:" + v[4:]
        return v

    # 2) строим карту ресурсов: cid/location -> data:...
    src_map: dict[str, str] = {}
    for ctype, content, cid, loc in resources:
        data_url = f"data:{ctype};base64,{base64.b64encode(content).decode('ascii')}"
        if cid:
            cid_clean = cid.strip().strip("<>").strip()
            # OneNote может ссылаться как "cid:XXXX" или просто "XXXX"
            for key in (cid_clean, f"cid:{cid_clean}", f"CID:{cid_clean}", norm(cid_clean), norm(f"cid:{cid_clean}")):
                src_map[key] = data_url

        if loc:
            loc_clean = loc.strip().strip("<>").strip()
            normalized = norm(loc_clean)
            for key in (loc_clean, normalized, f"cid:{loc_clean}", f"cid:{normalized}", f"CID:{loc_clean}"):
                src_map[key] = data_url
            basename = os.path.basename(normalized)
            if basename:
                for key in (basename, f"cid:{basename}", f"CID:{basename}", norm(basename)):
                    src_map[key] = data_url

    soup = BeautifulSoup(html_part, "html.parser")

    # 3) подмена src на data-url (картинки и т.п.)
    for tag in soup.find_all(src=True):
        src_val = tag.get("src", "")
        lookup = norm(src_val)

        if lookup in src_map:
            tag["src"] = src_map[lookup]
            continue

        # иногда OneNote даёт только basename
        base = os.path.basename(lookup)
        if base in src_map:
            tag["src"] = src_map[base]

    # 4) ссылки в новую вкладку
    for a in soup.find_all("a", href=True):
        a["target"] = "_blank"
        a["rel"] = "noopener noreferrer"

    # 5) делим на страницы
    pages = _split_onenote_html_into_pages(soup, filename)

    safe_pages: list[tuple[str, str]] = []
    base_title = filename.rsplit(".", 1)[0]

    for title, body_html in pages:
        safe_title = (title or "").strip() or base_title

        # 0) сохранить пустые строки OneNote (&nbsp; -> <br>)
        body_html2 = preserve_onenote_blank_paragraphs(body_html or "")

        # 1) убрать дату и время из тела
        body_no_dt = strip_onenote_datetime_block(body_html2)

        # 1.1) ещё раз сохранить пустые строки (после вырезания дат/обёрток)
        body_no_dt = preserve_onenote_blank_paragraphs(body_no_dt)

        # 2) нормализация rich-formatting (bold/highlight/underline)
        body_norm = normalize_onenote_rich_html(body_no_dt)

        # 3) безопасный санитайзер
        safe_body = sanitize_html_safe(body_norm)

        safe_pages.append((safe_title, safe_body))
    return safe_pages


def _decode_text_bytes(data: bytes) -> str:
    if not data:
        return ""
    # BOM для UTF-16 / UTF-8
    if data.startswith(b"\xff\xfe") or data.startswith(b"\xfe\xff"):
        return data.decode("utf-16").replace("\x00", "")
    if data.startswith(b"\xef\xbb\xbf"):
        return data.decode("utf-8-sig").replace("\x00", "")

    # эвристика: если много нулевых байт, это UTF-16
    if data.count(b"\x00") > max(2, len(data) // 20):
        try:
            return data.decode("utf-16").replace("\x00", "")
        except UnicodeDecodeError:
            pass

    # выбор между utf-8 и cp1251
    try:
        text_utf8 = data.decode("utf-8")
        if "\ufffd" not in text_utf8:
            return text_utf8.replace("\x00", "")
    except UnicodeDecodeError:
        text_utf8 = None

    try:
        return data.decode("cp1251").replace("\x00", "")
    except UnicodeDecodeError:
        pass

    for enc in ("utf-8-sig", "utf-16", "cp1251", "latin-1"):
        try:
            return data.decode(enc).replace("\x00", "")
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").replace("\x00", "")


def _html_escape_preserve_breaks(text: str) -> str:
    safe = html.escape(text or "")
    return safe.replace("\r\n", "\n").replace("\r", "\n").replace("\n", "<br>")


def _style_dataframe_table(table_html: str) -> str:
    if not table_html:
        return ""
    table_html = table_html.replace(
        '<table border="1" class="dataframe">',
        '<table style="border-collapse:collapse;width:100%;">',
    )
    table_html = table_html.replace(
        "<th>",
        '<th style="border:1px solid #d0d4da;padding:4px 6px;text-align:left;">',
    )
    table_html = table_html.replace(
        "<td>",
        '<td style="border:1px solid #d0d4da;padding:4px 6px;">',
    )
    return table_html


def _docx_bytes_to_html(data: bytes) -> str:
    if not data:
        return ""
    if not zipfile.is_zipfile(io.BytesIO(data)):
        raise ValueError("Файл не является корректным .docx (zip).")

    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        doc = None
    parts: list[str] = []

    try:
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.shared import RGBColor  # type: ignore
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        CT_Tbl = CT_P = Table = Paragraph = None  # type: ignore
        RGBColor = None  # type: ignore
        qn = None  # type: ignore
        WD_ALIGN_PARAGRAPH = None  # type: ignore

    def _build_numbering_map(doc_obj) -> dict[int, dict[int, dict[str, str | int | bool]]]:
        """
        Собираем сведения о списках: numId -> ilvl -> {num_fmt, lvl_text, start, is_bullet}.
        Этого хватает для выборки ul/ol и стартового значения.
        """
        res: dict[int, dict[int, dict[str, str | int | bool]]] = {}
        try:
            num_part = getattr(doc_obj.part, "numbering_part", None)
            if not num_part:
                return res
            root = num_part.element
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            abstract_map: dict[int, dict[int, dict[str, str | int | bool]]] = {}
            for abs_el in root.findall(".//w:abstractNum", ns):
                abs_id_raw = abs_el.get(qn("w:abstractNumId")) if qn else abs_el.get("w:abstractNumId")
                try:
                    abs_id = int(abs_id_raw)
                except Exception:
                    continue
                lvl_map: dict[int, dict[str, str | int | bool]] = {}
                for lvl in abs_el.findall("w:lvl", ns):
                    ilvl_raw = lvl.get(qn("w:ilvl")) if qn else lvl.get("w:ilvl")
                    try:
                        ilvl = int(ilvl_raw or 0)
                    except Exception:
                        ilvl = 0
                    num_fmt_el = lvl.find("w:numFmt", ns)
                    lvl_text_el = lvl.find("w:lvlText", ns)
                    start_el = lvl.find("w:start", ns)
                    num_fmt = (num_fmt_el.get(qn("w:val")) if qn and num_fmt_el is not None else None) or (num_fmt_el.get("w:val") if num_fmt_el is not None else None) or (num_fmt_el.get("val") if num_fmt_el is not None else None)
                    lvl_text = (lvl_text_el.get(qn("w:val")) if qn and lvl_text_el is not None else None) or (lvl_text_el.get("w:val") if lvl_text_el is not None else None) or (lvl_text_el.get("val") if lvl_text_el is not None else None) or ""
                    try:
                        start_val = int(
                            (start_el.get(qn("w:val")) if qn and start_el is not None else None)
                            or (start_el.get("w:val") if start_el is not None else None)
                            or (start_el.get("val") if start_el is not None else None)
                            or 1
                        )
                    except Exception:
                        start_val = 1

                    fmt_lower = (num_fmt or "").lower()
                    lvl_map[ilvl] = {
                        "num_fmt": fmt_lower,
                        "lvl_text": lvl_text or "",
                        "start": start_val,
                        "is_bullet": fmt_lower == "bullet" or ("•" in (lvl_text or "")),
                        "bullet_char": lvl_text or "",
                    }
                abstract_map[abs_id] = lvl_map

            for num_el in root.findall(".//w:num", ns):
                num_id_raw = num_el.get(qn("w:numId")) if qn else num_el.get("w:numId")
                abs_el = num_el.find("w:abstractNumId", ns)
                abs_id_raw = (abs_el.get(qn("w:val")) if qn and abs_el is not None else None) or (abs_el.get("w:val") if abs_el is not None else None) or (abs_el.get("val") if abs_el is not None else None)
                try:
                    num_id = int(num_id_raw)
                    abs_id = int(abs_id_raw)
                except Exception:
                    continue
                if abs_id in abstract_map:
                    res[num_id] = abstract_map.get(abs_id, {})
            return res
        except Exception:
            return res

    def _rgb_to_css(val) -> str | None:
        if val is None:
            return None
        try:
            hexval = str(val)
            if len(hexval) == 6:
                return f"#{hexval}"
        except Exception:
            return None
        return None

    def _highlight_to_css(val) -> str | None:
        if val is None:
            return None
        name = ""
        try:
            if hasattr(val, "name"):
                name = str(val.name)
            else:
                name = str(val)
        except Exception:
            name = str(val)
        name = (name or "").upper().replace("WD_COLOR_INDEX.", "")
        highlight_map = {
            "YELLOW": "#fff59d",
            "BRIGHT_GREEN": "#b9f6ca",
            "GREEN": "#c8e6c9",
            "LIME": "#b9f6ca",
            "CYAN": "#b2ebf2",
            "MAGENTA": "#f8bbd0",
            "BLUE": "#bbdefb",
            "RED": "#ffcdd2",
            "PINK": "#f8bbd0",
            "GRAY": "#e0e0e0",
            "LIGHTGRAY": "#e0e0e0",
            "DARKGRAY": "#bdbdbd",
        }
        if name in highlight_map:
            return highlight_map[name]
        try:
            num = int(str(val))
            num_map = {
                4: "#fff59d",
                5: "#b9f6ca",
                11: "#e0e0e0",
                12: "#bdbdbd",
            }
            return num_map.get(num)
        except Exception:
            return None

    def _extract_images_from_run(run) -> list[str]:
        """Return list of <img> html strings for images inside run."""
        out: list[str] = []
        try:
            ns = {
                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            }
            blips = run.element.xpath(".//w:drawing//a:blip", namespaces=ns)
        except Exception:
            blips = []
        for blip in blips:
            r_id = blip.get(qn("r:embed")) if qn else blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if not r_id:
                continue
            try:
                img_part = run.part.related_parts[r_id]
                img_bytes = img_part.blob
                content_type = getattr(img_part, "content_type", "image/png") or "image/png"
                b64 = base64.b64encode(img_bytes).decode("ascii")
                out.append(f'<img src="data:{content_type};base64,{b64}" style="max-width:100%;height:auto;" />')
            except Exception:
                continue
        return out

    def _render_run(run) -> str:
        out_parts: list[str] = []

        out_parts.extend(_extract_images_from_run(run))

        text = html.escape(run.text or "")
        if text:
            text = text.replace("\n", "<br>")
            color_css = None
            bg_css = None
            try:
                if run.font is not None:
                    color_css = _rgb_to_css(getattr(run.font.color, "rgb", None))
                    highlight_val = getattr(run.font, "highlight_color", None)
                    if highlight_val and not bg_css:
                        bg_css = _highlight_to_css(highlight_val)
                    try:
                        shd = run._element.xpath(".//w:shd", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                        if shd:
                            fill = shd[0].get(qn("w:fill")) if qn else shd[0].get("w:fill")
                            if fill and fill.lower() != "auto":
                                bg_css = f"#{fill}"
                    except Exception:
                        pass
            except Exception:
                pass

            tags: list[str] = []
            if run.bold:
                tags.append("strong")
            if run.italic:
                tags.append("em")
            if run.underline:
                tags.append("u")
            if run.font is not None:
                if run.font.superscript:
                    tags.append("sup")
                elif run.font.subscript:
                    tags.append("sub")

            style_attrs = []
            if color_css:
                style_attrs.append(f"color:{color_css}")
            if bg_css:
                style_attrs.append(f"background:{bg_css}")
            if style_attrs:
                text = f'<span style="{";".join(style_attrs)}">{text}</span>'
            for tag in tags:
                text = f"<{tag}>{text}</{tag}>"
            out_parts.append(text)

        if not out_parts:
            return ""
        return "".join(out_parts)

    def _para_alignment(paragraph):
        """Возвращает выравнивание параграфа с учетом формата и стиля (включая базовые стили)."""
        try:
            if paragraph.alignment:
                return paragraph.alignment
            pf = getattr(paragraph, "paragraph_format", None)
            if pf and getattr(pf, "alignment", None):
                return pf.alignment
            style = getattr(paragraph, "style", None)
            seen = set()
            while style is not None and id(style) not in seen:
                seen.add(id(style))
                fmt = getattr(style, "paragraph_format", None)
                if fmt and getattr(fmt, "alignment", None):
                    return fmt.alignment
                style = getattr(style, "base_style", None)
        except Exception:
            return None
        return None

    def _render_paragraph(paragraph, wrap_tag: str | None = "p") -> str:
        text = "".join(_render_run(run) for run in paragraph.runs)
        if not text:
            text = _html_escape_preserve_breaks(paragraph.text or "")
        style_parts: list[str] = []
        align_name = ""
        try:
            align = _para_alignment(paragraph)
            align_map = {
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
                WD_ALIGN_PARAGRAPH.LEFT: "left",
            }
            align_name = align_map.get(align) or str(getattr(align, "name", "")).lower()
            if not align_name and align is not None:
                try:
                    align_val = int(align)
                    align_name = {0: "left", 1: "center", 2: "right", 3: "justify"}.get(align_val, "")
                except Exception:
                    pass
            if align_name:
                style_parts.append(f"text-align:{align_name} !important")
                if align_name in ("center", "right"):
                    style_parts.append("width:100%")
        except Exception:
            pass
        align_style = f' style="{";".join(style_parts)}"' if style_parts else ""
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if wrap_tag is not None and style_name.startswith("heading"):
            level = 1
            for num in "123456":
                if num in style_name:
                    level = int(num)
                    break
            inner = f"<h{level}{align_style}>{text}</h{level}>"
            if align_name in ("center", "right"):
                return f'<div style="text-align:{align_name} !important;">{inner}</div>'
            return inner
        if wrap_tag is None:
            return text or "<br>"
        body = f"<{wrap_tag}{align_style}>{text or '<br>'}</{wrap_tag}>"
        if align_name in ("center", "right"):
            return f'<div style="text-align:{align_name} !important;">{body}</div>'
        return body

    def _render_table(table) -> str:
        rows_html: list[str] = []

        def _cell_bg(cell) -> str | None:
            try:
                shd_list = cell._tc.xpath(".//w:shd")
                if shd_list:
                    shd = shd_list[0]
                    fill = shd.get(qn("w:fill")) if qn else None
                    theme_fill = shd.get(qn("w:themeFill")) if qn else None
                    if fill and fill.lower() != "auto":
                        return f"#{fill}"
                    if theme_fill:
                        theme_map = {
                            "accent1": "#d9e2f3",
                            "accent2": "#e2efda",
                            "accent3": "#fce4d6",
                            "accent4": "#f3f0f7",
                            "accent5": "#fdebd3",
                            "accent6": "#f8cbad",
                        }
                        return theme_map.get(theme_fill.lower(), "#f2f2f2")
            except Exception:
                return None
            return None

        for row in table.rows:
            cells_html: list[str] = []
            for cell in row.cells:
                para_html = [_render_paragraph(p) for p in cell.paragraphs]
                cell_text = "".join(para_html) or "<br>"
                bg = _cell_bg(cell)
                bg_style = f"background:{bg};" if bg else ""
                cells_html.append(
                    f'<td style="border:1px solid #d0d4da;padding:4px 6px;vertical-align:top;{bg_style}">{cell_text}</td>'
                )
            rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
        return (
            '<table style="border-collapse:collapse;width:100%;margin:6px 0;">'
            + "".join(rows_html)
            + "</table>"
        )

    if doc is not None:
        num_map = _build_numbering_map(doc)
        num_counters: dict[int, list[int | None]] = {}

        def _extract_num_from_numpr(num_pr) -> tuple[int, int] | None:
            if num_pr is None:
                return None
            try:
                num_id = num_pr.numId.val if getattr(num_pr, "numId", None) is not None else None
                ilvl = num_pr.ilvl.val if getattr(num_pr, "ilvl", None) is not None else 0
                if num_id is None:
                    return None
                return int(num_id), int(ilvl or 0)
            except Exception:
                return None

        def _style_num_info(style_obj) -> tuple[int, int] | None:
            seen_ids = set()
            cur = style_obj
            while cur is not None and id(cur) not in seen_ids:
                seen_ids.add(id(cur))
                try:
                    num_pr_el = cur.element.find(".//w:pPr/w:numPr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                    if num_pr_el is not None:
                        num_id_el = num_pr_el.find(".//w:numId", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                        ilvl_el = num_pr_el.find(".//w:ilvl", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                        num_id = num_id_el.get(qn("w:val")) if qn and num_id_el is not None else (num_id_el.get("w:val") if num_id_el is not None else None)
                        ilvl = ilvl_el.get(qn("w:val")) if qn and ilvl_el is not None else (ilvl_el.get("w:val") if ilvl_el is not None else None)
                        if num_id is not None:
                            try:
                                return int(num_id), int(ilvl or 0)
                            except Exception:
                                return None
                except Exception:
                    pass
                try:
                    cur = cur.base_style
                except Exception:
                    break
            return None

        def _get_num_info(paragraph) -> tuple[int, int] | None:
            try:
                info = _extract_num_from_numpr(getattr(getattr(paragraph._p, "pPr", None), "numPr", None))
                if info:
                    return info
            except Exception:
                pass
            try:
                if paragraph.style:
                    info = _style_num_info(paragraph.style)
                    if info:
                        return info
            except Exception:
                pass
            return None

        def _format_alpha(n: int, upper: bool = False) -> str:
            if n <= 0:
                return ""
            letters = []
            while n > 0:
                n -= 1
                letters.append(chr(ord("A" if upper else "a") + (n % 26)))
                n //= 26
            return "".join(reversed(letters))

        def _format_roman(n: int, upper: bool = False) -> str:
            if n <= 0:
                return ""
            vals = [
                (1000, "M"),
                (900, "CM"),
                (500, "D"),
                (400, "CD"),
                (100, "C"),
                (90, "XC"),
                (50, "L"),
                (40, "XL"),
                (10, "X"),
                (9, "IX"),
                (5, "V"),
                (4, "IV"),
                (1, "I"),
            ]
            out = []
            for v, sym in vals:
                while n >= v:
                    out.append(sym)
                    n -= v
            roman = "".join(out)
            return roman if upper else roman.lower()

        def _format_num(val: int, fmt: str | None) -> str:
            fmt_l = (fmt or "").lower()
            if fmt_l == "bullet":
                return "•"
            if fmt_l in ("lowerletter", "loweralpha", "lower_letter", "lower_alpha"):
                return _format_alpha(val, upper=False)
            if fmt_l in ("upperletter", "upperalpha", "upper_letter", "upper_alpha"):
                return _format_alpha(val, upper=True)
            if fmt_l in ("lowerroman", "roman"):
                return _format_roman(val, upper=False)
            if fmt_l in ("upperroman", "upper_roman"):
                return _format_roman(val, upper=True)
            return str(val)

        def _next_list_label(num_id: int, ilvl: int) -> str:
            lvl_def = num_map.get(num_id, {}).get(ilvl, {})
            counters = num_counters.setdefault(num_id, [])
            while len(counters) <= ilvl:
                counters.append(None)

            levels_map = num_map.get(num_id, {})
            was_none_current = counters[ilvl] is None
            for j in range(ilvl + 1):
                if counters[j] is None:
                    start_j = int(levels_map.get(j, {}).get("start") or 1)
                    counters[j] = start_j

            start_val = int(lvl_def.get("start") or 1)
            if was_none_current:
                counters[ilvl] = start_val
            else:
                counters[ilvl] = (counters[ilvl] or 0) + 1

            for j in range(ilvl + 1, len(counters)):
                counters[j] = None

            template = str(lvl_def.get("lvl_text") or "")
            label = template
            for i in range(1, 10):
                ph = f"%{i}"
                if ph in label:
                    lvl_fmt = num_map.get(num_id, {}).get(i - 1, {}).get("num_fmt") or lvl_def.get("num_fmt")
                    val = counters[i - 1] or 0
                    label = label.replace(ph, _format_num(val, lvl_fmt))

            if not label.strip() and lvl_def.get("is_bullet"):
                label = lvl_def.get("bullet_char") or "•"

            return label.strip()

        def _list_item_html(paragraph, num_id: int, ilvl: int) -> str:
            label = _next_list_label(num_id, ilvl)
            text_html = _render_paragraph(paragraph, wrap_tag=None)
            margin_px = 0
            try:
                indent = paragraph.paragraph_format.left_indent
                if indent:
                    margin_px = max(0, float(indent.pt) * 1.333)
                else:
                    margin_px = max(0, ilvl * 18)
            except Exception:
                margin_px = max(0, ilvl * 18)

            label_html = ""
            if label:
                label_html = f'<span style="margin-right:6px;white-space:nowrap;">{html.escape(label)}</span>'

            return f'<div style="margin:4px 0 4px {margin_px}px;padding-left:0;">{label_html}{text_html}</div>'

        if CT_P and CT_Tbl:
            for child in doc.element.body.iterchildren():
                if isinstance(child, CT_P):
                    para_obj = Paragraph(child, doc)
                    info = _get_num_info(para_obj)

                    if info:
                        num_id, ilvl = info
                        parts.append(_list_item_html(para_obj, num_id, ilvl))
                    else:
                        parts.append(_render_paragraph(para_obj))
                elif isinstance(child, CT_Tbl):
                    parts.append(_render_table(Table(child, doc)))
        else:
            for paragraph in doc.paragraphs:
                info = _get_num_info(paragraph)
                if info:
                    num_id, ilvl = info
                    parts.append(_list_item_html(paragraph, num_id, ilvl))
                else:
                    parts.append(_render_paragraph(paragraph))
            for table in doc.tables:
                parts.append(_render_table(table))

        return "\n".join([p for p in parts if p])

    # Fallback: manual XML parsing with media extraction and table formatting
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            xml_bytes = zf.read("word/document.xml")
            rels_xml = zf.read("word/_rels/document.xml.rels")
            media_files = {name: zf.read(name) for name in zf.namelist() if name.startswith("word/media/")}

        rels_map: dict[str, str] = {}
        try:
            rels_soup = BeautifulSoup(rels_xml, "xml")
            for rel in rels_soup.find_all("Relationship"):
                rid = rel.get("Id")
                target = rel.get("Target")
                if rid and target:
                    rels_map[rid] = target
        except Exception:
            pass

        soup = BeautifulSoup(xml_bytes, "xml")
        body = soup.find(["w:body", "body"])
        if body is None:
            raise ValueError("Не найден body в .docx")

        def _strip_ns(name: str | None) -> str:
            return (name or "").split(":")[-1]

        def _run_style(run) -> str:
            rpr = run.find(["w:rPr", "rPr"], recursive=False)
            styles: list[str] = []
            if rpr:
                if rpr.find(["w:b", "b"]):
                    styles.append("font-weight:700;")
                if rpr.find(["w:i", "i"]):
                    styles.append("font-style:italic;")
                u = rpr.find(["w:u", "u"])
                if u is not None:
                    u_val = u.get("w:val") or u.get("val")
                    if not u_val or u_val != "none":
                        styles.append("text-decoration:underline;")
                color = rpr.find(["w:color", "color"])
                if color and (color.get("w:val") or color.get("val")):
                    styles.append(f"color:#{color.get('w:val') or color.get('val')};")
                bg = rpr.find(["w:highlight", "highlight"])
                if bg and (bg.get("w:val") or bg.get("val")):
                    val = (bg.get("w:val") or bg.get("val") or "").upper()
                    hl_map = {
                        "YELLOW": "#fff59d",
                        "BRIGHT_GREEN": "#b9f6ca",
                        "GREEN": "#c8e6c9",
                        "LIME": "#b9f6ca",
                        "CYAN": "#b2ebf2",
                        "MAGENTA": "#f8bbd0",
                        "BLUE": "#bbdefb",
                        "RED": "#ffcdd2",
                        "PINK": "#f8bbd0",
                        "GRAY": "#e0e0e0",
                        "LIGHTGRAY": "#e0e0e0",
                        "DARKGRAY": "#bdbdbd",
                    }
                    if val in hl_map:
                        styles.append(f"background:{hl_map[val]};")
                shd = rpr.find("w:shd")
                if shd and (shd.get("w:fill") or shd.get("fill")):
                    fill = (shd.get("w:fill") or shd.get("fill") or "").strip()
                    if fill and fill.lower() != "auto":
                        styles.append(f"background:#{fill};")
                sz = rpr.find(["w:sz", "sz"])
                if sz is not None:
                    val = sz.get("w:val") or sz.get("val")
                    try:
                        pt = float(val) / 2.0
                        styles.append(f"font-size:{pt:.1f}pt;")
                    except Exception:
                        pass
            return "".join(styles)

        def _collect_run_text(run) -> str:
            buf: list[str] = []
            for node in run.descendants:
                name = _strip_ns(getattr(node, "name", None))
                if name == "t":
                    buf.append(html.escape(node.get_text() or ""))
                elif name == "tab":
                    buf.append("&emsp;")
                elif name == "br":
                    buf.append("<br>")
            return "".join(buf)

        def _render_run_xml(run) -> str:
            text = _collect_run_text(run)
            imgs_html: list[str] = []
            for bl in run.find_all("a:blip"):
                rid = bl.get("r:embed")
                if rid:
                    target = rels_map.get(rid, "")
                    if target.startswith("../"):
                        target = target.replace("../", "")
                    media_key = None
                    for name in media_files:
                        if target and name.endswith(target):
                            media_key = name
                            break
                        if rid in name:
                            media_key = name
                            break
                    if media_key:
                        try:
                            img_bytes = media_files[media_key]
                            b64 = base64.b64encode(img_bytes).decode("ascii")
                            ext = media_key.split(".")[-1].lower()
                            mime = f"image/{ext if ext != 'jpg' else 'jpeg'}"
                            imgs_html.append(f'<img src=\"data:{mime};base64,{b64}\" style=\"max-width:100%;height:auto;\" />')
                        except Exception:
                            continue
            style = _run_style(run)
            if style and text:
                text = f'<span style=\"{style}\">{text}</span>'
            return "".join(imgs_html) + text

        def _render_paragraph_xml(p) -> str:
            parts_local: list[str] = []
            for child in p.find_all(recursive=False):
                if _strip_ns(getattr(child, "name", None)) == "r":
                    parts_local.append(_render_run_xml(child))
            align_style = ""
            ppr = p.find(["w:pPr", "pPr"], recursive=False)
            if ppr:
                jc = ppr.find(["w:jc", "jc"])
                if jc is not None:
                    align = (jc.get("w:val") or jc.get("val") or "").lower()
                    if align in ("center", "right", "left", "both", "justify"):
                        align_style = f' style="text-align:{"justify" if align == "both" else align};"'
            return f"<p{align_style}>{''.join(parts_local) or '<br>'}</p>"

        def _render_cell(tc) -> tuple[str, str]:
            paras = tc.find_all(["w:p", "p"], recursive=False)
            bg_style = ""
            shd = tc.find("w:shd")
            if shd and (shd.get("w:fill") or shd.get("fill")):
                fill = (shd.get("w:fill") or shd.get("fill") or "").strip()
                if fill and fill.lower() != "auto":
                    bg_style = f"background:#{fill};"
            inner = "".join(_render_paragraph_xml(p) for p in paras)
            return inner or "<br>", bg_style

        def _render_tbl(tbl) -> str:
            rows_html: list[str] = []
            for tr in tbl.find_all("w:tr", recursive=False):
                cells_html: list[str] = []
                for tc in tr.find_all("w:tc", recursive=False):
                    cell_text, bg_style = _render_cell(tc)
                    cells_html.append(
                        f'<td style=\"border:1px solid #d0d4da;padding:4px 6px;vertical-align:top;{bg_style}\">{cell_text}</td>'
                    )
                rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
            return (
                '<table style=\"border-collapse:collapse;width:100%;margin:6px 0;\">'
                + "".join(rows_html)
                + "</table>"
            )

        for child in body.find_all(recursive=False):
            if child.name in ("w:p", "p"):
                parts.append(_render_paragraph_xml(child))
            elif child.name in ("w:tbl", "tbl"):
                parts.append(_render_tbl(child))

        return "\n".join([p for p in parts if p])
    except Exception as exc:
        raise ValueError(f"Не удалось разобрать .docx: {exc}") from exc


def _excel_bytes_to_html(data: bytes) -> str:
    if not data:
        return ""
    from openpyxl import load_workbook
    from openpyxl.styles.colors import COLOR_INDEX
    import xml.etree.ElementTree as ET

    def _apply_tint_to_hex(hexval: str, tint: float | None) -> str:
        if not hexval:
            return hexval
        base = hexval[-6:]
        if not base:
            return base
        if not tint:
            return base.upper()
        try:
            r = int(base[0:2], 16)
            g = int(base[2:4], 16)
            b = int(base[4:6], 16)
        except Exception:
            return base.upper()
        if tint > 0:
            r = int(r + (255 - r) * tint)
            g = int(g + (255 - g) * tint)
            b = int(b + (255 - b) * tint)
        else:
            r = int(r * (1 + tint))
            g = int(g * (1 + tint))
            b = int(b * (1 + tint))
        r = max(0, min(255, r))
        g = max(0, min(255, g))
        b = max(0, min(255, b))
        return f"{r:02X}{g:02X}{b:02X}"

    def _build_theme_colors(wb) -> dict[int, str]:
        theme_bytes = getattr(wb, "loaded_theme", None)
        if not theme_bytes:
            return {}
        try:
            root = ET.fromstring(theme_bytes)
            ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
            scheme = root.find(".//a:clrScheme", ns)
            if scheme is None:
                return {}
            names = [
                "dk1",
                "lt1",
                "dk2",
                "lt2",
                "accent1",
                "accent2",
                "accent3",
                "accent4",
                "accent5",
                "accent6",
                "hlink",
                "folHlink",
            ]
            theme_map: dict[int, str] = {}
            for idx, name in enumerate(names):
                node = scheme.find(f"a:{name}", ns)
                if node is None:
                    continue
                val = None
                srgb = node.find("a:srgbClr", ns)
                if srgb is not None:
                    val = srgb.attrib.get("val")
                if not val:
                    sys = node.find("a:sysClr", ns)
                    if sys is not None:
                        val = sys.attrib.get("lastClr")
                if val:
                    theme_map[idx] = val
            return theme_map
        except Exception:
            return {}

    def _color_to_hex(color, theme_map: dict[int, str]) -> str | None:
        if color is None:
            return None
        color_type = getattr(color, "type", None)
        if color_type == "theme":
            idx = getattr(color, "theme", None)
            if idx in (0, 1):
                return None
            if idx is not None:
                base = theme_map.get(idx)
                if base:
                    return _apply_tint_to_hex(base, getattr(color, "tint", None))
        if color_type == "indexed":
            idx = getattr(color, "indexed", None)
            if idx is not None and 0 <= idx < len(COLOR_INDEX):
                val = COLOR_INDEX[idx]
                if isinstance(val, str):
                    return val[-6:]
        rgb = getattr(color, "rgb", None)
        if isinstance(rgb, str) and rgb:
            return rgb[-6:]
        value = getattr(color, "value", None)
        if isinstance(value, str) and value:
            return value[-6:]
        return None

    def _cell_to_style(cell, theme_map: dict[int, str]) -> str:
        styles: list[str] = []
        if cell.font is not None:
            font_color = getattr(cell.font, "color", None)
            if cell.font.bold:
                styles.append("font-weight:700;")
            if cell.font.italic:
                styles.append("font-style:italic;")
            if cell.font.underline:
                styles.append("text-decoration:underline;")
            if font_color is not None:
                rgb = _color_to_hex(font_color, theme_map)
                skip_auto = False
                theme_idx = getattr(font_color, "theme", None)
                raw_rgb = getattr(font_color, "rgb", None)
                if theme_idx in (0, 1) and not raw_rgb:
                    skip_auto = True
                if rgb and not skip_auto:
                    styles.append(f"color: #{rgb};")
        if cell.fill is not None and cell.fill.patternType == "solid":
            rgb = _color_to_hex(cell.fill.fgColor, theme_map)
            if rgb:
                styles.append(f"background-color: #{rgb};")
        if cell.alignment is not None:
            align = cell.alignment.horizontal
            if align in ("left", "center", "right", "justify"):
                styles.append(f"text-align:{align};")
            v_align = cell.alignment.vertical
            if v_align in ("top", "center", "bottom"):
                styles.append(f"vertical-align:{v_align};")
        return "".join(styles)

    wb = load_workbook(io.BytesIO(data), data_only=True)
    theme_map = _build_theme_colors(wb)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        max_row = ws.max_row or 0
        max_col = ws.max_column or 0
        if max_row == 0 or max_col == 0:
            continue

        non_empty_counts: list[tuple[int, int]] = []
        for row_idx in range(1, max_row + 1):
            count = 0
            for col_idx in range(1, max_col + 1):
                val = ws.cell(row=row_idx, column=col_idx).value
                if val is not None and str(val).strip() != "":
                    count += 1
            non_empty_counts.append((row_idx, count))

        header_idx, header_count = max(non_empty_counts, key=lambda x: x[1])
        if header_count == 0:
            continue

        header_cells = [ws.cell(row=header_idx, column=col_idx) for col_idx in range(1, max_col + 1)]
        header_values = [str(c.value).strip() if c.value is not None else "" for c in header_cells]
        keep_cols = [idx for idx, val in enumerate(header_values, start=1) if val and not val.startswith("Unnamed")]
        if not keep_cols:
            continue

        rows_html: list[str] = []
        titles_html: list[str] = []
        for row_idx in range(1, header_idx):
            row_cells = [ws.cell(row=row_idx, column=col_idx) for col_idx in keep_cols]
            if all((cell.value is None or str(cell.value).strip() == "") for cell in row_cells):
                continue
            spans: list[str] = []
            for cell in row_cells:
                if cell.value is None or str(cell.value).strip() == "":
                    continue
                text = html.escape(str(cell.value).strip())
                style = _cell_to_style(cell, theme_map)
                style_attr = f' style="{style}"' if style else ""
                spans.append(f"<span{style_attr}>{text}</span>")
            if spans:
                titles_html.append('<div style="margin:0 0 6px 0;">' + " ".join(spans) + "</div>")
        header_html: list[str] = []
        for col_idx in keep_cols:
            cell = ws.cell(row=header_idx, column=col_idx)
            text = html.escape(str(cell.value).strip() if cell.value is not None else "")
            style = _cell_to_style(cell, theme_map)
            style_attr = f' style="border:1px solid #d0d4da;padding:4px 6px;{style}"'
            header_html.append(f"<th{style_attr}>{text}</th>")
        rows_html.append("<tr>" + "".join(header_html) + "</tr>")

        for row_idx in range(header_idx + 1, max_row + 1):
            row_cells = [ws.cell(row=row_idx, column=col_idx) for col_idx in keep_cols]
            if all((cell.value is None or str(cell.value).strip() == "") for cell in row_cells):
                continue
            cell_html: list[str] = []
            for cell in row_cells:
                text = html.escape(str(cell.value) if cell.value is not None else "")
                style = _cell_to_style(cell, theme_map)
                style_attr = f' style="border:1px solid #d0d4da;padding:4px 6px;{style}"'
                cell_html.append(f"<td{style_attr}>{text}</td>")
            rows_html.append("<tr>" + "".join(cell_html) + "</tr>")

        table_html = (
            '<table style="border-collapse:collapse;width:100%;">'
            + "".join(rows_html)
            + "</table>"
        )
        parts.extend(titles_html)
        parts.append(table_html)
    return "\n".join(parts)


def _csv_bytes_to_html(data: bytes) -> str:
    if not data:
        return ""
    try:
        df = pd.read_csv(io.BytesIO(data), dtype=str, sep=None, engine="python")
    except Exception:
        df = pd.read_csv(io.BytesIO(data), dtype=str, encoding="cp1251")
    df = df.fillna("")
    return _style_dataframe_table(df.to_html(index=False, escape=True))


def _text_bytes_to_html(data: bytes) -> str:
    text = _decode_text_bytes(data)
    safe = html.escape(text or "")
    return '<pre style="white-space:pre-wrap;font-family:monospace;">' + safe + "</pre>"


def _png_bytes_to_html(data: bytes) -> str:
    if not data:
        return ""
    b64 = base64.b64encode(data).decode("ascii")
    return (
        '<div style="text-align:center;">'
        f'<img src="data:image/png;base64,{b64}" style="max-width:100%;height:auto;" />'
        "</div>"
    )
