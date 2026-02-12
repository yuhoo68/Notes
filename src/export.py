# vers 1.04
from __future__ import annotations

import base64
import io
import logging
import os
import re
import tempfile
from dataclasses import dataclass
from typing import Any, Callable

import pandas as pd
import requests
import streamlit as st
import streamlit.components.v1 as components
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

import config
from src.mail import send_mail

logger = logging.getLogger("notes_export")


@dataclass(frozen=True)
class ExportDeps:
    list_email_recipients: Callable[[], pd.DataFrame]
    get_user_signature: Callable[[str], tuple[str, str]]
    get_page_attachments: Callable[[int], pd.DataFrame]
    get_attachment_file: Callable[[int], tuple[bytes, str, str] | None]
    run_execute: Callable[[str], Any]
    run_scalar: Callable[[str], Any]
    add_event_log: Callable[..., Any]
    get_nb_section_id: Callable[[int], tuple[int, int]]
    build_page_deeplink: Callable[[int], str]
    escape: Callable[[str], str]
    email_recipients_table: str


def _safe_filename(title: str, ext: str) -> str:
    base = (title or "page").strip()
    base = re.sub(r"[^\w\-. ]+", "_", base)
    if not base:
        base = "page"
    return f"{base}.{ext}"


def _norm_newlines_global(s: str) -> str:
    return (s or "").replace("\r\n", "\n").replace("\r", "\n")


def _indent_level_with_root(tag: Tag, root: Tag) -> int:
    cur: Tag | None = tag
    while isinstance(cur, Tag):
        classes = cur.get("class") or []
        for c in classes:
            m = re.match(r"ql-indent-(\d+)", str(c))
            if m:
                return int(m.group(1))
        if cur is root:
            break
        cur = cur.parent if isinstance(cur.parent, Tag) else None
    return 0


def _has_desc_block(tag: Tag, block_tags: set[str]) -> bool:
    for child in tag.find_all(list(block_tags), recursive=True):
        if child is not tag:
            return True
    return False


def _text_of_leaf(tag: Tag, _norm_newlines_global) -> str:
    parts: list[str] = []
    for elem in tag.descendants:
        if isinstance(elem, NavigableString):
            parts.append(str(elem))
    raw = "".join(parts)
    raw = _norm_newlines_global(raw).replace("\xa0", " ")
    return "\n".join(line.rstrip() for line in raw.split("\n"))


def _format_file_size(size: int | float | None) -> str:
    if not size or size <= 0:
        return ""
    units = ["B", "KB", "MB", "GB"]
    value = float(size)
    for unit in units:
        if value < 1024 or unit == units[-1]:
            return f"{value:.1f} {unit}"
        value /= 1024
    return f"{value:.1f} B"


def _merge_email_lists(*lists: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for items in lists:
        for item in items:
            if not item or item in seen:
                continue
            seen.add(item)
            out.append(item)
    return out


def _is_valid_email(value: str) -> bool:
    val = (value or "").strip()
    if not val:
        return False
    return re.match(r"^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$", val) is not None


def _sanitize_attachment_filename(name: str) -> str:
    base = os.path.basename(name or "attachment")
    base = _transliterate_ru_filename(base)
    base = re.sub(r"[^A-Za-z0-9_.-]+", "_", base)
    base = re.sub(r"_+", "_", base).strip("._- ")
    return base or "attachment"


def _transliterate_ru_filename(text: str) -> str:
    if not text:
        return ""

    m = {
        "а": "a",
        "б": "b",
        "в": "v",
        "г": "g",
        "д": "d",
        "е": "e",
        "ё": "yo",
        "ж": "zh",
        "з": "z",
        "и": "i",
        "й": "y",
        "к": "k",
        "л": "l",
        "м": "m",
        "н": "n",
        "о": "o",
        "п": "p",
        "р": "r",
        "с": "s",
        "т": "t",
        "у": "u",
        "ф": "f",
        "х": "kh",
        "ц": "ts",
        "ч": "ch",
        "ш": "sh",
        "щ": "shch",
        "ъ": "",
        "ы": "y",
        "ь": "",
        "э": "e",
        "ю": "yu",
        "я": "ya",
    }
    out: list[str] = []
    for ch in text:
        low = ch.lower()
        if low in m:
            rep = m[low]
            if ch.isupper() and rep:
                rep = rep[0].upper() + rep[1:]
            out.append(rep)
        else:
            out.append(ch)
    return "".join(out)


def export_html_to_docx_bytes(html: str, title: str) -> io.BytesIO:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

    _XML_BAD = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")

    def _xml_safe(s: str) -> str:
        if not s:
            return ""
        s = s.replace("\xa0", " ")
        return _XML_BAD.sub("", s)



    def _norm_title(t: str) -> str:
        # убираем переносы и лишние пробелы, чтобы "Test 5" не разъезжал на 2 строки
        return re.sub(r"\s+", " ", (t or "").replace("\xa0", " ")).strip()

    def _set_paragraph_spacing(p):
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _add_hyperlink(paragraph, url: str, text: str, bold=False, italic=False, underline=True):
        """
        Создаёт настоящую hyperlink-ссылку в docx (а не просто текст).
        """
        if not url:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            run.underline = underline
            return run

        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Стиль Hyperlink (если есть) + подчёркивание
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)

        if bold:
            b = OxmlElement("w:b")
            rPr.append(b)
        if italic:
            i = OxmlElement("w:i")
            rPr.append(i)

        # underline
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single" if underline else "none")
        rPr.append(u)

        # (опционально) цвет как у ссылок
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "0000FF")
        rPr.append(c)

        r.append(rPr)

        t = OxmlElement("w:t")
        t.text = text
        r.append(t)

        hyperlink.append(r)
        paragraph._p.append(hyperlink)
        return hyperlink

    doc = Document()

    # Нормальные интервалы
    normal_style = doc.styles["Normal"]
    nf = normal_style.paragraph_format
    nf.space_before = Pt(0)
    nf.space_after = Pt(4)
    nf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    safe_title = _xml_safe(_norm_title(title))

    # Не добавляем автозаголовок в DOCX (по требованию — без тайтла наверху)
    add_heading = False


    soup = BeautifulSoup(html or "", "html.parser")
    body = soup.body or soup

    # <br> превращаем в безопасный маркер (НЕ \u0000)
    BR_TOKEN = "\uE000"
    for br in body.find_all("br"):
        br.replace_with(BR_TOKEN)

    def _normalize_text_for_docx(s: str) -> str:
        """
        Главный фикс для вашего кейса:
        OneNote/Word вставляет \n прямо в текст внутри <span>/<p>.
        В docx это превращается в переносы строк -> "по слову на строку".
        """
        if not s:
            return ""
        s = _xml_safe(s)

        # сохраняем BR_TOKEN, остальное \n/\r/\t превращаем в пробелы
        s = s.replace(BR_TOKEN, "\uE001")
        s = s.replace("\r\n", " ").replace("\r", " ").replace("\n", " ").replace("\t", " ")
        # сжимаем “служебные” пробелы от переносов
        s = re.sub(r"[ ]{2,}", " ", s)
        s = s.replace("\uE001", BR_TOKEN)
        return s

    def _parse_color(val: str) -> tuple[int, int, int] | None:
        if not val:
            return None
        raw = val.strip().lower()
        # Hex
        m_hex = re.match(r"#([0-9a-f]{6})", raw)
        if m_hex:
            hx = m_hex.group(1)
            return int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
        # rgb(...)
        m_rgb = re.match(r"rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", raw)
        if m_rgb:
            return int(m_rgb.group(1)), int(m_rgb.group(2)), int(m_rgb.group(3))

        names = {
            "black": (0, 0, 0),
            "white": (255, 255, 255),
            "gray": (128, 128, 128),
            "grey": (128, 128, 128),
            "lightgray": (211, 211, 211),
            "lightgrey": (211, 211, 211),
            "yellow": (255, 255, 0),
            "lime": (0, 255, 0),
            "green": (0, 128, 0),
            "red": (255, 0, 0),
            "blue": (0, 0, 255),
        }
        return names.get(raw)


    def _parse_style(
        style: str,
    ) -> tuple[
        float | None,
        tuple[int, int, int] | None,
        tuple[int, int, int] | None,
        bool | None,
        bool | None,
        bool | None,
        bool | None,
    ]:
        """
        Берём минимум из inline-style OneNote:
        font-size:10.0pt / color:gray / color:#2E74B5
        """
        if not style:
            return None, None, None, None, None, None, None
        st = style.lower()

        fs = None
        m = re.search(r"font-size\s*:\s*([0-9.]+)\s*pt", st)
        if m:
            try:
                fs = float(m.group(1))
            except Exception:
                fs = None

        col = None
        m2 = re.search(r"(?:^|;)\s*color\s*:\s*([^;]+)", st)
        if m2:
            col = _parse_color(m2.group(1))

        bg = None
        m5 = re.search(r"(?:^|;)\s*(background-color|background|mso-highlight)\s*:\s*([^;]+)", st)
        if m5:
            bg = _parse_color(m5.group(2))

        bold = None
        italic = None
        underline = None
        strike = None
        m8 = re.search(r"font-weight\s*:\s*([0-9]+|bold)", st)
        if m8:
            val = m8.group(1)
            bold = val == "bold" or (val.isdigit() and int(val) >= 600)
        if "font-style:italic" in st:
            italic = True
        if "text-decoration" in st:
            underline = "underline" in st
            strike = "line-through" in st or "strike" in st

        return fs, col, bg, bold, italic, underline, strike

    def _parse_paragraph_style(style: str) -> dict:
        if not style:
            return {}
        st = style.lower()
        out: dict[str, str] = {}
        m = re.search(r"text-align\s*:\s*([a-z]+)", st)
        if m:
            out["align"] = m.group(1)
        m = re.search(r"margin-left\s*:\s*([0-9.]+)\s*px", st)
        if m:
            out["margin_left_px"] = m.group(1)
        m = re.search(r"text-indent\s*:\s*([0-9.]+)\s*px", st)
        if m:
            out["text_indent_px"] = m.group(1)
        m = re.search(r"margin-top\s*:\s*([0-9.]+)\s*px", st)
        if m:
            out["margin_top_px"] = m.group(1)
        m = re.search(r"margin-bottom\s*:\s*([0-9.]+)\s*px", st)
        if m:
            out["margin_bottom_px"] = m.group(1)
        return out

    def _px_to_pt(px: str | None) -> float | None:
        if not px:
            return None
        try:
            return float(px) / 1.333
        except Exception:
            return None

    def _set_run_shading(run, rgb: tuple[int, int, int]) -> None:
        r, g, b = rgb
        fill = f"{r:02X}{g:02X}{b:02X}"
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        rPr.append(shd)

    def _apply_run_background(run, rgb: tuple[int, int, int]) -> None:
        highlight_map = {
            (255, 245, 157): WD_COLOR_INDEX.YELLOW,
            (165, 214, 167): WD_COLOR_INDEX.BRIGHT_GREEN,
            (0, 255, 0): WD_COLOR_INDEX.BRIGHT_GREEN,
            (255, 255, 0): WD_COLOR_INDEX.YELLOW,
            (144, 202, 249): WD_COLOR_INDEX.BLUE,
            (239, 154, 154): WD_COLOR_INDEX.RED,
            (206, 147, 216): WD_COLOR_INDEX.PINK,
            (224, 224, 224): WD_COLOR_INDEX.GRAY_25,
            (189, 189, 189): WD_COLOR_INDEX.GRAY_50,
        }
        highlight = highlight_map.get(rgb)
        if highlight is not None:
            run.font.highlight_color = highlight
        else:
            _set_run_shading(run, rgb)


    def _apply_run_format(run, bold=False, italic=False, underline=False, strike=False):
        run.bold = bool(bold)
        run.italic = bool(italic)
        run.underline = bool(underline)
        run.font.strike = bool(strike)



    def _walk_inline(
        node,
        paragraph,
        bold=False,
        italic=False,
        underline=False,
        strike=False,
        font_size_pt: float | None = None,
        font_rgb: tuple[int, int, int] | None = None,
        bg_rgb: tuple[int, int, int] | None = None,
    ):
        """
        Рекурсивный обход inline-узлов:
        - фикс OneNote переносов (\n -> пробел)
        - поддержка bold/italic/underline/strike
        - минимальная поддержка font-size/color из inline-style
        - ссылки -> real hyperlink
        """
        if isinstance(node, NavigableString):
            text = _normalize_text_for_docx(str(node))
            if not text:
                return

            parts = text.split(BR_TOKEN)
            for i, part in enumerate(parts):
                if part:
                    run = paragraph.add_run(part)
                    _apply_run_format(run, bold, italic, underline, strike)
                    if font_size_pt:
                        run.font.size = Pt(font_size_pt)
                    if font_rgb:
                        run.font.color.rgb = RGBColor(*font_rgb)
                    elif bg_rgb:
                        # Если есть фон, а цвет текста не задан, принудительно делаем текст чёрным для контраста.
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    if bg_rgb:
                        _apply_run_background(run, bg_rgb)
                if i < len(parts) - 1:
                    paragraph.add_run().add_break()
            return

        if not isinstance(node, Tag):
            return

        name = (node.name or "").lower()

        # перенос форматирования
        nb, ni, nu, ns = bold, italic, underline, strike
        if name in ("b", "strong"):
            nb = True
        if name in ("i", "em"):
            ni = True
        if name == "u":
            nu = True
        if name in ("s", "strike", "del"):
            ns = True

        # перенос стиля текста (font-size / color)
        cur_size, cur_col, cur_bg = font_size_pt, font_rgb, bg_rgb
        fs2, col2, bg2, b2, i2, u2, s2 = _parse_style(node.get("style") or "")
        if fs2:
            cur_size = fs2
        if col2:
            cur_col = col2
        if bg2:
            cur_bg = bg2
        if b2 is True:
            nb = True
        if i2 is True:
            ni = True
        if u2 is True:
            nu = True
        if s2 is True:
            ns = True

        # картинки
        if name == "img":
            src = (node.get("src") or "").strip()
            if not src:
                return
            try:
                img_bytes = None
                if src.startswith("data:image"):
                    _, b64data = src.split(",", 1)
                    img_bytes = base64.b64decode(b64data)
                elif src.startswith("http://") or src.startswith("https://"):
                    resp = requests.get(src, timeout=8)
                    resp.raise_for_status()
                    img_bytes = resp.content

                if img_bytes:
                    img_stream = io.BytesIO(img_bytes)
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(5))
            except Exception as e:
                logger.warning("Не удалось вставить изображение в DOCX: %s", e)
            return

        # гиперссылка
        if name == "a":
            href = (node.get("href") or "").strip()
            link_text = _normalize_text_for_docx(node.get_text("", strip=False))
            link_text = link_text.strip()
            if link_text:
                _add_hyperlink(paragraph, href, link_text, bold=nb, italic=ni, underline=True)
            return

        for child in node.children:
            _walk_inline(child, paragraph, nb, ni, nu, ns, cur_size, cur_col, cur_bg)




    def _apply_paragraph_style(p, style: str | None) -> None:
        styles = _parse_paragraph_style(style or "")
        if not styles:
            return
        if "align" in styles:
            align = styles["align"]
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            if align in align_map:
                p.alignment = align_map[align]
        pf = p.paragraph_format
        ml_pt = _px_to_pt(styles.get("margin_left_px"))
        if ml_pt is not None:
            pf.left_indent = Pt(ml_pt)
        ti_pt = _px_to_pt(styles.get("text_indent_px"))
        if ti_pt is not None:
            pf.first_line_indent = Pt(ti_pt)
        mt_pt = _px_to_pt(styles.get("margin_top_px"))
        if mt_pt is not None:
            pf.space_before = Pt(mt_pt)
        mb_pt = _px_to_pt(styles.get("margin_bottom_px"))
        if mb_pt is not None:
            pf.space_after = Pt(mb_pt)


    def _add_paragraph_from_tag(tag: Tag, style: str | None = None):
        # Пустые абзацы из <p><br/></p> — это осознанные пустые строки. Сохраняем их как пустые параграфы.
        raw_txt = tag.get_text("", strip=False)
        if isinstance(raw_txt, str):
            br_only = raw_txt.replace("\xa0", "").strip() == BR_TOKEN
            txt_no_br = raw_txt.replace(BR_TOKEN, "").replace("\xa0", "").strip()
            if not txt_no_br and not tag.find(["img", "table"]):
                # количество пустых строк = количество <br> внутри параграфа (минимум 1)
                br_count = raw_txt.count(BR_TOKEN)
                for _ in range(max(1, br_count)):
                    p_empty = doc.add_paragraph(style=style)
                    _set_paragraph_spacing(p_empty)
                return

        p = doc.add_paragraph(style=style)
        _set_paragraph_spacing(p)
        _apply_paragraph_style(p, tag.get("style") or "")
        for child in tag.children:
            _walk_inline(child, p)

    def _add_pre(tag: Tag):
        # pre оставляем с переносами строк
        p = doc.add_paragraph()
        _set_paragraph_spacing(p)
        run = p.add_run(tag.get_text().replace("\xa0", " "))
        run.font.name = "Consolas"
        # чуть уменьшить, чтобы было похоже на код
        run.font.size = Pt(10)




    def _add_table_from_tag(table_tag: Tag):
        rows = table_tag.find_all("tr")
        if not rows:
            return
        max_cols = 0
        row_cells = []
        for tr in rows:
            cells = tr.find_all(["td", "th"], recursive=False)
            row_cells.append(cells)
            max_cols = max(max_cols, len(cells))

        t = doc.add_table(rows=len(rows), cols=max_cols)
        t.style = "Table Grid"
        for i, cells in enumerate(row_cells):
            for j, cell_tag in enumerate(cells):
                cell = t.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                _set_paragraph_spacing(p)
                # apply cell background if present
                style_attr = (cell_tag.get("style") or "").lower()
                bg_val = ""
                for part in style_attr.split(";"):
                    if ":" not in part:
                        continue
                    k, v = part.split(":", 1)
                    if k.strip() in ("background", "background-color"):
                        bg_val = v.strip()
                        break
                if bg_val:
                    try:
                        tc_pr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        hex_color = bg_val.replace("#", "").strip()
                        shd.set(qn("w:fill"), hex_color)
                        tc_pr.append(shd)
                    except Exception:
                        pass
                for child in cell_tag.children:
                    _walk_inline(child, p)

    block_tags = {
        "p",
        "div",
        "table",
        "ul",
        "ol",
        "pre",
        "blockquote",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
    }

    def _handle_block(node):
        if isinstance(node, NavigableString):
            # игнорируем пустые межтеговые переносы
            if not str(node).strip():
                return
            p = doc.add_paragraph(str(node).strip())
            _set_paragraph_spacing(p)
            return

        if not isinstance(node, Tag):
            return

        name = (node.name or "").lower()

        if name in ("div", "span") and not _has_desc_block(node, block_tags):
            _add_paragraph_from_tag(node)
            return

        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(name[1])
            _add_paragraph_from_tag(node, style=f"Heading {min(level, 3)}")
            return

        if name == "p":
            _add_paragraph_from_tag(node)
            return

        if name == "pre":
            _add_pre(node)
            return

        if name == "blockquote":
            _add_paragraph_from_tag(node, style="Intense Quote")
            return
        if name in ("ul", "ol"):
            is_ul = (name == "ul")

            children = [c for c in node.children if isinstance(c, Tag)]

            # If a list has no <li> (broken HTML), use any child tags as items to avoid losing text
            items = [c for c in children if (c.name or "").lower() == "li"]
            if not items:
                for child in node.children:
                    _handle_block(child)
                return

            num = 1
            for child in children:
                if not isinstance(child, Tag):
                    continue
                child_name = (child.name or "").lower()
                if child_name != "li":
                    _handle_block(child)
                    continue

                li = child
                p = doc.add_paragraph()
                _set_paragraph_spacing(p)
                # apply list and item styles (indent/alignment)
                _apply_paragraph_style(p, node.get("style") or "")
                _apply_paragraph_style(p, li.get("style") or "")

                if is_ul:
                    p.add_run("- ")
                else:
                    v = li.get("value")
                    if v is not None:
                        try:
                            num = int(str(v))
                        except Exception:
                            pass
                    p.add_run(f"{num}. ")
                    num += 1

                for child in li.children:
                    if not isinstance(child, Tag):
                        _walk_inline(child, p)
                        continue
                    cname = (child.name or "").lower()
                    if cname == "table":
                        _add_table_from_tag(child)
                    elif cname in ("ol", "ul"):
                        _handle_block(child)
                    else:
                        _walk_inline(child, p)
            return

            num = 1
            for li in items:
                if not isinstance(li, Tag):
                    continue

                p = doc.add_paragraph()
                _set_paragraph_spacing(p)

                if is_ul:
                    p.add_run("- ")
                else:
                    v = li.get("value")
                    if v is not None:
                        try:
                            num = int(str(v))
                        except Exception:
                            pass
                    p.add_run(f"{num}. ")
                    num += 1

                for child in li.children:
                    if isinstance(child, Tag) and (child.name or "").lower() == "table":
                        _add_table_from_tag(child)
                    else:
                        _walk_inline(child, p)
            return


        if name == "table":
            _add_table_from_tag(node)
            return

        # контейнеры — раскрываем внутрь
        for child in node.children:
            _handle_block(child)

    for child in body.children:
        _handle_block(child)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# =========================
# FIX: сохраняем отступы/пробелы для .sql экспорта + корректный текст для буфера
# =========================
def _html_to_plain_preserving_layout(html: str, indent_spaces: int = 4) -> str:
    if not html:
        return ""

    soup = BeautifulSoup(html, "html.parser")

    for br in soup.find_all("br"):
        br.replace_with("\n")

    body = soup.body or soup

    block_tags = {"p", "div", "li", "pre", "h1", "h2", "h3", "h4", "h5", "h6"}
    out_lines: list[str] = []

    blocks_all = body.find_all(list(block_tags))
    # берем "листовые" блоки: те, у которых нет вложенных блоков
    blocks = [t for t in blocks_all if not _has_desc_block(t, block_tags)]
    if blocks:
        for el in blocks:
            txt = _text_of_leaf(el, _norm_newlines_global)

            if txt.replace("\n", "").strip() == "":
                out_lines.append("")
                continue

            lvl = _indent_level_with_root(el, body)
            prefix = (" " * (lvl * indent_spaces)) if lvl > 0 else ""

            for line in txt.split("\n"):
                if line == "":
                    out_lines.append("")
                else:
                    out_lines.append(prefix + line)
    else:
        txt = _text_of_leaf(body, _norm_newlines_global)
        out_lines.extend(txt.split("\n"))

    result = "\n".join(out_lines)
    result = _norm_newlines_global(result)
    result = re.sub(r"\n{3,}", "\n\n", result).rstrip()
    return result


def _html_to_plain_for_email(html: str, indent_spaces: int = 4) -> str:
    """
    Plain-text для eMail:
    - нормализует "служебные" переносы OneNote/Word, чтобы не было "по слову на строку"
    - сохраняет ссылки из HTML как: "Текст ссылки - URL"
    - <br> -> перенос строки, <pre> -> сохраняем переносы
    """
    if not html:
        return ""

    soup = BeautifulSoup(html, "html.parser")

    for bad in soup.find_all(["script", "style"]):
        bad.decompose()

    # безопасный маркер переноса (не управляющий символ)
    br_token = "\uE000"
    for br in soup.find_all("br"):
        br.replace_with(br_token)

    body = soup.body or soup
    block_tags = {"p", "li", "pre", "h1", "h2", "h3", "h4", "h5", "h6"}

    def _normalize_inline_text(s: str) -> str:
        """
        Нормализация текста внутри обычных блоков:
        - br_token пока оставляем (потом превратим в \n)
        - любые \n/\t из исходного HTML считаем пробелами
        """
        s = _norm_newlines_global(s).replace("\xa0", " ")
        s = s.replace(br_token, "\n")  # реальные переносы от <br>
        # все прочие переносы/табуляции схлопываем
        s = s.replace("\n", "\n")  # уже "настоящие" переносы
        # но часто OneNote приносит \n как текстовые переносы - они уже схлопнутся ниже
        # (между словами будет пробел)
        lines = s.split("\n")
        cleaned_lines = []
        for line in lines:
            line = line.replace("\t", " ")
            line = re.sub(r"[ \t\f\v]+", " ", line).strip()
            cleaned_lines.append(line)
        # сохраняем пустые строки как пустые
        return "\n".join(cleaned_lines)

    def _replace_links(root: Tag) -> None:
        """
        <a href="URL">Text</a> -> Text - URL
        Делает это ДО get_text(), чтобы URL не потерялся.
        """
        for a in list(root.find_all("a", href=True)):
            href = (a.get("href") or "").strip()
            if not href:
                continue

            # Берём видимый текст ссылки (с учётом вложенных span/strong и т.д.)
            text = a.get_text(" ", strip=True)
            text = re.sub(r"[ \t\f\v]+", " ", (text or "").replace("\xa0", " ")).strip()

            # Если текста нет — используем URL как текст
            if not text:
                repl = href
            else:
                # если текст уже равен URL — не дублируем
                if text.strip().lower() == href.strip().lower():
                    repl = href
                else:
                    repl = f"{text} - {href}"

            a.replace_with(NavigableString(repl))

    def _text_of(tag: Tag) -> str:
        # pre: сохраняем переносы строк
        if tag.name and tag.name.lower() == "pre":
            # ссылки внутри pre тоже можно сохранить (редко, но бывает)
            tag_copy = BeautifulSoup(str(tag), "html.parser")
            pre_tag = tag_copy.find("pre") or tag_copy
            _replace_links(pre_tag)

            raw = pre_tag.get_text()
            raw = raw.replace("\xa0", " ")
            raw = raw.replace(br_token, "\n")
            raw = _norm_newlines_global(raw)
            return "\n".join(line.rstrip() for line in raw.split("\n")).rstrip()

        # обычный блок: подменяем ссылки на "Text - URL"
        tag_copy = BeautifulSoup(str(tag), "html.parser")
        root = tag_copy.find(tag.name) or tag_copy
        _replace_links(root)

        raw = root.get_text(separator=" ", strip=False)
        raw = raw.replace("\xa0", " ")
        raw = raw.replace(br_token, "\n")
        raw = _norm_newlines_global(raw)

        # Важно: служебные переносы/табуляции -> пробелы, кроме тех, что пришли от <br>
        # Для этого схлопнем whitespace внутри каждой строки отдельно.
        raw = _normalize_inline_text(raw)

        # подчистим пробелы вокруг переносов
        raw = re.sub(r" *\n *", "\n", raw)
        return raw.strip()

    blocks_all = body.find_all(list(block_tags))
    blocks = [t for t in blocks_all if not _has_desc_block(t, block_tags)]

    out_lines: list[str] = []
    for el in blocks:
        txt = _text_of(el)
        if txt is None:
            continue

        txt = _norm_newlines_global(txt)
        # если блок реально пустой
        if txt.strip() == "":
            out_lines.append("")
            continue

        lvl = _indent_level_with_root(el, body)
        prefix = (" " * (lvl * indent_spaces)) if lvl > 0 else ""

        if el.name and el.name.lower() == "li":
            for line in txt.split("\n"):
                line = line.strip()
                if not line:
                    continue
                out_lines.append(prefix + "- " + line)
            continue

        for line in txt.split("\n"):
            if line.strip() == "":
                out_lines.append("")
            else:
                out_lines.append(prefix + line.rstrip())

    result = "\n".join(out_lines)
    result = _norm_newlines_global(result)
    result = re.sub(r"\n{3,}", "\n\n", result).rstrip()
    return result


def _sql_text_from_html(html: str, title: str) -> str:
    plain = _html_to_plain_preserving_layout(html or "")
    header = f"-- {title or ''}".rstrip()
    return header + ("\n\n" + plain if plain else "\n")


def export_html_to_sql_bytes(html: str, title: str, encoding: str = "utf-8") -> bytes:
    full_text = _sql_text_from_html(html or "", title or "")
    return full_text.encode(encoding, errors="replace")


@st.dialog("Новое сообщение", width="large")
def email_message_dialog(
    page_id: int,
    page_title: str,
    page_html: str,
    sender_login: str,
    page_path: str,
    deps: ExportDeps,
):
    st.session_state["email_dialog_page_id"] = int(page_id)
    recipients_df = deps.list_email_recipients()
    recipient_options = list(recipients_df.itertuples(index=False)) if not recipients_df.empty else []
    sender_full_name, sender_job_title = deps.get_user_signature(sender_login)

    # если ранее попросили сбросить поля добавления получателя — сделаем это ДО создания виджетов
    reset_flag_key = f"email_reset_fields_{page_id}"
    add_expanded_key = f"email_add_recipient_expanded_{page_id}"
    if st.session_state.pop(reset_flag_key, False):
        st.session_state[f"email_new_email_{page_id}"] = ""
        st.session_state[f"email_new_fio_{page_id}"] = ""
        st.session_state[f"email_new_gender_{page_id}"] = "муж"
        st.session_state[f"email_new_job_title_{page_id}"] = ""
        st.session_state[f"email_new_salutation_{page_id}"] = ""
        st.session_state[add_expanded_key] = False

    def _recipient_label(row) -> str:
        fio = getattr(row, "fio", "") or ""
        email_addr = getattr(row, "email", "") or ""
        job_title = getattr(row, "job_title", "") or ""
        if fio and email_addr:
            label = f"{fio} <{email_addr}>"
        else:
            label = fio or email_addr
        if job_title:
            label = f"{label} ({job_title})"
        return label

    def _attachment_label(row) -> str:
        size_txt = _format_file_size(getattr(row, "file_size", None))
        suffix = f" ({size_txt})" if size_txt else ""
        return f"{row.file_name}{suffix}"

    attachments_df = deps.get_page_attachments(page_id)
    file_records = [
        row for row in attachments_df.itertuples(index=False) if getattr(row, "attachment_type", "") == "file"
    ]
    link_records = [
        row
        for row in attachments_df.itertuples(index=False)
        if getattr(row, "attachment_type", "") == "link" and getattr(row, "url", None)
    ]

    def _build_links_block() -> str:
        if not link_records:
            return ""
        lines = []
        for row in link_records:
            title = (getattr(row, "file_name", "") or "").strip()
            url = (getattr(row, "url", "") or "").strip()
            if not url:
                continue
            if title:
                lines.append(f"- {title}: {url}")
            else:
                lines.append(f"- {url}")
        if not lines:
            return ""
        return "Ссылки:\n" + "\n".join(lines)

    subject_default = page_title or f"Страница {page_id}"
    body_plain = _html_to_plain_for_email(page_html or "")
    greeting = "Добрый день!"
    links_block = _build_links_block()
    page_info = f"{page_path}".strip()
    page_link = deps.build_page_deeplink(page_id)
    signature = f"С уважением!\n{sender_full_name}\n{sender_job_title}".rstrip()

    body_parts = [greeting]
    if body_plain:
        body_parts.append(body_plain)
    if links_block:
        body_parts.append(links_block)
    if page_path:
        body_parts.append(page_info)
        if page_link:
            body_parts.append(page_link)
    body_parts.append(signature)
    body_default = "\n\n".join(part for part in body_parts if part)

    with st.form(f"email_form_{page_id}", clear_on_submit=False):
        subject = st.text_input("Тема", value=subject_default)

        if not recipient_options:
            st.caption("Справочник получателей пуст.")

        to_dir = st.multiselect(
            "Получатели (справочник)",
            options=recipient_options,
            format_func=_recipient_label,
            key=f"email_to_dir_{page_id}",
        )
        cc_dir = st.multiselect(
            "Копия (справочник)",
            options=recipient_options,
            format_func=_recipient_label,
            key=f"email_cc_dir_{page_id}",
        )

        bcc_dir = st.multiselect(
            "Скрытая копия (справочник)",
            options=recipient_options,
            format_func=_recipient_label,
            key=f"email_bcc_dir_{page_id}",
        )

        body = st.text_area("Тело письма", value=body_default, height=200)
        important = st.checkbox("Важное", value=False)

        if file_records:
            selected_files = st.multiselect(
                "Вложения страницы",
                options=file_records,
                format_func=_attachment_label,
                key=f"email_files_{page_id}",
            )
        else:
            st.caption("Вложения страницы отсутствуют.")
            selected_files = []

        send_col, cancel_col, add_col = st.columns([1, 1, 2])
        with send_col:
            submitted = st.form_submit_button("Отправить")
        with cancel_col:
            canceled = st.form_submit_button("Отмена")
        with add_col:
            with st.expander("Добавить получателя", expanded=st.session_state.get(add_expanded_key, False)):
                new_email = st.text_input("eMail", key=f"email_new_email_{page_id}")
                new_fio = st.text_input("ФИО", key=f"email_new_fio_{page_id}")
                new_gender = st.selectbox(
                    "Пол",
                    options=["муж", "жен"],
                    key=f"email_new_gender_{page_id}",
                )
                new_job_title = st.text_input("Должность", key=f"email_new_job_title_{page_id}")
                new_salutation = st.text_input("Как обращаться", key=f"email_new_salutation_{page_id}")
                add_recipient = st.form_submit_button("Добавить")

    if canceled:
        st.session_state["email_dialog_page_id"] = None
        st.rerun()
    elif add_recipient:
        email_trim = new_email.strip().lower()
        fio_trim = new_fio.strip()

        if not email_trim or not fio_trim:
            st.error("Заполните eMail и ФИО.")
            return

        if not _is_valid_email(email_trim):
            st.error("Некорректный формат eMail.")
            return

        exists = deps.run_scalar(
            f"""
            SELECT 1
            FROM {deps.email_recipients_table}
            WHERE lower(email) = lower('{deps.escape(email_trim)}')
            LIMIT 1
            """
        )
        if exists:
            st.error("eMail уже существует.")
            return

        deps.run_execute(
            f"""
            INSERT INTO {deps.email_recipients_table}
                (email, fio, job_title, salutation, gender, created_by)
            VALUES
                ('{deps.escape(email_trim)}',
                 '{deps.escape(fio_trim)}',
                 '{deps.escape(new_job_title.strip())}',
                 '{deps.escape(new_salutation.strip())}',
                 '{deps.escape(new_gender)}',
                 '{deps.escape(sender_login)}')
            """
        )
        _notebook_id, _section_id = deps.get_nb_section_id(page_id)
        deps.add_event_log(
            topic="RECIPIENT",
            subtopic="CREATE",
            notebook_id=_notebook_id,
            section_id=_section_id,
            page_id=page_id,
            event="create_recipient: fio = " + deps.escape(fio_trim) + "; email = " + deps.escape(email_trim),
            body_html="",
        )
        st.success("Получатель добавлен")
        # очистить поля ввода и свернуть экспандер (на следующем ререндере)
        st.session_state[f"email_reset_fields_{page_id}"] = True
        st.session_state["email_dialog_page_id"] = page_id
        st.rerun()
    elif submitted:
        to_list = [r.email for r in to_dir if getattr(r, "email", None)]
        cc_list = [r.email for r in cc_dir if getattr(r, "email", None)]
        bcc_list = [r.email for r in bcc_dir if getattr(r, "email", None)]

        recipients = _merge_email_lists(to_list)
        cc = _merge_email_lists(cc_list)
        bcc = _merge_email_lists(bcc_list)

        if not recipients:
            st.error("Укажите получателей.")
            return

        temp_paths: list[str] = []
        try:
            base_dir = os.path.dirname(__file__)
            temp_root = os.path.join(base_dir, config.temp_attachments_dir)
            os.makedirs(temp_root, exist_ok=True)
            with tempfile.TemporaryDirectory(prefix="notes_email_", dir=temp_root) as tmpdir:
                for row in selected_files:
                    payload = deps.get_attachment_file(int(row.id))
                    if not payload:
                        st.warning(
                            f"Не удалось подгрузить вложение: {row.file_name}"
                        )
                        continue
                    data, file_name, _mime = payload
                    safe_name = _sanitize_attachment_filename(file_name)
                    file_path = os.path.join(tmpdir, safe_name)
                    with open(file_path, "wb") as handle:
                        handle.write(data)
                    temp_paths.append(file_path)

                body_text = body or ""
                if not body_text.strip():
                    body_text = greeting

                if not body_text.lstrip().startswith(greeting):
                    body_text = f"{greeting}\n\n{body_text.lstrip()}"

                links_block = _build_links_block()
                page_info = f"{page_path}".strip()
                signature = f"С уважением!\n{sender_full_name}\n{sender_job_title}".rstrip()

                if links_block and "Ссылки:" not in body_text:
                    if "С уважением!" in body_text:
                        body_text = body_text.replace("С уважением!", f"{links_block}\n\nС уважением!", 1)
                    else:
                        body_text = f"{body_text.rstrip()}\n\n{links_block}"

                def _insert_before_signature(text: str, insert: str) -> str:
                    if not insert:
                        return text
                    if "С уважением!" in text:
                        return text.replace("С уважением!", f"{insert}\n\nС уважением!", 1)
                    return f"{text.rstrip()}\n\n{insert}"

                if page_path and page_info not in body_text:
                    block = page_info
                    page_link = deps.build_page_deeplink(page_id)
                    if page_link:
                        block = f"{block}\n{page_link}"
                    body_text = _insert_before_signature(body_text, block)
                else:
                    page_link = deps.build_page_deeplink(page_id)
                    if page_link and page_link not in body_text:
                        body_text = _insert_before_signature(body_text, page_link)

                if "С уважением!" not in body_text:
                    body_text = f"{body_text.rstrip()}\n\n{signature}"

                send_mail(
                    subject=subject or "",
                    recipients=recipients,
                    cc=cc,
                    bcc=bcc,
                    body=body_text,
                    important=bool(important),
                    files=temp_paths,
                )

            st.success("Сообщение отправлено")
            st.session_state["email_dialog_page_id"] = None
            st.rerun()
        except Exception as exc:
            st.error(f"Ошибка отправки: {exc}")


def render_export_section(
    *,
    container,
    page_id: int,
    current_title: str,
    current_html: str,
    selected_login: str,
    page_path: str,
    deps: ExportDeps,
) -> None:
    with container:
        exp_export_nonce_key = f"exp_export_nonce_{page_id}"
        st.session_state.setdefault(exp_export_nonce_key, 0)

        # состояние docx на страницу
        docx_state_key = f"docx_state_{page_id}"  # dict | None
        st.session_state.setdefault(docx_state_key, {"data": None, "file_name": None, "mime": None, "error": None})

        def _collapse_export():
            st.session_state[exp_export_nonce_key] += 1

        exp_export_label = "Экспорт" + ("\u200b" * int(st.session_state[exp_export_nonce_key]))

        with st.expander(exp_export_label, expanded=False):
            safe_title2 = current_title or f"Страница_{page_id}"

            # ---------- ОДНА КНОПКА: .docx ----------
            # при нажатии: генерируем bytes, сохраняем в state и автокликаем download
            clicked_docx = st.button(".docx", key=f"btn_docx_{page_id}", use_container_width=True)

            if clicked_docx:
                try:
                    buf = export_html_to_docx_bytes(current_html or "", safe_title2)
                    data_bytes = buf.getvalue() if hasattr(buf, "getvalue") else bytes(buf)

                    st.session_state[docx_state_key] = {
                        "data": data_bytes,
                        "file_name": _safe_filename(safe_title2, "docx"),
                        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "error": None,
                    }
                except Exception as exc:
                    st.session_state[docx_state_key] = {
                        "data": None,
                        "file_name": _safe_filename(safe_title2, "docx"),
                        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "error": str(exc),
                    }

            docx_state = st.session_state.get(docx_state_key) or {}
            if docx_state.get("error"):
                st.error(f"Ошибка формирования DOCX: {docx_state['error']}")

            # Если есть готовые bytes — рисуем download_button и автокликаем его (один клик пользователя)
            if docx_state.get("data"):
                dl_key = f"dl_docx_hidden_{page_id}"

                # сам download_button можно сделать невидимым, но он должен быть в DOM
                st.download_button(
                    label="Скачать DOCX",
                    data=docx_state["data"],
                    file_name=docx_state.get("file_name") or _safe_filename(safe_title2, "docx"),
                    mime=docx_state.get("mime")
                    or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=dl_key,
                    on_click=_collapse_export,
                    use_container_width=True,
                )

                # авто-клик по кнопке download_button
                # (ищем DOM-кнопку по data-testid + тексту; работает стабильно в Streamlit)
                components.html(
                    f"""
                    <script>
                    (function() {{
                      try {{
                        const root = window.parent.document;
                        const buttons = root.querySelectorAll('button');
                        let target = null;
                        buttons.forEach(b => {{
                          if ((b.innerText || '').trim() === 'Скачать DOCX') target = b;
                        }});
                        if (target) {{
                          setTimeout(() => target.click(), 30);
                        }}
                      }} catch(e) {{}}
                    }})();
                    </script>
                    """,
                    height=0,
                )

                # чтобы не автокликало при следующем rerun (и не скачивало повторно)
                # сбрасываем data сразу после постановки автоклика
                st.session_state[docx_state_key]["data"] = None


            if st.button("eMail", key=f"open_email_dialog_{page_id}", use_container_width=True):
                st.session_state["email_dialog_page_id"] = page_id

            if st.session_state.get("email_dialog_page_id") == page_id:
                email_message_dialog(page_id, safe_title2, current_html or "", selected_login, page_path, deps)
                _collapse_export()

            # ---------- SQL (как было) ----------
            st.download_button(
                ".sql (utf-8)",
                data=export_html_to_sql_bytes(current_html or "", safe_title2, encoding="utf-8"),
                file_name=_safe_filename(safe_title2, "sql"),
                mime="text/plain; charset=utf-8",
                on_click=_collapse_export,
                key=f"dl_sql_u8_{page_id}",
            )


            # Не удалять!!!!!!
            
            # st.download_button(
            #     ".sql (cp1251)",
            #     data=export_html_to_sql_bytes(current_html or "", safe_title2, encoding="cp1251"),
            #     file_name=_safe_filename(safe_title2, "sql"),
            #     mime="text/plain; charset=windows-1251",
            #     on_click=_collapse_export,
            #     key=f"dl_sql_1251_{page_id}",
            # )
