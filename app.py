import logging
import os
import streamlit as st
import streamlit.components.v1 as components
import subprocess
import sys
import tempfile

# if "reqs_installed" not in st.session_state:
#     req_path = os.path.join(os.path.dirname(__file__), "requirements.txt")
#     if os.path.exists(req_path):
#         try:
#             subprocess.check_call([
#                 sys.executable, "-m", "pip", "install", "--user", "-r", req_path
#             ])
#             st.session_state["reqs_installed"] = True
#         except subprocess.CalledProcessError as e:
#             print("Ошибка при установке зависимостей:", e)
#             st.error("Не удалось установить зависимости. См. терминал.")

import base64
import email
import html
import io
import json
import re
import zipfile
import urllib.parse
import datetime
from typing import Optional

import pandas as pd
import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from st_aggrid import AgGrid, GridOptionsBuilder, JsCode
from st_aggrid.shared import GridUpdateMode, DataReturnMode
from streamlit_quill import st_quill

import config
from src.database_utils_DRP import get_execute, get_fetch, test_connection
from src.mail import send_mail

SCHEMA = "sbx_dfip_ocpp"
USERS_TABLE = f"{SCHEMA}.notes_users"
NOTEBOOKS_TABLE = f"{SCHEMA}.notes_notebooks"
SECTIONS_TABLE = f"{SCHEMA}.notes_sections"
PAGES_TABLE = f"{SCHEMA}.notes_pages"
OWNERS_TABLE = f"{SCHEMA}.notes_notebook_owners"
DEPARTMENTS_TABLE = f"{SCHEMA}.notes_departments"
ATTACHMENTS_TABLE = f"{SCHEMA}.notes_page_attachments"
EMAIL_RECIPIENTS_TABLE = f"{SCHEMA}.notes_email_recipients"
EVENT_LOGS_TABLE = f"{SCHEMA}.notes_event_logs"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("notes_app")


def _escape(val: str) -> str:
    """Минимальное экранирование строк для SQL."""
    safe = (val or "").replace("\x00", "")
    return safe.replace("'", "''")



def _escape_like(val: str) -> str:
    """
    Экранируем спецсимволы для ILIKE-шаблонов (% и _), чтобы пользователь
    не мог ломать шаблон. Используем ESCAPE '\\' в SQL.
    """
    s = (val or "")
    s = s.replace("\\", "\\\\")
    s = s.replace("%", "\\%")
    s = s.replace("_", "\\_")
    return _escape(s)


# -------------------------
# Advanced search parser (Google-like)
# -------------------------
class _Tok:
    def __init__(self, typ: str, val: str = ""):
        self.typ = typ  # WORD, PHRASE, LPAREN, RPAREN, OR, PLUS, MINUS
        self.val = val


def _lex_search(q: str) -> list[_Tok]:
    """
    Лексер:
      - "фраза" -> PHRASE(val)
      - ( )     -> LPAREN/RPAREN
      - |       -> OR
      - +       -> PLUS (унарный)
      - -       -> MINUS (унарный)
      - остальное -> WORD (может содержать '*')
    """
    s = (q or "").strip()
    out: list[_Tok] = []
    i = 0
    n = len(s)
    while i < n:
        ch = s[i]

        if ch.isspace():
            i += 1
            continue

        if ch == '"':
            i += 1
            buf = []
            while i < n and s[i] != '"':
                buf.append(s[i])
                i += 1
            if i < n and s[i] == '"':
                i += 1
            phrase = "".join(buf).strip()
            if phrase:
                out.append(_Tok("PHRASE", phrase))
            continue

        if ch == "(":
            out.append(_Tok("LPAREN", ch))
            i += 1
            continue
        if ch == ")":
            out.append(_Tok("RPAREN", ch))
            i += 1
            continue
        if ch == "|":
            out.append(_Tok("OR", ch))
            i += 1
            continue
        if ch == "+":
            out.append(_Tok("PLUS", ch))
            i += 1
            continue
        if ch == "-":
            out.append(_Tok("MINUS", ch))
            i += 1
            continue

        # WORD
        buf = []
        while i < n:
            ch2 = s[i]
            if ch2.isspace() or ch2 in ['"', "(", ")", "|"]:
                break
            # + и - считаем частью слова, если они НЕ в начале токена
            if ch2 in ["+", "-"] and not buf:
                break
            buf.append(ch2)
            i += 1
        word = "".join(buf).strip()
        if word:
            out.append(_Tok("WORD", word))
        else:
            # если упёрлись в + / - в начале — обработаем как оператор
            if i < n and s[i] == "+":
                out.append(_Tok("PLUS", "+"))
                i += 1
            elif i < n and s[i] == "-":
                out.append(_Tok("MINUS", "-"))
                i += 1
            else:
                i += 1

    return out




def _fields_like_expr_text(like_pattern_sql: str) -> str:
    # В Postgres ESCAPE должен быть пустым или состоять из ОДНОГО байта.
    # Надёжный вариант: ESCAPE E'\\' (это ровно один символ backslash).
    esc = " ESCAPE E'\\\\' "
    return (
        f"(COALESCE(p.title, '') ILIKE {like_pattern_sql}{esc}"
        f" OR COALESCE(p.body_html, '') ILIKE {like_pattern_sql}{esc}"
        f" OR COALESCE(p.tag, '') ILIKE {like_pattern_sql}{esc})"
    )


def _fields_like_expr_tags(like_pattern_sql: str) -> str:
    esc = " ESCAPE E'\\\\' "
    return f"(COALESCE(p.tag, '') ILIKE {like_pattern_sql}{esc})"



def _term_to_sql(tok: _Tok, mode: str = "text") -> str:
    """
    WORD / PHRASE -> SQL expr.
    WORD: поддержка wildcard '*' -> '%'

    mode:
      - "text": поиск по title/body_html
      - "tags": поиск по p.tag
    """
    def _fields_expr(pat_sql: str) -> str:
        return _fields_like_expr_tags(pat_sql) if mode == "tags" else _fields_like_expr_text(pat_sql)

    if tok.typ == "PHRASE":
        pat = f"'%{_escape_like(tok.val)}%'"
        return _fields_expr(pat)

    if tok.typ == "WORD":
        w = tok.val.strip()
        if not w:
            return "TRUE"

        if "*" in w:
            # '*' разрешаем как wildcard. В _escape_like '*' не экранируется, но на всякий случай:
            safe = _escape_like(w).replace("\\*", "*")
            safe = safe.replace("*", "%")
            if not safe.startswith("%"):
                safe = "%" + safe
            if not safe.endswith("%"):
                safe = safe + "%"
            return _fields_expr(f"'{safe}'")

        pat = f"'%{_escape_like(w)}%'"
        return _fields_expr(pat)

    return "TRUE"



class _Parser:
    """
    Грамматика (OR ниже AND, как Google):
      expr     := or_expr
      or_expr  := and_expr (OR and_expr)*
      and_expr := unary (unary)*            # неявный AND по соседству
      unary    := (PLUS|MINUS)* primary
      primary  := TERM | '(' expr ')'
      TERM     := WORD | PHRASE
    """
    def __init__(self, toks: list[_Tok], term_sql_func):
        self.toks = toks
        self.i = 0
        self.term_sql_func = term_sql_func  # TERM -> SQL

    def _peek(self) -> _Tok | None:
        return self.toks[self.i] if self.i < len(self.toks) else None

    def _eat(self, typ: str) -> _Tok | None:
        t = self._peek()
        if t and t.typ == typ:
            self.i += 1
            return t
        return None

    def parse(self) -> str:
        if not self.toks:
            return ""
        return self._parse_or()

    def _parse_or(self) -> str:
        left = self._parse_and()
        while self._eat("OR"):
            right = self._parse_and()
            left = f"({left} OR {right})"
        return left

    def _starts_unary_or_primary(self, t: _Tok | None) -> bool:
        return bool(t and t.typ in ("PLUS", "MINUS", "WORD", "PHRASE", "LPAREN"))

    def _parse_and(self) -> str:
        left = self._parse_unary()
        while True:
            t = self._peek()
            if not self._starts_unary_or_primary(t):
                break
            if t and t.typ in ("RPAREN", "OR"):
                break
            right = self._parse_unary()
            left = f"({left} AND {right})"
        return left

    def _parse_unary(self) -> str:
        neg = False
        while True:
            if self._eat("PLUS"):
                continue
            if self._eat("MINUS"):
                neg = not neg
                continue
            break

        prim = self._parse_primary()
        return f"(NOT {prim})" if neg else prim

    def _parse_primary(self) -> str:
        if self._eat("LPAREN"):
            inside = self._parse_or()
            self._eat("RPAREN")
            return f"({inside})"

        t = self._peek()
        if t and t.typ in ("WORD", "PHRASE"):
            self.i += 1
            return self.term_sql_func(t)

        if t:
            self.i += 1
        return "TRUE"





def build_advanced_search_where(search_raw: str, mode: str = "text") -> str:
    toks = _lex_search(search_raw)
    if not toks:
        return ""

    def _term_func(t: _Tok) -> str:
        return _term_to_sql(t, mode=mode)

    parser = _Parser(toks, term_sql_func=_term_func)
    expr_sql = parser.parse().strip()
    if not expr_sql:
        return ""
    return " AND " + expr_sql



# --- SIMPLE HTML SANITIZER (no bleach) ---

_DANGEROUS_TAGS = {
    "script", "iframe", "object", "embed", "link", "meta", "base",
    "form", "input", "button", "textarea", "select", "option",
    "svg", "math",
}

_ALLOWED_TAGS = {
    "div","p","span","br",
    "strong","b","em","i","u","s",
    "a","img",
    "ul","ol","li",
    "h1","h2","h3","h4","h5","h6",
    "blockquote","pre","code",
    "table","thead","tbody","tfoot","tr","td","th","colgroup","col",
    "mark",
}


_ALLOWED_ATTRS_COMMON = {"style", "class"}


_ALLOWED_ATTRS = {
    "a": {"href", "title", "target", "rel"} | _ALLOWED_ATTRS_COMMON,
    "img": {"src", "alt", "title", "width", "height", "style"} | _ALLOWED_ATTRS_COMMON,
    "td": {"colspan", "rowspan"} | _ALLOWED_ATTRS_COMMON,
    "th": {"colspan", "rowspan"} | _ALLOWED_ATTRS_COMMON,
    "col": {"span", "width"} | _ALLOWED_ATTRS_COMMON,
    "mark": _ALLOWED_ATTRS_COMMON,
    "*": _ALLOWED_ATTRS_COMMON,
}



_ALLOWED_CSS_PROPS = {
    "font-family", "font-size", "font-weight", "font-style",
    "text-decoration", "color", "line-height",
    "background", "background-color",
    "text-align",
    "margin", "margin-top", "margin-right", "margin-bottom", "margin-left",
    "padding", "padding-top", "padding-right", "padding-bottom", "padding-left",
    "border", "border-width", "border-style", "border-color",
    "border-collapse",
    "vertical-align",
    "width",
    "white-space",
}

_RE_STYLE_DECL = re.compile(r"\s*([-\w]+)\s*:\s*([^;]+)\s*;?", re.UNICODE)

def _is_bad_url(url: str) -> bool:
    u = (url or "").strip().lower()
    return (
        u.startswith("javascript:")
        or u.startswith("vbscript:")
        or u.startswith("data:text/html")
        or u.startswith("data:application/xhtml+xml")
    )

def _sanitize_style(style: str) -> str:
    if not style:
        return ""
    s = style
    s = re.sub(r"expression\s*\([^)]*\)", "", s, flags=re.I)
    s = re.sub(r"javascript\s*:", "", s, flags=re.I)
    s = re.sub(r"url\s*\(\s*['\"]?\s*javascript:[^)]*\)", "", s, flags=re.I)
    s = re.sub(r"[\x00-\x1f\x7f]", "", s)

    out: list[str] = []
    for prop, val in _RE_STYLE_DECL.findall(s):
        prop_l = prop.strip().lower()
        val = val.strip()

        # OneNote
        if prop_l == "mso-highlight":
            prop_l = "background-color"

        if prop_l not in _ALLOWED_CSS_PROPS:
            continue

        # запретим любые url(...) в значениях
        if re.search(r"url\s*\(", val, flags=re.I):
            continue

        out.append(f"{prop_l}: {val}")

    return "; ".join(out)

def sanitize_html_safe(html_in: str) -> str:
    """
    Универсальная санитизация без bleach:
    - удаляет опасные теги
    - оставляет только whitelist тегов/атрибутов
    - режет on* обработчики
    - фильтрует href/src
    - чистит style по whitelist CSS свойств
    """
    if not html_in:
        return ""

    soup = BeautifulSoup(html_in, "html.parser")
    root = soup.body or soup

    # 1) удалить опасные теги целиком
    for bad in list(root.find_all(list(_DANGEROUS_TAGS))):
        bad.decompose()

    # 2) нормализация и чистка
    for tag in list(root.find_all(True)):
        name = (tag.name or "").lower()

        # неизвестные теги — unwrap (сохраняем текст)
        if name not in _ALLOWED_TAGS:
            tag.unwrap()
            continue

        allowed_attrs = _ALLOWED_ATTRS.get(name, _ALLOWED_ATTRS["*"])

        for attr in list(tag.attrs.keys()):
            al = str(attr).lower()

            # убрать обработчики событий
            if al.startswith("on"):
                del tag.attrs[attr]
                continue

            # whitelist атрибутов
            if attr not in allowed_attrs:
                del tag.attrs[attr]
                continue

        # href/src safety
        if tag.has_attr("href") and _is_bad_url(tag.get("href", "")):
            del tag.attrs["href"]
        if tag.has_attr("src") and _is_bad_url(tag.get("src", "")):
            del tag.attrs["src"]

        # style
        if tag.has_attr("style"):
            safe_style = _sanitize_style(tag.get("style", ""))
            if safe_style:
                tag["style"] = safe_style
            else:
                del tag.attrs["style"]

        # ссылки
        if name == "a":
            tag["target"] = "_blank"
            tag["rel"] = "noopener noreferrer"

    return str(root)





def preserve_onenote_blank_paragraphs(html_in: str) -> str:
    """
    OneNote пустые строки между абзацами делает как <p>&nbsp;</p>.
    Мы превращаем такие абзацы в <p><br/></p>, чтобы они не удалялись
    нормализацией/санитизацией и корректно отображались в редакторе.
    """
    if not html_in:
        return ""

    soup = BeautifulSoup(html_in, "html.parser")
    root = soup.body or soup

    for p in root.find_all("p"):
        if not isinstance(p, Tag):
            continue

        # если внутри есть реальная структура — не трогаем
        if p.find(["img", "table", "a", "br"]):
            continue

        txt = p.get_text("", strip=False) or ""
        # только пробелы/переводы + NBSP
        if txt.replace("\xa0", "").strip() == "":
            p.clear()
            p.append(soup.new_tag("br"))

    return str(root)




def _onenote_keep_blank_lines_in_div(page_div: Tag) -> None:
    """
    OneNote пустые строки = <p>&nbsp;</p>.
    Превращаем их в <p><br/></p>, чтобы _normalize_onenote_page_div не удалил.
    """
    if not isinstance(page_div, Tag):
        return

    soup = page_div if hasattr(page_div, "new_tag") else None  # обычно это BeautifulSoup Tag

    for p in page_div.find_all("p"):
        if not isinstance(p, Tag):
            continue

        # если в <p> уже есть структура — не трогаем
        if p.find(["img", "table", "a", "br"]):
            continue

        # текст без пробелов, но NBSP считаем пустотой
        txt = p.get_text("", strip=False) or ""
        if txt.replace("\xa0", "").strip() == "":
            p.clear()
            # создаём <br>
            br = p._parent.new_tag("br") if getattr(p, "_parent", None) else None
            if br is None:
                # fallback
                br = BeautifulSoup("", "html.parser").new_tag("br")
            p.append(br)




def sanitize_html_simple_keep_formatting(html_in: str) -> str:
    if not html_in:
        return ""

    soup = BeautifulSoup(html_in, "html.parser")
    root = soup.body or soup

    # 1) убрать опасные теги целиком
    for bad in list(root.find_all(_DANGEROUS_TAGS)):
        bad.decompose()

    # 2) пройти по всем тегам
    for tag in list(root.find_all(True)):
        name = (tag.name or "").lower()

        # неразрешённые теги разворачиваем, сохраняя текст
        if name not in _ALLOWED_TAGS:
            tag.unwrap()
            continue

        allowed = _ALLOWED_ATTRS.get(name, _ALLOWED_ATTRS.get("*", set()))
        allowed_all = set(allowed) | set(_ALLOWED_ATTRS.get("*", set()))

        # attrs whitelist + remove on*
        for attr in list(tag.attrs.keys()):
            al = str(attr).lower()
            if al.startswith("on"):
                del tag.attrs[attr]
                continue
            if attr not in allowed_all:
                del tag.attrs[attr]
                continue

        # href/src safety
        if tag.has_attr("href") and _is_bad_url(tag.get("href", "")):
            del tag.attrs["href"]
        if tag.has_attr("src") and _is_bad_url(tag.get("src", "")):
            del tag.attrs["src"]

        # clean style
        if tag.has_attr("style"):
            safe_style = _sanitize_style(tag.get("style", ""))
            if safe_style:
                tag["style"] = safe_style
            else:
                del tag.attrs["style"]

        # normalize links
        if name == "a":
            tag["target"] = "_blank"
            tag["rel"] = "noopener noreferrer"

    return str(root)


_IMPORT_ALLOWED_PROTOCOLS = ["http", "https", "mailto", "data"]

_RE_STYLE_DECL2 = re.compile(r"\s*([-\w]+)\s*:\s*([^;]+)\s*;?", re.UNICODE)



_MONTHS_RU = r"(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)"
_RE_DATE_RU = re.compile(rf"^\s*\d{{1,2}}\s+{_MONTHS_RU}\s+\d{{4}}\s*г\.?\s*$", re.IGNORECASE)
_RE_TIME = re.compile(r"^\s*\d{1,2}:\d{2}\s*$")

def strip_onenote_datetime_block(html: str) -> str:
    if not html:
        return ""

    soup = BeautifulSoup(html, "html.parser")
    root = soup.body or soup

    # Сначала удаляем p со временем/датой
    for p in root.find_all("p"):
        txt = (p.get_text(" ", strip=True) or "").replace("\xa0", " ").strip()
        if not txt:
            continue

        if _RE_TIME.match(txt) or _RE_DATE_RU.match(txt):
            p.decompose()

    # Затем подчистим пустые div, которые могли остаться
    for div in root.find_all("div"):
        # если в div нет текста и нет картинок/таблиц/ссылок — убрать
        text = div.get_text(" ", strip=True).replace("\xa0", " ").strip()
        has_struct = div.find(["img", "table", "a"])
        if (not text) and (not has_struct):
            div.decompose()

    return str(root)



def _format_file_size(size: int | float | None) -> str:
    if not size or size <= 0:
        return ""
    units = ["B", "KB", "MB", "GB"]
    value = float(size)
    for unit in units:
        if value < 1024 or unit == units[-1]:
            return f"{value:.1f} {unit}"
        value /= 1024
    return f"{value:.1f} B"


def _looks_like_imported_html(html: str) -> bool:
    if not html:
        return False
    h = html.lstrip().lower()
    if h.startswith("<div") and "direction:ltr" in h and "font-family" in h:
        return True
    if "mso-" in h or "margin-left:0in" in h:
        return True
    return False





def _merge_email_lists(*lists: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for items in lists:
        for item in items:
            if not item or item in seen:
                continue
            seen.add(item)
            out.append(item)
    return out


def _is_valid_email(value: str) -> bool:
    val = (value or "").strip()
    if not val:
        return False
    return re.match(r"^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$", val) is not None


def _sanitize_attachment_filename(name: str) -> str:
    base = os.path.basename(name or "attachment")
    base = _transliterate_ru_filename(base)
    base = re.sub(r"[^A-Za-z0-9_.-]+", "_", base)
    base = re.sub(r"_+", "_", base).strip("._- ")
    return base or "attachment"


def _transliterate_ru_filename(text: str) -> str:
    if not text:
        return ""

    m = {
        "а": "a",
        "б": "b",
        "в": "v",
        "г": "g",
        "д": "d",
        "е": "e",
        "ё": "yo",
        "ж": "zh",
        "з": "z",
        "и": "i",
        "й": "y",
        "к": "k",
        "л": "l",
        "м": "m",
        "н": "n",
        "о": "o",
        "п": "p",
        "р": "r",
        "с": "s",
        "т": "t",
        "у": "u",
        "ф": "f",
        "х": "kh",
        "ц": "ts",
        "ч": "ch",
        "ш": "sh",
        "щ": "shch",
        "ъ": "",
        "ы": "y",
        "ь": "",
        "э": "e",
        "ю": "yu",
        "я": "ya",
    }
    out: list[str] = []
    for ch in text:
        low = ch.lower()
        if low in m:
            rep = m[low]
            if ch.isupper() and rep:
                rep = rep[0].upper() + rep[1:]
            out.append(rep)
        else:
            out.append(ch)
    return "".join(out)


def ensure_db_credentials() -> dict[str, str]:
    """Запрос логина/пароля к БД один раз за сессию."""
    creds = st.session_state.get("db_credentials")
    if creds and creds.get("user") and creds.get("password"):
        st.session_state.setdefault("current_user_login", creds["user"])
        return creds

    # если уже открыт другой диалог, сначала закрываем его и запускаем rerun,
    # чтобы окно авторизации показалось отдельно (Streamlit запрещает вложенные диалоги)
    if st.session_state.get("dialog_open"):
        st.session_state["dialog_open"] = False
        st.rerun()
        st.stop()

    @st.dialog("Подключение к базе", width="small")
    def _ask_credentials():
        st.session_state["dialog_open"] = True
        st.write("Введите логин и пароль.")
        with st.form("db_login_form", clear_on_submit=False):
            user = st.text_input("Логин", key="db_login")
            pwd = st.text_input("Пароль", type="password", key="db_password")
            submitted = st.form_submit_button("Подключиться")
            if submitted:
                if not user or not pwd:
                    st.error("Укажите логин и пароль.")
                    return
                if not test_connection(user, pwd):
                    st.error("Не удалось подключиться. Проверьте данные.")
                    return
                st.session_state["db_credentials"] = {"user": user, "password": pwd}
                st.session_state["current_user_login"] = user
                add_event_log(topic="LOGON", subtopic="LOGON",notebook_id=0, section_id=0, page_id=0, event="LOGON user " + user, body_html="")
                st.success("Подключение установлено. Обновляем страницу...")
                st.session_state["dialog_open"] = False
                st.rerun()

    _ask_credentials()
    st.session_state["dialog_open"] = False
    st.stop()


def cleanup_docs_dir() -> None:
    base_dir = os.path.dirname(__file__)
    docs_dir = os.path.join(base_dir, config.temp_attachments_dir)
    if not os.path.isdir(docs_dir):
        return

    today = datetime.date.today()
    for name in os.listdir(docs_dir):
        if name in ("__init__", "__init__.py"):
            continue
        path = os.path.join(docs_dir, name)
        if not os.path.isfile(path):
            continue
        try:
            mtime = datetime.date.fromtimestamp(os.path.getmtime(path))
        except OSError:
            continue
        if mtime < today:
            try:
                os.remove(path)
                logger.info("Удален устаревший файл: %s", path)
            except OSError as exc:
                logger.warning("Не удалось удалить файл %s: %s", path, exc)


st.session_state.setdefault("edit_dialog_page_id", None)
st.session_state.setdefault("download_payload", None)  # tuple[bytes, str, str] | None
st.session_state.setdefault("download_att_id", None)  # int | None
st.session_state.setdefault("download_error", None)  # str | None
st.session_state.setdefault("open_editor_once_for_page", None)  # int|None
st.session_state.setdefault("open_editor_once_edit_mode", False)  # bool: открыть диалог сразу в режиме редактирования
st.session_state.setdefault("dialog_open", False)  # защита от вложенных диалогов
st.session_state.setdefault("email_dialog_page_id", None)  # int | None


def _creds() -> tuple[str, str]:
    creds = ensure_db_credentials()
    return creds["user"], creds["password"]


def run_fetch_df(query: str) -> pd.DataFrame:
    user, pwd = _creds()
    try:
        df = get_fetch(query, user, pwd)
        if df is None:
            return pd.DataFrame()
        return df
    except Exception as e:
        logger.exception("DB fetch error: %s", e)
        st.error(f"Ошибка запроса к БД: {e}")
        return pd.DataFrame()



def run_execute(query: str) -> int | None:
    user, pwd = _creds()
    return get_execute(query, user, pwd)


def run_scalar(query: str):
    df = run_fetch_df(query)
    if df.empty:
        return None
    return df.iat[0, 0]


def list_users() -> pd.DataFrame:
    return run_fetch_df(
        f"""
        SELECT login, full_name, department_id
        FROM {USERS_TABLE}
        ORDER BY COALESCE(full_name, login)
        """
    )


# По page_id возвращает notebook_id, section_id
def get_nb_section_id(page_id: int) -> tuple[int, int]:
    if page_id == 0 or page_id is None:
        return 0,0
    
    df = run_fetch_df(
    f"""
        SELECT s.notebook_id, p.section_id
        FROM {PAGES_TABLE} p
            INNER JOIN {SECTIONS_TABLE} s ON p.section_id = s.id
        WHERE p.id = {page_id}
        LIMIT 1
        """
    )
    return df.at[0, "notebook_id"], df.at[0, "section_id"]


def get_user_signature(login: str) -> tuple[str, str]:
    df = run_fetch_df(
        f"""
        SELECT full_name, job_title
        FROM {USERS_TABLE}
        WHERE login = '{_escape(login)}'
        LIMIT 1
        """
    )
    if df.empty:
        return login, ""
    full_name = str(df.at[0, "full_name"] or "").strip() or login
    job_title = str(df.at[0, "job_title"] or "").strip()
    return full_name, job_title


def list_email_recipients() -> pd.DataFrame:
    return run_fetch_df(
        f"""
        SELECT email, fio, job_title, salutation
        FROM {EMAIL_RECIPIENTS_TABLE}
        ORDER BY COALESCE(fio, email)
        """
    )


def get_departments() -> pd.DataFrame:
    """Справочник подразделений (включая '99 Создатель книги')."""
    return run_fetch_df(
        f"""
        SELECT department_id, name_department
        FROM {DEPARTMENTS_TABLE}
        ORDER BY name_department
        """
    )



def add_notebook_owner(notebook_id: int, user_login: str) -> None:
    # logger.info("Добавление владельца: notebook_id=%s user=%s", notebook_id, user_login)
    run_execute(
        f"""
        INSERT INTO {OWNERS_TABLE} (notebook_id, user_login)
        VALUES ({int(notebook_id)}, '{_escape(user_login)}')
        """
    )


def add_event_log( topic: str, subtopic: str,notebook_id: int, section_id: int, page_id: int , event: str, body_html : str) -> None:
    run_execute(
        f"""
        INSERT INTO {EVENT_LOGS_TABLE} (user_name, topic, subtopic, notebook_id, section_id, page_id, "event", body_html)
        VALUES ('{st.session_state["current_user_login"]}','{_escape(topic)}','{_escape(subtopic)}',{int(notebook_id)},{int(section_id)},{int(page_id)},'{_escape(event)}','{_escape(body_html)}' )
        """
    )


def remove_notebook_owner(notebook_id: int, user_login: str) -> None:
    """Удалить пользователя из владельцев книги."""
    run_execute(
        f"""
        DELETE FROM {OWNERS_TABLE}
        WHERE notebook_id = {int(notebook_id)}
          AND user_login = '{_escape(user_login)}'
        """
    )


def set_notebook_department(notebook_id: int, department_id: str | None) -> None:
    """
    Область видимости книги определяется department_id книги.

    - '00' или NULL -> видна всем
    - '01.01'       -> видна пользователям из '01.01' и ниже ('01.01.*')
    - '99'          -> видна только владельцам (owners)
    """
    dep_value = "NULL" if not department_id else f"'{_escape(department_id)}'"
    run_execute(
        f"""
        UPDATE {NOTEBOOKS_TABLE}
        SET department_id = {dep_value}, updated_at = NOW()
        WHERE id = {int(notebook_id)}
        """
    )


def is_notebook_owner(notebook_id: int, user_login: str) -> bool:
    result = run_scalar(
        f"""
        SELECT 1
        FROM {OWNERS_TABLE}
        WHERE notebook_id = {int(notebook_id)} AND user_login = '{_escape(user_login)}'
        LIMIT 1
        """
    )
    return bool(result)


def get_notebook_owners(notebook_id: int) -> pd.DataFrame:
    return run_fetch_df(
        f"""
        SELECT o.user_login AS login, u.full_name
        FROM {OWNERS_TABLE} o
        LEFT JOIN {USERS_TABLE} u ON u.login = o.user_login
        WHERE o.notebook_id = {int(notebook_id)}
        ORDER BY COALESCE(u.full_name, o.user_login)
        """
    )


def get_notebooks(user_login: str, user_department_id: str | None) -> pd.DataFrame:
    """
    Возвращает книги, которые пользователь может хотя бы читать.

    Правила видимости:
      - department_id IS NULL или '00' -> видна всем
      - department_id = '99'           -> видна только владельцам книги (notes_notebook_owners)
      - department_id = 'X'            -> видна пользователям из 'X' и всем подчинённым (X.*)
      - владелец книги видит её всегда
    """
    ulogin = _escape(user_login)
    user_dep = _escape(user_department_id or "")

    owner_condition = f"""
        n.id IN (
            SELECT notebook_id
            FROM {OWNERS_TABLE}
            WHERE user_login = '{ulogin}'
        )
    """

    if user_dep:
        dept_visibility_non99 = f"""
            (
                n.department_id IS NULL OR n.department_id = '00'
                OR '{user_dep}' = n.department_id
                OR '{user_dep}' LIKE (n.department_id || '.%')
            )
        """
    else:
        dept_visibility_non99 = " (n.department_id IS NULL OR n.department_id = '00') "

    return run_fetch_df(
        f"""
        SELECT n.id,
               n.name,
               n.department_id,
               n.created_at,
               n.updated_at,
               n.created_by
        FROM {NOTEBOOKS_TABLE} n
        WHERE
            (n.department_id = '99' AND ({owner_condition}))
            OR
            ((n.department_id IS NULL OR n.department_id <> '99')
                AND ( {owner_condition} OR {dept_visibility_non99} )
            )
        ORDER BY n.name
        """
    )


def get_owned_notebooks(user_login: str) -> pd.DataFrame:
    return run_fetch_df(
        f"""
        SELECT n.id,
               n.name,
               n.department_id,
               n.created_at,
               n.updated_at,
               n.created_by
        FROM {NOTEBOOKS_TABLE} n
        JOIN {OWNERS_TABLE} o
          ON o.notebook_id = n.id
        WHERE o.user_login = '{_escape(user_login)}'
        ORDER BY n.name
        """
    )


def get_sections(notebook_id: int | None) -> pd.DataFrame:
    query = f"""
        SELECT id, notebook_id, name, created_at, updated_at, created_by
        FROM {SECTIONS_TABLE}
    """
    if notebook_id:
        query += f" WHERE notebook_id = {int(notebook_id)}"
    query += " ORDER BY name"
    return run_fetch_df(query)


def load_pages_df(
    notebook_id: int | None,
    section_id: int | None,
    allowed_notebook_ids: list[int],
    search_text: str | None,
    search_tags_only: bool,
) -> pd.DataFrame:
    if not allowed_notebook_ids:
        return pd.DataFrame()

    # ✅ ДОБАВИЛИ p.status
    query = f"""
        SELECT
            p.id,
            p.title,
            p.status
        FROM {PAGES_TABLE} p
            JOIN {SECTIONS_TABLE} s ON p.section_id = s.id
            JOIN {NOTEBOOKS_TABLE} n ON s.notebook_id = n.id
        WHERE 1=1
    """

    allowed_csv = ", ".join(str(int(x)) for x in allowed_notebook_ids)
    query += f" AND n.id IN ({allowed_csv})"

    if section_id:
        query += f" AND s.id = {int(section_id)}"

    if search_text:
        if search_tags_only:
            # ✅ расширенный поиск (Google-like) по тегам (p.tag)
            query += build_advanced_search_where(search_text, mode="tags")
        else:
            # ✅ расширенный поиск по title/body_html
            query += build_advanced_search_where(search_text, mode="text")

    # query += " ORDER BY p.updated_at DESC, p.id DESC"
    query += " ORDER BY p.id DESC"

    # logger.info("PAGES SEARCH SQL:\n%s", query)

    df = run_fetch_df(query)
    return df if isinstance(df, pd.DataFrame) else pd.DataFrame()



def load_page_detail_df(page_id: int) -> pd.DataFrame:
    query = f"""
        SELECT
            p.id,
            p.title,
            p.status,
            p.tag,
            p.body_html,
            s.id AS section_id,
            s.name AS section_name,
            n.id AS notebook_id,
            n.name AS notebook_name,
            n.department_id AS notebook_department_id
        FROM {PAGES_TABLE} p
            JOIN {SECTIONS_TABLE} s ON p.section_id = s.id
            JOIN {NOTEBOOKS_TABLE} n ON s.notebook_id = n.id
        WHERE p.id = {int(page_id)}
        LIMIT 1
    """
    df = run_fetch_df(query)
    return df if isinstance(df, pd.DataFrame) else pd.DataFrame()


def create_notebook(name: str, user_login: str, department_id: str | None) -> int:
    cleaned = name.strip() or "Новая книга"
    dept_value = f"'{_escape(department_id)}'" if department_id else "NULL"
    new_id = run_scalar(
        f"""
        INSERT INTO {NOTEBOOKS_TABLE} (name, created_by, department_id)
        VALUES ('{_escape(cleaned)}', '{_escape(user_login)}', {dept_value})
        RETURNING id
        """
    )
    if new_id is None:
        raise RuntimeError("Не удалось создать книгу")
    add_notebook_owner(int(new_id), user_login)
    return int(new_id)


def create_section(notebook_id: int, name: str, user_login: str) -> int:
    cleaned = name.strip() or "Новый раздел"
    new_id = run_scalar(
        f"""
        INSERT INTO {SECTIONS_TABLE} (notebook_id, name, created_by)
        VALUES ({int(notebook_id)}, '{_escape(cleaned)}', '{_escape(user_login)}')
        RETURNING id
        """
    )
    if new_id is None:
        raise RuntimeError("Не удалось создать раздел")
    return int(new_id)


def rename_section(section_id: int, new_name: str) -> None:
    cleaned = (new_name or "").strip() or "Новый раздел"
    run_execute(
        f"""
        UPDATE {SECTIONS_TABLE}
        SET name = '{_escape(cleaned)}', updated_at = NOW()
        WHERE id = {int(section_id)}
        """
    )


def get_section_pages_count(section_id: int) -> int:
    cnt = run_scalar(
        f"""
        SELECT COUNT(*)
        FROM {PAGES_TABLE}
        WHERE section_id = {int(section_id)}
        """
    )
    return int(cnt or 0)


def delete_section(section_id: int) -> None:
    run_execute(f"DELETE FROM {SECTIONS_TABLE} WHERE id = {int(section_id)}")


def create_page(section_id: int, user_login: str, title: str | None = None) -> int:
    page_title = (title or "").strip() or "Новая страница"
    new_id = run_scalar(
        f"""
        INSERT INTO {PAGES_TABLE} (section_id, title, tag, body_html, created_by)
        VALUES ({int(section_id)}, '{_escape(page_title)}', '', '', '{_escape(user_login)}')
        RETURNING id
        """
    )
    if new_id is None:
        raise RuntimeError("Не удалось создать страницу")
    return int(new_id)


def insert_page_with_content(section_id: int, title: str, body_html: str, user_login: str) -> int:
    raw = body_html or ""

    if _looks_like_imported_html(raw) or "mso-" in raw.lower():
        raw = strip_onenote_datetime_block(raw)
        raw = normalize_onenote_rich_html(raw)

    safe_html = sanitize_html_safe(raw)

    new_id = run_scalar(
        f"""
        INSERT INTO {PAGES_TABLE} (section_id, title, tag, body_html, created_by)
        VALUES ({int(section_id)},
                '{_escape(title.strip() or 'Untitled')}',
                '',
                '{_escape(safe_html)}',
                '{_escape(user_login)}')
        RETURNING id
        """
    )
    if new_id is None:
        raise RuntimeError("Не удалось импортировать страницу")
    return int(new_id)
STATUS_COLOR_OPTIONS = [
    ("#FFFFFF", "Без цвета"),
    ("#FFF59D", "Жёлтый"),
    ("#90CAF9", "Голубой"),
    ("#A5D6A7", "Зелёный"),
    ("#EF9A9A", "Красный"),
    ("#CE93D8", "Фиолетовый"),
]


def _normalize_status(val: str | None) -> str:
    v = (val or "").strip().upper()
    allowed = {c for c, _ in STATUS_COLOR_OPTIONS}
    if v not in allowed:
        return "#FFFFFF"
    return v


def update_page(page_id: int, title: str, body_html: str, tag: str, status: str) -> None:
    raw = body_html or ""
    status_norm = _normalize_status(status)

    if _looks_like_imported_html(raw) or "mso-" in raw.lower():
        raw = strip_onenote_datetime_block(raw)
        raw = normalize_onenote_rich_html(raw)

    safe_html = sanitize_html_safe(raw)

    run_execute(
        f"""
        UPDATE {PAGES_TABLE}
        SET title = '{_escape(title.strip() or 'Без названия')}',
            tag = '{_escape(tag)}',
            status = '{_escape(status_norm)}',
            body_html = '{_escape(safe_html)}',
            updated_at = NOW()
        WHERE id = {int(page_id)}
        """
    )



def delete_page(page_id: int) -> None:
    run_execute(f"DELETE FROM {PAGES_TABLE} WHERE id = {int(page_id)}")
    _bump_data_version()


def get_page_attachments(page_id: int) -> pd.DataFrame:
    params = {"page_id": int(page_id)}

    def _fetch() -> pd.DataFrame:
        return run_fetch_df(
            f"""
            SELECT id,
                   attachment_type,
                   file_name,
                   mime_type,
                   file_size,
                   url,
                   created_at,
                   created_by,
                   question_number
            FROM {ATTACHMENTS_TABLE}
            WHERE page_id = {int(page_id)}
            ORDER BY COALESCE(question_number,''), created_at DESC, id DESC
            """
        )

    return _cached_df("cache_page_attachments", params, _fetch)


def delete_attachment(attachment_id: int) -> None:
    run_execute(f"DELETE FROM {ATTACHMENTS_TABLE} WHERE id = {int(attachment_id)}")
    _bump_data_version()


def get_attachment_file(attachment_id: int) -> tuple[bytes, str, str] | None:
    df = run_fetch_df(
        f"""
        SELECT
            encode(file_data, 'base64') AS file_b64,
            file_name,
            COALESCE(mime_type, 'application/octet-stream') AS mime_type,
            file_size
        FROM {ATTACHMENTS_TABLE}
        WHERE id = {int(attachment_id)} AND attachment_type = 'file'
        LIMIT 1
        """
    )
    if df.empty:
        return None

    file_b64 = df.at[0, "file_b64"]
    file_name = str(df.at[0, "file_name"])
    mime_type = str(df.at[0, "mime_type"])
    file_size = df.at[0, "file_size"]

    if file_b64 is None:
        return None

    if isinstance(file_b64, (bytes, bytearray, memoryview)):
        file_b64 = file_b64.decode("ascii", errors="ignore")
    else:
        file_b64 = str(file_b64)

    file_b64 = file_b64.replace("\n", "").replace("\r", "").strip()

    try:
        data = base64.b64decode(file_b64, validate=False)
    except Exception as e:
        logger.error("Не удалось раскодировать base64 для вложения %s: %s", attachment_id, e)
        return None

    try:
        if file_size and int(file_size) != len(data):
            logger.warning(
                "Размер файла не совпадает: file_size=%s, len(data)=%s, attachment_id=%s",
                file_size,
                len(data),
                attachment_id,
            )
    except Exception:
        pass

    return data, file_name, mime_type



@st.dialog("Новое сообщение", width="large")
def email_message_dialog(page_id: int, page_title: str, page_html: str, sender_login: str, page_path: str):
    recipients_df = list_email_recipients()
    recipient_options = list(recipients_df.itertuples(index=False)) if not recipients_df.empty else []
    sender_full_name, sender_job_title = get_user_signature(sender_login)

    # если ранее попросили сбросить поля добавления получателя — сделаем это ДО создания виджетов
    reset_flag_key = f"email_reset_fields_{page_id}"
    add_expanded_key = f"email_add_recipient_expanded_{page_id}"
    if st.session_state.pop(reset_flag_key, False):
        st.session_state[f"email_new_email_{page_id}"] = ""
        st.session_state[f"email_new_fio_{page_id}"] = ""
        st.session_state[f"email_new_gender_{page_id}"] = "муж"
        st.session_state[f"email_new_job_title_{page_id}"] = ""
        st.session_state[f"email_new_salutation_{page_id}"] = ""
        st.session_state[add_expanded_key] = False
    def _recipient_label(row) -> str:
        fio = getattr(row, "fio", "") or ""
        email_addr = getattr(row, "email", "") or ""
        job_title = getattr(row, "job_title", "") or ""
        if fio and email_addr:
            label = f"{fio} <{email_addr}>"
        else:
            label = fio or email_addr
        if job_title:
            label = f"{label} ({job_title})"
        return label

    def _attachment_label(row) -> str:
        size_txt = _format_file_size(getattr(row, "file_size", None))
        suffix = f" ({size_txt})" if size_txt else ""
        return f"{row.file_name}{suffix}"

    attachments_df = get_page_attachments(page_id)
    file_records = [
        row for row in attachments_df.itertuples(index=False) if getattr(row, "attachment_type", "") == "file"
    ]
    link_records = [
        row
        for row in attachments_df.itertuples(index=False)
        if getattr(row, "attachment_type", "") == "link" and getattr(row, "url", None)
    ]

    def _build_links_block() -> str:
        if not link_records:
            return ""
        lines = []
        for row in link_records:
            title = (getattr(row, "file_name", "") or "").strip()
            url = (getattr(row, "url", "") or "").strip()
            if not url:
                continue
            if title:
                lines.append(f"- {title}: {url}")
            else:
                lines.append(f"- {url}")
        if not lines:
            return ""
        return "Ссылки:\n" + "\n".join(lines)

    subject_default = page_title or f"Страница {page_id}"
    body_plain = _html_to_plain_for_email(page_html or "")
    greeting = "Добрый день!"
    links_block = _build_links_block()
    page_info = f"{page_path}".strip()
    page_link = build_page_deeplink(page_id)
    signature = f"С уважением!\n{sender_full_name}\n{sender_job_title}".rstrip()

    body_parts = [greeting]
    if body_plain:
        body_parts.append(body_plain)
    if links_block:
        body_parts.append(links_block)
    if page_path:
        body_parts.append(page_info)
        if page_link:
            body_parts.append(page_link)
    body_parts.append(signature)
    body_default = "\n\n".join(part for part in body_parts if part)

    with st.form(f"email_form_{page_id}", clear_on_submit=False):
        subject = st.text_input("Тема", value=subject_default)

        if not recipient_options:
            st.caption("Справочник получателей пуст.")

        to_dir = st.multiselect(
            "Получатели (справочник)",
            options=recipient_options,
            format_func=_recipient_label,
            key=f"email_to_dir_{page_id}",
        )
        cc_dir = st.multiselect(
            "Копия (справочник)",
            options=recipient_options,
            format_func=_recipient_label,
            key=f"email_cc_dir_{page_id}",
        )

        bcc_dir = st.multiselect(
            "Скрытая копия (справочник)",
            options=recipient_options,
            format_func=_recipient_label,
            key=f"email_bcc_dir_{page_id}",
        )

        body = st.text_area("Тело письма", value=body_default, height=200)
        important = st.checkbox("Важное", value=False)

        if file_records:
            selected_files = st.multiselect(
                "Вложения страницы",
                options=file_records,
                format_func=_attachment_label,
                key=f"email_files_{page_id}",
            )
        else:
            st.caption("Вложения страницы отсутствуют.")
            selected_files = []

        send_col, cancel_col, add_col = st.columns([1, 1, 2])
        with send_col:
            submitted = st.form_submit_button("Отправить")
        with cancel_col:
            canceled = st.form_submit_button("Отмена")
        with add_col:
            with st.expander("Добавить получателя", expanded=st.session_state.get(add_expanded_key, False)):
                new_email = st.text_input("eMail", key=f"email_new_email_{page_id}")
                new_fio = st.text_input("ФИО", key=f"email_new_fio_{page_id}")
                new_gender = st.selectbox(
                    "Пол",
                    options=["муж", "жен"],
                    key=f"email_new_gender_{page_id}",
                )
                new_job_title = st.text_input("Должность", key=f"email_new_job_title_{page_id}")
                new_salutation = st.text_input("Как обращаться", key=f"email_new_salutation_{page_id}")
                add_recipient = st.form_submit_button("Добавить")

    if canceled:
        st.session_state["email_dialog_page_id"] = None
        st.rerun()
    elif add_recipient:
        email_trim = new_email.strip().lower()
        fio_trim = new_fio.strip()

        if not email_trim or not fio_trim:
            st.error("Заполните eMail и ФИО.")
            return

        if not _is_valid_email(email_trim):
            st.error("Некорректный формат eMail.")
            return

        exists = run_scalar(
            f"""
            SELECT 1
            FROM {EMAIL_RECIPIENTS_TABLE}
            WHERE lower(email) = lower('{_escape(email_trim)}')
            LIMIT 1
            """
        )
        if exists:
            st.error("eMail уже существует.")
            return

        run_execute(
            f"""
            INSERT INTO {EMAIL_RECIPIENTS_TABLE}
                (email, fio, job_title, salutation, gender, created_by)
            VALUES
                ('{_escape(email_trim)}',
                 '{_escape(fio_trim)}',
                 '{_escape(new_job_title.strip())}',
                 '{_escape(new_salutation.strip())}',
                 '{_escape(new_gender)}',
                 '{_escape(sender_login)}')
            """
        )
        _notebook_id,_section_id = get_nb_section_id(page_id)
        add_event_log(topic="RECIPIENT", subtopic="CREATE",notebook_id=_notebook_id, section_id=_section_id, page_id=page_id, event="create_recipient: fio = " + _escape(fio_trim) + "; email = " + _escape(email_trim), body_html="")
        st.success("Получатель добавлен")
        # очистить поля ввода и свернуть экспандер (на следующем ререндере)
        st.session_state[f"email_reset_fields_{page_id}"] = True
        st.session_state["email_dialog_page_id"] = page_id
        st.rerun()
    elif submitted:
        to_list = [r.email for r in to_dir if getattr(r, "email", None)]
        cc_list = [r.email for r in cc_dir if getattr(r, "email", None)]
        bcc_list = [r.email for r in bcc_dir if getattr(r, "email", None)]

        recipients = _merge_email_lists(to_list)
        cc = _merge_email_lists(cc_list)
        bcc = _merge_email_lists(bcc_list)

        if not recipients:
            st.error("Укажите получателей.")
            return

        temp_paths: list[str] = []
        try:
            base_dir = os.path.dirname(__file__)
            temp_root = os.path.join(base_dir, config.temp_attachments_dir)
            os.makedirs(temp_root, exist_ok=True)
            with tempfile.TemporaryDirectory(prefix="notes_email_", dir=temp_root) as tmpdir:
                for row in selected_files:
                    payload = get_attachment_file(int(row.id))
                    if not payload:
                        st.warning(
                            f"Не удалось подгрузить вложение: {row.file_name}"
                        )
                        continue
                    data, file_name, _mime = payload
                    safe_name = _sanitize_attachment_filename(file_name)
                    file_path = os.path.join(tmpdir, safe_name)
                    with open(file_path, "wb") as handle:
                        handle.write(data)
                    temp_paths.append(file_path)

                body_text = body or ""
                if not body_text.strip():
                    body_text = greeting

                if not body_text.lstrip().startswith(greeting):
                    body_text = f"{greeting}\n\n{body_text.lstrip()}"

                links_block = _build_links_block()
                page_info = f"{page_path}".strip()
                signature = f"С уважением!\n{sender_full_name}\n{sender_job_title}".rstrip()

                if links_block and "Ссылки:" not in body_text:
                    if "С уважением!" in body_text:
                        body_text = body_text.replace("С уважением!", f"{links_block}\n\nС уважением!", 1)
                    else:
                        body_text = f"{body_text.rstrip()}\n\n{links_block}"

                def _insert_before_signature(text: str, insert: str) -> str:
                    if not insert:
                        return text
                    if "С уважением!" in text:
                        return text.replace("С уважением!", f"{insert}\n\nС уважением!", 1)
                    return f"{text.rstrip()}\n\n{insert}"

                if page_path and page_info not in body_text:
                    block = page_info
                    page_link = build_page_deeplink(page_id)
                    if page_link:
                        block = f"{block}\n{page_link}"
                    body_text = _insert_before_signature(body_text, block)
                else:
                    page_link = build_page_deeplink(page_id)
                    if page_link and page_link not in body_text:
                        body_text = _insert_before_signature(body_text, page_link)

                if "С уважением!" not in body_text:
                    body_text = f"{body_text.rstrip()}\n\n{signature}"

                send_mail(
                    subject=subject or "",
                    recipients=recipients,
                    cc=cc,
                    bcc=bcc,
                    body=body_text,
                    important=bool(important),
                    files=temp_paths,
                )

            st.success("Сообщение отправлено")
            st.session_state["email_dialog_page_id"] = None
            st.rerun()
        except Exception as exc:
            st.error(f"Ошибка отправки: {exc}")


def save_file_attachment(page_id: int, uploaded_file, user_login: str, question_number: str | None = None) -> None:
    if uploaded_file is None:
        return
    content = uploaded_file.getvalue()
    if not content:
        raise ValueError("Файл пустой или не удалось прочитать.")
    mime_type = uploaded_file.type or "application/octet-stream"
    encoded = base64.b64encode(content).decode("ascii")
    question_number_clean = (question_number or "").strip()
    question_number_sql = f"'{_escape(question_number_clean)}'" if question_number_clean else "NULL"
    run_execute(
        f"""
        INSERT INTO {ATTACHMENTS_TABLE}
            (page_id, attachment_type, file_name, mime_type, file_size, file_data, created_by, question_number)
        VALUES
            ({int(page_id)},
             'file',
             '{_escape(uploaded_file.name)}',
             '{_escape(mime_type)}',
             {len(content)},
             decode('{encoded}', 'base64'),
             '{_escape(user_login)}',
             {question_number_sql})
        """
    )
    _notebook_id,_section_id = get_nb_section_id(int(page_id))
    add_event_log(topic="ATTACHMENT", subtopic="UPLOAD",notebook_id=_notebook_id, section_id=_section_id, page_id=int(page_id), event="load_attachment: attachment_type = 'file'; file_name = " + _escape(uploaded_file.name) + ": file_size = " + str(len(content)) , body_html="")
    _bump_data_version()

def save_link_attachment(page_id: int, url: str, title: str, user_login: str, question_number: str | None = None) -> None:
    cleaned_url = (url or "").strip()
    if not cleaned_url:
        raise ValueError("URL не указан.")
    low = cleaned_url.lower()
    if low.startswith(("javascript:", "vbscript:", "data:text/html", "data:application/xhtml+xml")):
        raise ValueError("Запрещённый URL (возможная XSS-атака).")

    name = (title or "").strip() or cleaned_url
    question_number_clean = (question_number or "").strip()
    question_number_sql = f"'{_escape(question_number_clean)}'" if question_number_clean else "NULL"
    run_execute(
        f"""
        INSERT INTO {ATTACHMENTS_TABLE}
            (page_id, attachment_type, file_name, url, created_by, question_number)
        VALUES
            ({int(page_id)},
             'link',
             '{_escape(name)}',
             '{_escape(cleaned_url)}',
             '{_escape(user_login)}',
             {question_number_sql})
        """
    )
    _notebook_id,_section_id = get_nb_section_id(int(page_id))
    add_event_log(topic="ATTACHMENT", subtopic="UPLOAD",notebook_id=_notebook_id, section_id=_section_id, page_id=int(page_id), event="load_attachment: attachment_type = 'link'; name = " + _escape(name), body_html=_escape(cleaned_url))
    _bump_data_version()


def html_to_body(text: str, fallback_title: str):
    soup = BeautifulSoup(text, "html.parser")
    title = soup.title.string.strip() if soup.title and soup.title.string else fallback_title
    body = str(soup.body or soup)
    return title, body


def _split_onenote_html_into_pages(soup: BeautifulSoup, filename: str):
    base_title = filename.rsplit(".", 1)[0]
    page_divs = soup.find_all("div", style=lambda v: v and "border-width:100%" in v)

    pages: list[tuple[str, str]] = []

    if not page_divs:
        title, body_html = html_to_body(str(soup), base_title)
        pages.append((title, body_html))
        return pages

    for idx, div in enumerate(page_divs, start=1):
        # Заголовок страницы OneNote обычно в первом <p> (часто внутри <a>)
        title_p = div.find("p")
        title_text = title_p.get_text(" ", strip=True) if title_p else ""
        if not title_text:
            title_text = f"{base_title} {idx}"

        # ✅ ВАЖНО: НЕ удаляем заголовок/дату (как раньше делал _strip_onenote_header),
        # а нормализуем их wrapper-div, чтобы не было "узкого" переноса.
        _unwrap_onenote_header_wrappers(div)
        _remove_onenote_datetime(div)

        # ✅ СОХРАНЯЕМ ПУСТЫЕ СТРОКИ ДО НОРМАЛИЗАЦИИ
        _onenote_keep_blank_lines_in_div(div)

        # дальше как раньше: нормализуем страницу от мусорных обёрток/пустот
        _normalize_onenote_page_div(div)


        body_html = str(div)
        pages.append((title_text, body_html))

    return pages



def _is_onenote_datetime_div(div: Tag) -> bool:
    if not isinstance(div, Tag):
        return False
    ps = div.find_all("p")
    if not ps:
        return False
    text = div.get_text(" ", strip=True)
    if len(text) > 40:
        return False
    for p in ps:
        style = (p.get("style") or "").lower()
        if "font-size:10" not in style and "font-size:9" not in style and "font-size:11" not in style:
            return False
        if "color:#767676" not in style and "color: #767676" not in style:
            return False
    return True


def _is_onenote_datetime_p(p: Tag) -> bool:
    if not isinstance(p, Tag):
        return False
    text = p.get_text(" ", strip=True)
    if not text or len(text) > 40:
        return False
    style = (p.get("style") or "").lower()
    if "font-size:10" not in style and "font-size:9" not in style and "font-size:11" not in style:
        return False
    if "color:#767676" not in style and "color: #767676" not in style and "color:gray" not in style:
        return False
    has_time = re.search(r"\b\d{1,2}:\d{2}\b", text) is not None
    has_date = re.search(r"\b\d{1,2}\s+[A-Za-zА-Яа-я]+\s+\d{4}\b", text) is not None
    return bool(has_time or has_date)


def _remove_onenote_datetime(page_div: Tag) -> None:
    if not isinstance(page_div, Tag):
        return
    for div in list(page_div.find_all("div")):
        if _is_onenote_datetime_div(div):
            div.decompose()
    for p in list(page_div.find_all("p")):
        if _is_onenote_datetime_p(p):
            p.decompose()


def _unwrap_onenote_header_wrappers(page_div: Tag) -> None:
    """
    OneNote экспортирует заголовок/дату в узких wrapper-div с width:...in,
    из-за чего в редакторе/preview может быть плохой перенос.
    Здесь мы НЕ удаляем заголовок/дату, а превращаем эти div в нормальные <p>:
      <div style="...width:1.7in"><p>...</p></div>  ->  <p>...</p>
    Так сохраняется гиперссылка в заголовке (<a>...</a>) и даты/время.
    """
    if not isinstance(page_div, Tag):
        return

    # берем только прямых детей, т.к. header/date у OneNote обычно сверху
    for div in list(page_div.find_all("div", recursive=False)):
        if not isinstance(div, Tag):
            continue

        # если внутри нет вложенных div и есть только p (1..N) — это типичный wrapper
        inner_divs = div.find_all("div", recursive=False)
        if inner_divs:
            continue

        ps = div.find_all("p", recursive=False)
        if not ps:
            continue

        # wrapper должен быть "простым": кроме <p> и пробелов ничего
        ok = True
        for c in div.contents:
            if isinstance(c, NavigableString) and not str(c).strip():
                continue
            if isinstance(c, Tag) and c.name and c.name.lower() == "p":
                continue
            ok = False
            break
        if not ok:
            continue

        # разворачиваем wrapper: переносим <p> на уровень page_div
        insert_before = div
        for p in ps:
            # вынимаем p из div и вставляем перед div
            p.extract()
            insert_before.insert_before(p)

        # удаляем пустой wrapper
        div.decompose()



def _strip_onenote_header(page_div: Tag, title_p: Tag | None) -> None:
    # OneNote .mht export uses narrow wrappers for title/date that cause bad wrapping in the editor.
    if not isinstance(page_div, Tag):
        return

    title_text = ""
    if title_p and isinstance(title_p, Tag):
        title_text = title_p.get_text(" ", strip=True)

    for div in list(page_div.find_all("div")):
        if _is_onenote_datetime_div(div):
            div.decompose()
            continue

        if not title_text:
            continue
        ps = div.find_all("p", recursive=False)
        if len(ps) != 1:
            continue
        p = ps[0]
        p_text = p.get_text(" ", strip=True)
        if p_text != title_text:
            continue
        style = (p.get("style") or "").lower()
        if "font-size:20" in style or "font-size:19" in style or "font-family:\"calibri light\"" in style:
            div.decompose()



def _normalize_onenote_page_div(page_div: Tag) -> None:
    if not isinstance(page_div, Tag):
        return

    def _has_visible_content(tag: Tag) -> bool:
        if tag.find(["img", "br"]):
            return True
        text = tag.get_text("", strip=False)
        # OneNote often stores spaces inside language/formatting spans; keep them.
        if tag.name == "span" and text:
            return True
        return bool((text or "").strip())

    def _has_non_whitespace_text(tag: Tag) -> bool:
        for child in tag.contents:
            if isinstance(child, NavigableString) and child.strip():
                return True
        return False

    def _remove_whitespace_text_nodes(tag: Tag) -> None:
        for child in list(tag.contents):
            if isinstance(child, NavigableString) and not child.strip():
                # Preserve explicit spaces that OneNote splits into separate nodes.
                child.replace_with(" ")

    for tag in list(page_div.find_all(["div", "p", "span"])):
        if _has_visible_content(tag):
            continue
        tag.decompose()

    for tag in page_div.find_all(True):
        _remove_whitespace_text_nodes(tag)

    def _merge_wrapper_style(curr_style: str, child_style: str) -> str:
        if not child_style:
            return curr_style
        if not curr_style:
            return child_style
        if child_style in curr_style:
            return curr_style
        return f"{curr_style.rstrip(';')}; {child_style}"




# --- 2) Нормализация OneNote inline-style -> семантические теги (устойчиво к санитизации/рендеру) ---

def _parse_style_attr(style: str) -> dict[str, str]:
    out: dict[str, str] = {}
    s = (style or "").strip()
    if not s:
        return out
    # простейший разбор "k:v; k2:v2"
    parts = [p.strip() for p in s.split(";") if p.strip()]
    for p in parts:
        if ":" not in p:
            continue
        k, v = p.split(":", 1)
        k = k.strip().lower()
        v = v.strip()
        if k:
            out[k] = v
    return out


def _build_style_attr(style_map: dict[str, str]) -> str:
    # соберём обратно; порядок не критичен
    items = []
    for k, v in style_map.items():
        v2 = (v or "").strip()
        if v2:
            items.append(f"{k}:{v2}")
    return "; ".join(items)


def normalize_onenote_rich_html(html_in: str) -> str:
    """
    OneNote часто кладёт форматирование в <span style="...">:
      - font-weight:bold
      - text-decoration:underline
      - color:#000099
      - background:lime / mso-highlight:lime

    Чтобы форматирование НЕ терялось:
      - превращаем bold/italic/underline/strike в <strong>/<em>/<u>/<s>
      - подсветку переносим в <mark style="background-color:...">
      - mso-highlight / background -> background-color
    """
    if not html_in:
        return ""

    soup = BeautifulSoup(html_in, "html.parser")
    root = soup.body or soup

    def _wrap_children(node: Tag, wrapper_name: str) -> None:
        wrapper = soup.new_tag(wrapper_name)
        for ch in list(node.contents):
            wrapper.append(ch.extract())
        node.append(wrapper)

    # Идём по span (и иногда p), потому что OneNote активно кладёт стили туда
    for node in root.find_all(["span", "p"]):
        style_map = _parse_style_attr(node.get("style", ""))
        if not style_map:
            continue

        # --- подсветка: mso-highlight / background -> background-color ---
        # OneNote: background:lime; mso-highlight:lime
        if "mso-highlight" in style_map and "background-color" not in style_map:
            style_map["background-color"] = style_map.get("mso-highlight", "")
        if "background" in style_map and "background-color" not in style_map:
            style_map["background-color"] = style_map.get("background", "")

        # --- семантика ---
        make_bold = style_map.get("font-weight", "").strip().lower() in {"bold", "700", "800", "900"}
        make_italic = style_map.get("font-style", "").strip().lower() == "italic"
        tdec = style_map.get("text-decoration", "").strip().lower()

        make_underline = "underline" in tdec
        make_strike = ("line-through" in tdec) or ("strike" in tdec)

        # Удаляем преобразованные свойства, чтобы не зависеть от CSS в дальнейшем
        if "font-weight" in style_map and make_bold:
            style_map.pop("font-weight", None)
        if "font-style" in style_map and make_italic:
            style_map.pop("font-style", None)
        if "text-decoration" in style_map and (make_underline or make_strike):
            style_map.pop("text-decoration", None)

        # --- подсветку переносим в <mark> внутрь узла ---
        bg = (style_map.get("background-color") or "").strip()
        if bg:
            # уберём background/background-color/mso-highlight из style, а визуализацию сделаем mark-ом
            style_map.pop("background-color", None)
            style_map.pop("background", None)
            style_map.pop("mso-highlight", None)

            mark = soup.new_tag("mark")
            mark["style"] = f"background-color:{bg};"

            # перенесём текущие children внутрь mark
            for ch in list(node.contents):
                mark.append(ch.extract())
            node.append(mark)

        # применим обёртки (важно: оборачиваем содержимое node целиком)
        # порядок: strike -> underline -> italic -> bold (можно любой, визуально ок)
        if make_strike:
            _wrap_children(node, "s")
        if make_underline:
            _wrap_children(node, "u")
        if make_italic:
            _wrap_children(node, "em")
        if make_bold:
            _wrap_children(node, "strong")

        # восстановим style без “проблемных” свойств
        new_style = _build_style_attr(style_map)
        if new_style:
            node["style"] = new_style
        else:
            node.attrs.pop("style", None)

    return str(root)




def parse_mht_to_pages(data: bytes, filename: str):
    """
    Импорт OneNote .mht:
    - используем email (ОК) для корректного decode quoted-printable/base64
    - НЕ используем quopri и bleach
    - встраиваем ресурсы (img и пр.) как data: URL
    - делим на страницы
    - сохраняем форматирование: normalize_onenote_rich_html()
    - простая своя санитизация: sanitize_html_simple_keep_formatting()
    """
    msg = email.message_from_bytes(data)

    html_bytes: bytes | None = None
    html_charset: str = "utf-8"
    resources: list[tuple[str, bytes, str | None, str | None]] = []

    # 1) достаём HTML и ресурсы
    for part in msg.walk():
        ctype = (part.get_content_type() or "").lower()
        if ctype == "text/html" and html_bytes is None:
            html_charset = part.get_content_charset() or "utf-8"
            # ВАЖНО: decode=True -> email сам снимет quoted-printable (=3D, =\r\n) и base64
            html_bytes = part.get_payload(decode=True) or b""
        else:
            cid = part.get("Content-ID")
            loc = part.get("Content-Location")
            payload = part.get_payload(decode=True) or b""
            if (cid or loc) and payload:
                resources.append((ctype, payload, cid, loc))

    if not html_bytes:
        raise ValueError("В .mht не найден HTML-контент")

    try:
        html_part = html_bytes.decode(html_charset, errors="replace")
    except Exception:
        html_part = html_bytes.decode("utf-8", errors="replace")

    def norm(val: str) -> str:
        v = urllib.parse.unquote(val or "").strip()
        v = v.replace("\\", "/")
        if v.lower().startswith("cid:"):
            v = "cid:" + v[4:]
        return v

    # 2) строим карту ресурсов: cid/location -> data:...
    src_map: dict[str, str] = {}
    for ctype, content, cid, loc in resources:
        data_url = f"data:{ctype};base64,{base64.b64encode(content).decode('ascii')}"
        if cid:
            cid_clean = cid.strip().strip("<>").strip()
            # OneNote может ссылаться как "cid:XXXX" или просто "XXXX"
            for key in (cid_clean, f"cid:{cid_clean}", f"CID:{cid_clean}", norm(cid_clean), norm(f"cid:{cid_clean}")):
                src_map[key] = data_url

        if loc:
            loc_clean = loc.strip().strip("<>").strip()
            normalized = norm(loc_clean)
            for key in (loc_clean, normalized, f"cid:{loc_clean}", f"cid:{normalized}", f"CID:{loc_clean}"):
                src_map[key] = data_url
            basename = os.path.basename(normalized)
            if basename:
                for key in (basename, f"cid:{basename}", f"CID:{basename}", norm(basename)):
                    src_map[key] = data_url

    soup = BeautifulSoup(html_part, "html.parser")

    # 3) подмена src на data-url (картинки и т.п.)
    for tag in soup.find_all(src=True):
        src_val = tag.get("src", "")
        lookup = norm(src_val)

        if lookup in src_map:
            tag["src"] = src_map[lookup]
            continue

        # иногда OneNote даёт только basename
        base = os.path.basename(lookup)
        if base in src_map:
            tag["src"] = src_map[base]

    # 4) ссылки в новую вкладку
    for a in soup.find_all("a", href=True):
        a["target"] = "_blank"
        a["rel"] = "noopener noreferrer"

    # 5) делим на страницы (ваша функция)
    pages = _split_onenote_html_into_pages(soup, filename)

    safe_pages: list[tuple[str, str]] = []
    base_title = filename.rsplit(".", 1)[0]

    for title, body_html in pages:
        safe_title = (title or "").strip() or base_title

        # 0) сохранить пустые строки OneNote (&nbsp; -> <br>)
        body_html2 = preserve_onenote_blank_paragraphs(body_html or "")

        # 1) убрать дату и время из тела
        body_no_dt = strip_onenote_datetime_block(body_html2)

        # 1.1) ещё раз сохранить пустые строки (после вырезания дат/обёрток)
        body_no_dt = preserve_onenote_blank_paragraphs(body_no_dt)

        # 2) нормализация rich-formatting (bold/highlight/underline)
        body_norm = normalize_onenote_rich_html(body_no_dt)

        # 3) ваш текущий безопасный санитайзер
        safe_body = sanitize_html_safe(body_norm)

        safe_pages.append((safe_title, safe_body))
    return safe_pages




def _safe_filename(title: str, ext: str) -> str:
    base = (title or "page").strip()
    base = re.sub(r"[^\w\-. ]+", "_", base)
    if not base:
        base = "page"
    return f"{base}.{ext}"


def _norm_newlines_global(s: str) -> str:
    return (s or "").replace("\r\n", "\n").replace("\r", "\n")


def _indent_level_with_root(tag: Tag, root: Tag) -> int:
    cur: Tag | None = tag
    while isinstance(cur, Tag):
        classes = cur.get("class") or []
        for c in classes:
            m = re.match(r"ql-indent-(\d+)", str(c))
            if m:
                return int(m.group(1))
        if cur is root:
            break
        cur = cur.parent if isinstance(cur.parent, Tag) else None
    return 0


def _has_desc_block(tag: Tag, block_tags: set[str]) -> bool:
    for child in tag.find_all(list(block_tags), recursive=True):
        if child is not tag:
            return True
    return False


def _text_of_leaf(tag: Tag, _norm_newlines_global) -> str:
    parts: list[str] = []
    for elem in tag.descendants:
        if isinstance(elem, NavigableString):
            parts.append(str(elem))
    raw = "".join(parts)
    raw = _norm_newlines_global(raw).replace("\xa0", " ")
    return "\n".join(line.rstrip() for line in raw.split("\n"))


def _bump_data_version() -> None:
    st.session_state["data_version"] = int(st.session_state.get("data_version", 0)) + 1


def _cached_df(cache_key: str, params, fetcher):
    """
    Простое кэширование DataFrame в session_state с учетом data_version.
    """
    version = int(st.session_state.get("data_version", 0))
    cache = st.session_state.get(cache_key) or {}
    if cache.get("params") == params and cache.get("version") == version and cache.get("data") is not None:
        try:
            return cache["data"].copy()
        except Exception:
            pass

    df = fetcher()
    if df is None:
        df = pd.DataFrame()
    st.session_state[cache_key] = {"params": params, "version": version, "data": df}
    return df


def _decode_text_bytes(data: bytes) -> str:
    if not data:
        return ""
    # BOM для UTF-16 / UTF-8
    if data.startswith(b"\xff\xfe") or data.startswith(b"\xfe\xff"):
        return data.decode("utf-16").replace("\x00", "")
    if data.startswith(b"\xef\xbb\xbf"):
        return data.decode("utf-8-sig").replace("\x00", "")

    # эвристика: если много нулевых байт, это UTF-16
    if data.count(b"\x00") > max(2, len(data) // 20):
        try:
            return data.decode("utf-16").replace("\x00", "")
        except UnicodeDecodeError:
            pass

    # выбор между utf-8 и cp1251
    try:
        text_utf8 = data.decode("utf-8")
        # если нет странных замен, считаем utf-8
        if "\ufffd" not in text_utf8:
            return text_utf8.replace("\x00", "")
    except UnicodeDecodeError:
        text_utf8 = None

    try:
        return data.decode("cp1251").replace("\x00", "")
    except UnicodeDecodeError:
        pass

    for enc in ("utf-8-sig", "utf-16", "cp1251", "latin-1"):
        try:
            return data.decode(enc).replace("\x00", "")
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").replace("\x00", "")


def _html_escape_preserve_breaks(text: str) -> str:
    safe = html.escape(text or "")
    return safe.replace("\r\n", "\n").replace("\r", "\n").replace("\n", "<br>")


def _style_dataframe_table(table_html: str) -> str:
    if not table_html:
        return ""
    table_html = table_html.replace(
        '<table border="1" class="dataframe">',
        '<table style="border-collapse:collapse;width:100%;">',
    )
    table_html = table_html.replace(
        "<th>",
        '<th style="border:1px solid #d0d4da;padding:4px 6px;text-align:left;">',
    )
    table_html = table_html.replace(
        "<td>",
        '<td style="border:1px solid #d0d4da;padding:4px 6px;">',
    )
    return table_html


def _docx_bytes_to_html(data: bytes) -> str:
    if not data:
        return ""
    if not zipfile.is_zipfile(io.BytesIO(data)):
        raise ValueError("Файл не является корректным .docx (zip).")

    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        doc = None
    parts: list[str] = []

    try:
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except Exception:
        CT_Tbl = CT_P = Table = Paragraph = None  # type: ignore

    def _render_run(run) -> str:
        text = html.escape(run.text or "")
        if not text:
            return ""
        text = text.replace("\n", "<br>")
        tags: list[str] = []
        if run.bold:
            tags.append("strong")
        if run.italic:
            tags.append("em")
        if run.underline:
            tags.append("u")
        if run.font is not None:
            if run.font.superscript:
                tags.append("sup")
            elif run.font.subscript:
                tags.append("sub")
        for tag in tags:
            text = f"<{tag}>{text}</{tag}>"
        return text

    def _render_paragraph(paragraph) -> str:
        text = "".join(_render_run(run) for run in paragraph.runs)
        if not text:
            text = _html_escape_preserve_breaks(paragraph.text or "")
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            level = 1
            for num in "123456":
                if num in style_name:
                    level = int(num)
                    break
            return f"<h{level}>{text}</h{level}>"
        return f"<p>{text or '<br>'}</p>"

    def _render_table(table) -> str:
        rows_html: list[str] = []
        for row in table.rows:
            cells_html: list[str] = []
            for cell in row.cells:
                cell_text = _html_escape_preserve_breaks(cell.text or "")
                cells_html.append(
                    f'<td style="border:1px solid #d0d4da;padding:4px 6px;vertical-align:top;">{cell_text}</td>'
                )
            rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
        return (
            '<table style="border-collapse:collapse;width:100%;margin:6px 0;">'
            + "".join(rows_html)
            + "</table>"
        )

    if doc is not None:
        if CT_P and CT_Tbl:
            for child in doc.element.body.iterchildren():
                if isinstance(child, CT_P):
                    parts.append(_render_paragraph(Paragraph(child, doc)))
                elif isinstance(child, CT_Tbl):
                    parts.append(_render_table(Table(child, doc)))
        else:
            for paragraph in doc.paragraphs:
                parts.append(_render_paragraph(paragraph))
            for table in doc.tables:
                parts.append(_render_table(table))

        return "\n".join([p for p in parts if p])

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            xml_bytes = zf.read("word/document.xml")
        soup = BeautifulSoup(xml_bytes, "xml")
        body = soup.find(["w:body", "body"])
        if body is None:
            raise ValueError("Не найден body в .docx")

        def _strip_ns(name: str | None) -> str:
            return (name or "").split(":")[-1]

        def _twips_to_px(value: str | None) -> str | None:
            if not value:
                return None
            try:
                twips = float(value)
            except Exception:
                return None
            pt = twips / 20.0
            px = pt * 1.333
            return f"{px:.1f}px"

        def _run_style(run) -> str:
            rpr = run.find(["w:rPr", "rPr"], recursive=False)
            styles: list[str] = []
            if rpr:
                if rpr.find(["w:b", "b"]):
                    styles.append("font-weight:700;")
                if rpr.find(["w:i", "i"]):
                    styles.append("font-style:italic;")
                u = rpr.find(["w:u", "u"])
                if u is not None:
                    u_val = u.get("w:val") or u.get("val")
                    if not u_val or u_val != "none":
                        styles.append("text-decoration:underline;")
                color = rpr.find(["w:color", "color"])
                if color is not None:
                    val = color.get("w:val") or color.get("val")
                    if val and val.lower() != "auto":
                        styles.append(f"color: #{val};")
                highlight = rpr.find(["w:highlight", "highlight"])
                if highlight is not None:
                    val = (highlight.get("w:val") or highlight.get("val") or "").lower()
                    highlight_map = {
                        "yellow": "#fff59d",
                        "green": "#a5d6a7",
                        "cyan": "#80deea",
                        "magenta": "#f48fb1",
                        "blue": "#90caf9",
                        "red": "#ef9a9a",
                        "gray": "#e0e0e0",
                        "darkgray": "#bdbdbd",
                    }
                    if val in highlight_map:
                        styles.append(f"background-color: {highlight_map[val]};")
                size = rpr.find(["w:sz", "sz"])
                if size is not None:
                    val = size.get("w:val") or size.get("val")
                    try:
                        pt = float(val) / 2.0
                        styles.append(f"font-size:{pt:.1f}pt;")
                    except Exception:
                        pass
            return "".join(styles)

        def _render_runs(parent) -> str:
            fragments: list[str] = []

            def _collect_run_text(run) -> str:
                buf: list[str] = []
                for node in run.descendants:
                    name = _strip_ns(getattr(node, "name", None))
                    if name == "t":
                        buf.append(html.escape(node.get_text() or ""))
                    elif name == "tab":
                        buf.append("&emsp;")
                    elif name == "br":
                        buf.append("<br>")
                return "".join(buf)

            for child in parent.find_all(recursive=False):
                tag = _strip_ns(getattr(child, "name", None))
                if tag == "r":
                    text = _collect_run_text(child)
                    if not text:
                        continue
                    style = _run_style(child)
                    if style:
                        fragments.append(f'<span style="{style}">{text}</span>')
                    else:
                        fragments.append(text)
                elif tag == "hyperlink":
                    for run in child.find_all(["w:r", "r"], recursive=False):
                        text = _collect_run_text(run)
                        if not text:
                            continue
                        style = _run_style(run)
                        if style:
                            fragments.append(f'<span style="{style}">{text}</span>')
                        else:
                            fragments.append(text)
            return "".join(fragments)

        def _render_paragraph_xml(p) -> str:
            ppr = p.find(["w:pPr", "pPr"], recursive=False)
            styles: list[str] = []
            if ppr:
                jc = ppr.find(["w:jc", "jc"])
                if jc is not None:
                    align = (jc.get("w:val") or jc.get("val") or "").lower()
                    if align in ("center", "right", "left", "both", "justify"):
                        styles.append(f"text-align:{'justify' if align == 'both' else align};")
                ind = ppr.find(["w:ind", "ind"])
                if ind is not None:
                    left = _twips_to_px(ind.get("w:left") or ind.get("left"))
                    if left:
                        styles.append(f"margin-left:{left};")
                    first = _twips_to_px(ind.get("w:firstLine") or ind.get("firstLine"))
                    hanging = _twips_to_px(ind.get("w:hanging") or ind.get("hanging"))
                    if first:
                        styles.append(f"text-indent:{first};")
                    elif hanging:
                        styles.append(f"text-indent:-{hanging};")
                spacing = ppr.find(["w:spacing", "spacing"])
                if spacing is not None:
                    before = _twips_to_px(spacing.get("w:before") or spacing.get("before"))
                    after = _twips_to_px(spacing.get("w:after") or spacing.get("after"))
                    if before:
                        styles.append(f"margin-top:{before};")
                    if after:
                        styles.append(f"margin-bottom:{after};")
            text = _render_runs(p)
            style_attr = f' style="{"".join(styles)}"' if styles else ""
            return f"<p{style_attr}>{text or '<br>'}</p>"

        def _render_tbl(tbl) -> str:
            rows_html: list[str] = []
            for tr in tbl.find_all("w:tr", recursive=False):
                cells_html: list[str] = []
                for tc in tr.find_all("w:tc", recursive=False):
                    cell_parts: list[str] = []
                    for p in tc.find_all(["w:p", "p"], recursive=False):
                        cell_parts.append(_render_paragraph_xml(p))
                    cell_text = "".join(cell_parts) or "&nbsp;"
                    cells_html.append(
                        f'<td style="border:1px solid #d0d4da;padding:4px 6px;vertical-align:top;">{cell_text}</td>'
                    )
                rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
            return (
                '<table style="border-collapse:collapse;width:100%;margin:6px 0;">'
                + "".join(rows_html)
                + "</table>"
            )

        for child in body.find_all(recursive=False):
            if child.name in ("w:p", "p"):
                parts.append(_render_paragraph_xml(child))
            elif child.name in ("w:tbl", "tbl"):
                parts.append(_render_tbl(child))

        return "\n".join([p for p in parts if p])
    except Exception as exc:
        raise ValueError(f"Не удалось разобрать .docx: {exc}") from exc


def _excel_bytes_to_html(data: bytes) -> str:
    if not data:
        return ""
    from openpyxl import load_workbook

    def _color_to_hex(color) -> str | None:
        if color is None:
            return None
        rgb = getattr(color, "rgb", None)
        if isinstance(rgb, str) and rgb:
            return rgb[-6:]
        value = getattr(color, "value", None)
        if isinstance(value, str) and value:
            return value[-6:]
        return None

    def _cell_to_style(cell) -> str:
        styles: list[str] = []
        if cell.font is not None:
            if cell.font.bold:
                styles.append("font-weight:700;")
            if cell.font.italic:
                styles.append("font-style:italic;")
            if cell.font.underline:
                styles.append("text-decoration:underline;")
            if cell.font.color is not None:
                rgb = _color_to_hex(cell.font.color)
                if rgb:
                    styles.append(f"color: #{rgb};")
        if cell.fill is not None and cell.fill.patternType == "solid":
            rgb = _color_to_hex(cell.fill.fgColor)
            if rgb:
                styles.append(f"background-color: #{rgb};")
        if cell.alignment is not None:
            align = cell.alignment.horizontal
            if align in ("left", "center", "right", "justify"):
                styles.append(f"text-align:{align};")
            v_align = cell.alignment.vertical
            if v_align in ("top", "center", "bottom"):
                styles.append(f"vertical-align:{v_align};")
        return "".join(styles)

    wb = load_workbook(io.BytesIO(data), data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        max_row = ws.max_row or 0
        max_col = ws.max_column or 0
        if max_row == 0 or max_col == 0:
            continue

        non_empty_counts: list[tuple[int, int]] = []
        for row_idx in range(1, max_row + 1):
            count = 0
            for col_idx in range(1, max_col + 1):
                val = ws.cell(row=row_idx, column=col_idx).value
                if val is not None and str(val).strip() != "":
                    count += 1
            non_empty_counts.append((row_idx, count))

        header_idx, header_count = max(non_empty_counts, key=lambda x: x[1])
        if header_count == 0:
            continue

        header_cells = [ws.cell(row=header_idx, column=col_idx) for col_idx in range(1, max_col + 1)]
        header_values = [str(c.value).strip() if c.value is not None else "" for c in header_cells]
        keep_cols = [idx for idx, val in enumerate(header_values, start=1) if val and not val.startswith("Unnamed")]
        if not keep_cols:
            continue

        rows_html: list[str] = []
        header_html: list[str] = []
        for col_idx in keep_cols:
            cell = ws.cell(row=header_idx, column=col_idx)
            text = html.escape(str(cell.value).strip() if cell.value is not None else "")
            style = _cell_to_style(cell)
            style_attr = f' style="border:1px solid #d0d4da;padding:4px 6px;text-align:left;{style}"'
            header_html.append(f"<th{style_attr}>{text}</th>")
        rows_html.append("<tr>" + "".join(header_html) + "</tr>")

        for row_idx in range(header_idx + 1, max_row + 1):
            row_cells = [ws.cell(row=row_idx, column=col_idx) for col_idx in keep_cols]
            if all((cell.value is None or str(cell.value).strip() == "") for cell in row_cells):
                continue
            cell_html: list[str] = []
            for cell in row_cells:
                text = html.escape(str(cell.value) if cell.value is not None else "")
                style = _cell_to_style(cell)
                style_attr = f' style="border:1px solid #d0d4da;padding:4px 6px;{style}"'
                cell_html.append(f"<td{style_attr}>{text}</td>")
            rows_html.append("<tr>" + "".join(cell_html) + "</tr>")

        table_html = (
            '<table style="border-collapse:collapse;width:100%;">'
            + "".join(rows_html)
            + "</table>"
        )
        parts.append(f"<h3>{html.escape(sheet_name)}</h3>")
        parts.append(table_html)
    return "\n".join(parts)


def _csv_bytes_to_html(data: bytes) -> str:
    if not data:
        return ""
    try:
        df = pd.read_csv(io.BytesIO(data), dtype=str, sep=None, engine="python")
    except Exception:
        df = pd.read_csv(io.BytesIO(data), dtype=str, encoding="cp1251")
    df = df.fillna("")
    return _style_dataframe_table(df.to_html(index=False, escape=True))


def _text_bytes_to_html(data: bytes) -> str:
    text = _decode_text_bytes(data)
    safe = html.escape(text or "")
    return '<pre style="white-space:pre-wrap;font-family:monospace;">' + safe + "</pre>"


def _png_bytes_to_html(data: bytes) -> str:
    if not data:
        return ""
    b64 = base64.b64encode(data).decode("ascii")
    return (
        '<div style="text-align:center;">'
        f'<img src="data:image/png;base64,{b64}" style="max-width:100%;height:auto;" />'
        "</div>"
    )



def export_html_to_docx_bytes(html: str, title: str) -> io.BytesIO:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

    _XML_BAD = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")

    def _xml_safe(s: str) -> str:
        if not s:
            return ""
        s = s.replace("\xa0", " ")
        return _XML_BAD.sub("", s)



    def _norm_title(t: str) -> str:
        # убираем переносы и лишние пробелы, чтобы "Test 5" не разъезжал на 2 строки
        return re.sub(r"\s+", " ", (t or "").replace("\xa0", " ")).strip()

    def _set_paragraph_spacing(p):
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _add_hyperlink(paragraph, url: str, text: str, bold=False, italic=False, underline=True):
        """
        Создаёт настоящую hyperlink-ссылку в docx (а не просто текст).
        """
        if not url:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            run.underline = underline
            return run

        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Стиль Hyperlink (если есть) + подчёркивание
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)

        if bold:
            b = OxmlElement("w:b")
            rPr.append(b)
        if italic:
            i = OxmlElement("w:i")
            rPr.append(i)

        # underline
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single" if underline else "none")
        rPr.append(u)

        # (опционально) цвет как у ссылок
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "0000FF")
        rPr.append(c)

        r.append(rPr)

        t = OxmlElement("w:t")
        t.text = text
        r.append(t)

        hyperlink.append(r)
        paragraph._p.append(hyperlink)
        return hyperlink

    doc = Document()

    # Нормальные интервалы
    normal_style = doc.styles["Normal"]
    nf = normal_style.paragraph_format
    nf.space_before = Pt(0)
    nf.space_after = Pt(4)
    nf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    safe_title = _xml_safe(_norm_title(title))

    # если html уже похож на OneNote/Word и содержит "первый заголовок" — не дублируем
    add_heading = True
    try:
        first_p = (BeautifulSoup(html or "", "html.parser").find("p"))
        if first_p and (first_p.get_text(strip=True) or "").strip():
            # если title совпадает с первым p — не добавляем второй раз
            if (first_p.get_text(" ", strip=True) or "").strip() == (safe_title or "").strip():
                add_heading = False
    except Exception:
        pass

    if safe_title and add_heading:
        h = doc.add_heading(safe_title, level=1)
        _set_paragraph_spacing(h)


    soup = BeautifulSoup(html or "", "html.parser")
    body = soup.body or soup

    # <br> превращаем в безопасный маркер (НЕ \u0000)
    BR_TOKEN = "\uE000"
    for br in body.find_all("br"):
        br.replace_with(BR_TOKEN)

    def _normalize_text_for_docx(s: str) -> str:
        """
        Главный фикс для вашего кейса:
        OneNote/Word вставляет \n прямо в текст внутри <span>/<p>.
        В docx это превращается в переносы строк -> "по слову на строку".
        """
        if not s:
            return ""
        s = _xml_safe(s)

        # сохраняем BR_TOKEN, остальное \n/\r/\t превращаем в пробелы
        s = s.replace(BR_TOKEN, "\uE001")
        s = s.replace("\r\n", " ").replace("\r", " ").replace("\n", " ").replace("\t", " ")
        # сжимаем “служебные” пробелы от переносов
        s = re.sub(r"[ ]{2,}", " ", s)
        s = s.replace("\uE001", BR_TOKEN)
        return s

    def _parse_color(val: str) -> tuple[int, int, int] | None:
        if not val:
            return None
        raw = val.strip().lower()
        # Hex
        m_hex = re.match(r"#([0-9a-f]{6})", raw)
        if m_hex:
            hx = m_hex.group(1)
            return int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
        # rgb(...)
        m_rgb = re.match(r"rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", raw)
        if m_rgb:
            return int(m_rgb.group(1)), int(m_rgb.group(2)), int(m_rgb.group(3))

        names = {
            "black": (0, 0, 0),
            "white": (255, 255, 255),
            "gray": (128, 128, 128),
            "grey": (128, 128, 128),
            "lightgray": (211, 211, 211),
            "lightgrey": (211, 211, 211),
            "yellow": (255, 255, 0),
            "lime": (0, 255, 0),
            "green": (0, 128, 0),
            "red": (255, 0, 0),
            "blue": (0, 0, 255),
        }
        return names.get(raw)


    def _parse_style(
        style: str,
    ) -> tuple[
        float | None,
        tuple[int, int, int] | None,
        tuple[int, int, int] | None,
        bool | None,
        bool | None,
        bool | None,
        bool | None,
    ]:
        """
        Берём минимум из inline-style OneNote:
        font-size:10.0pt / color:gray / color:#2E74B5
        """
        if not style:
            return None, None, None, None, None, None, None
        st = style.lower()

        fs = None
        m = re.search(r"font-size\s*:\s*([0-9.]+)\s*pt", st)
        if m:
            try:
                fs = float(m.group(1))
            except Exception:
                fs = None

        col = None
        m2 = re.search(r"(?:^|;)\s*color\s*:\s*([^;]+)", st)
        if m2:
            col = _parse_color(m2.group(1))

        bg = None
        m5 = re.search(r"(?:^|;)\s*(background-color|background|mso-highlight)\s*:\s*([^;]+)", st)
        if m5:
            bg = _parse_color(m5.group(2))

        bold = None
        italic = None
        underline = None
        strike = None
        m8 = re.search(r"font-weight\s*:\s*([0-9]+|bold)", st)
        if m8:
            val = m8.group(1)
            bold = val == "bold" or (val.isdigit() and int(val) >= 600)
        if "font-style:italic" in st:
            italic = True
        if "text-decoration" in st:
            underline = "underline" in st
            strike = "line-through" in st or "strike" in st

        return fs, col, bg, bold, italic, underline, strike

    def _parse_paragraph_style(style: str) -> dict:
        if not style:
            return {}
        st = style.lower()
        out: dict[str, str] = {}
        m = re.search(r"text-align\s*:\s*([a-z]+)", st)
        if m:
            out["align"] = m.group(1)
        m = re.search(r"margin-left\s*:\s*([0-9.]+)\s*px", st)
        if m:
            out["margin_left_px"] = m.group(1)
        m = re.search(r"text-indent\s*:\s*([0-9.]+)\s*px", st)
        if m:
            out["text_indent_px"] = m.group(1)
        m = re.search(r"margin-top\s*:\s*([0-9.]+)\s*px", st)
        if m:
            out["margin_top_px"] = m.group(1)
        m = re.search(r"margin-bottom\s*:\s*([0-9.]+)\s*px", st)
        if m:
            out["margin_bottom_px"] = m.group(1)
        return out

    def _px_to_pt(px: str | None) -> float | None:
        if not px:
            return None
        try:
            return float(px) / 1.333
        except Exception:
            return None

    def _set_run_shading(run, rgb: tuple[int, int, int]) -> None:
        r, g, b = rgb
        fill = f"{r:02X}{g:02X}{b:02X}"
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        rPr.append(shd)

    def _apply_run_background(run, rgb: tuple[int, int, int]) -> None:
        highlight_map = {
            (255, 245, 157): WD_COLOR_INDEX.YELLOW,
            (165, 214, 167): WD_COLOR_INDEX.BRIGHT_GREEN,
            (0, 255, 0): WD_COLOR_INDEX.BRIGHT_GREEN,
            (255, 255, 0): WD_COLOR_INDEX.YELLOW,
            (144, 202, 249): WD_COLOR_INDEX.BLUE,
            (239, 154, 154): WD_COLOR_INDEX.RED,
            (206, 147, 216): WD_COLOR_INDEX.PINK,
            (224, 224, 224): WD_COLOR_INDEX.GRAY_25,
            (189, 189, 189): WD_COLOR_INDEX.GRAY_50,
        }
        highlight = highlight_map.get(rgb)
        if highlight is not None:
            run.font.highlight_color = highlight
        else:
            _set_run_shading(run, rgb)


    def _apply_run_format(run, bold=False, italic=False, underline=False, strike=False):
        run.bold = bool(bold)
        run.italic = bool(italic)
        run.underline = bool(underline)
        run.font.strike = bool(strike)



    def _walk_inline(
        node,
        paragraph,
        bold=False,
        italic=False,
        underline=False,
        strike=False,
        font_size_pt: float | None = None,
        font_rgb: tuple[int, int, int] | None = None,
        bg_rgb: tuple[int, int, int] | None = None,
    ):
        """
        Рекурсивный обход inline-узлов:
        - фикс OneNote переносов (\n -> пробел)
        - поддержка bold/italic/underline/strike
        - минимальная поддержка font-size/color из inline-style
        - ссылки -> real hyperlink
        """
        if isinstance(node, NavigableString):
            text = _normalize_text_for_docx(str(node))
            if not text:
                return

            parts = text.split(BR_TOKEN)
            for i, part in enumerate(parts):
                if part:
                    run = paragraph.add_run(part)
                    _apply_run_format(run, bold, italic, underline, strike)
                    if font_size_pt:
                        run.font.size = Pt(font_size_pt)
                    if font_rgb:
                        run.font.color.rgb = RGBColor(*font_rgb)
                    elif bg_rgb:
                        # Если есть фон, а цвет текста не задан, принудительно делаем текст чёрным для контраста.
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    if bg_rgb:
                        _apply_run_background(run, bg_rgb)
                if i < len(parts) - 1:
                    paragraph.add_run().add_break()
            return

        if not isinstance(node, Tag):
            return

        name = (node.name or "").lower()

        # перенос форматирования
        nb, ni, nu, ns = bold, italic, underline, strike
        if name in ("b", "strong"):
            nb = True
        if name in ("i", "em"):
            ni = True
        if name == "u":
            nu = True
        if name in ("s", "strike", "del"):
            ns = True

        # перенос стиля текста (font-size / color)
        cur_size, cur_col, cur_bg = font_size_pt, font_rgb, bg_rgb
        fs2, col2, bg2, b2, i2, u2, s2 = _parse_style(node.get("style") or "")
        if fs2:
            cur_size = fs2
        if col2:
            cur_col = col2
        if bg2:
            cur_bg = bg2
        if b2 is True:
            nb = True
        if i2 is True:
            ni = True
        if u2 is True:
            nu = True
        if s2 is True:
            ns = True

        # картинки
        if name == "img":
            src = (node.get("src") or "").strip()
            if not src:
                return
            try:
                img_bytes = None
                if src.startswith("data:image"):
                    _, b64data = src.split(",", 1)
                    img_bytes = base64.b64decode(b64data)
                elif src.startswith("http://") or src.startswith("https://"):
                    resp = requests.get(src, timeout=8)
                    resp.raise_for_status()
                    img_bytes = resp.content

                if img_bytes:
                    img_stream = io.BytesIO(img_bytes)
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(5))
            except Exception as e:
                logger.warning("Не удалось вставить изображение в DOCX: %s", e)
            return

        # гиперссылка
        if name == "a":
            href = (node.get("href") or "").strip()
            link_text = _normalize_text_for_docx(node.get_text("", strip=False))
            link_text = link_text.strip()
            if link_text:
                _add_hyperlink(paragraph, href, link_text, bold=nb, italic=ni, underline=True)
            return

        for child in node.children:
            _walk_inline(child, paragraph, nb, ni, nu, ns, cur_size, cur_col, cur_bg)




    def _add_paragraph_from_tag(tag: Tag, style: str | None = None):
        # Пустые абзацы из <p><br/></p> — это осознанные пустые строки. Сохраняем их как пустые параграфы.
        raw_txt = tag.get_text("", strip=False)
        if isinstance(raw_txt, str):
            br_only = raw_txt.replace("\xa0", "").strip() == BR_TOKEN
            txt_no_br = raw_txt.replace(BR_TOKEN, "").replace("\xa0", "").strip()
            if not txt_no_br and not tag.find(["img", "table"]):
                # количество пустых строк = количество <br> внутри параграфа (минимум 1)
                br_count = raw_txt.count(BR_TOKEN)
                for _ in range(max(1, br_count)):
                    p_empty = doc.add_paragraph(style=style)
                    _set_paragraph_spacing(p_empty)
                return

        p = doc.add_paragraph(style=style)
        _set_paragraph_spacing(p)
        styles = _parse_paragraph_style(tag.get("style") or "")
        if styles:
            if "align" in styles:
                align = styles["align"]
                align_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                if align in align_map:
                    p.alignment = align_map[align]
            pf = p.paragraph_format
            ml_pt = _px_to_pt(styles.get("margin_left_px"))
            if ml_pt is not None:
                pf.left_indent = Pt(ml_pt)
            ti_pt = _px_to_pt(styles.get("text_indent_px"))
            if ti_pt is not None:
                pf.first_line_indent = Pt(ti_pt)
            mt_pt = _px_to_pt(styles.get("margin_top_px"))
            if mt_pt is not None:
                pf.space_before = Pt(mt_pt)
            mb_pt = _px_to_pt(styles.get("margin_bottom_px"))
            if mb_pt is not None:
                pf.space_after = Pt(mb_pt)
        for child in tag.children:
            _walk_inline(child, p)

    def _add_pre(tag: Tag):
        # pre оставляем с переносами строк
        p = doc.add_paragraph()
        _set_paragraph_spacing(p)
        run = p.add_run(tag.get_text().replace("\xa0", " "))
        run.font.name = "Consolas"
        # чуть уменьшить, чтобы было похоже на код
        run.font.size = Pt(10)





    def _add_table_from_tag(table_tag: Tag):
        rows = table_tag.find_all("tr")
        if not rows:
            return
        max_cols = 0
        row_cells = []
        for tr in rows:
            cells = tr.find_all(["td", "th"], recursive=False)
            row_cells.append(cells)
            max_cols = max(max_cols, len(cells))

        t = doc.add_table(rows=len(rows), cols=max_cols)
        t.style = "Table Grid"
        for i, cells in enumerate(row_cells):
            for j, cell_tag in enumerate(cells):
                cell = t.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                _set_paragraph_spacing(p)
                for child in cell_tag.children:
                    _walk_inline(child, p)

    def _handle_block(node):
        if isinstance(node, NavigableString):
            # игнорируем пустые межтеговые переносы
            if not str(node).strip():
                return
            p = doc.add_paragraph(str(node).strip())
            _set_paragraph_spacing(p)
            return

        if not isinstance(node, Tag):
            return

        name = (node.name or "").lower()

        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(name[1])
            _add_paragraph_from_tag(node, style=f"Heading {min(level, 3)}")
            return

        if name == "p":
            _add_paragraph_from_tag(node)
            return

        if name == "pre":
            _add_pre(node)
            return

        if name == "blockquote":
            _add_paragraph_from_tag(node, style="Intense Quote")
            return

        if name in ("ul", "ol"):
            is_ul = (name == "ul")

            # старт для <ol> (учитываем value у первого li, если есть)
            num = 1
            for li in node.find_all("li", recursive=False):
                if not isinstance(li, Tag):
                    continue

                if not is_ul:
                    # OneNote часто ставит value="7" value="8" и дробит ol на куски
                    v = li.get("value")
                    if v is not None:
                        try:
                            num = int(str(v))
                        except Exception:
                            pass

                    p = doc.add_paragraph()
                    _set_paragraph_spacing(p)
                    # "7. " + текст
                    p.add_run(f"{num}. ")
                    for child in li.children:
                        _walk_inline(child, p)
                    num += 1
                else:
                    p = doc.add_paragraph()
                    _set_paragraph_spacing(p)
                    p.add_run("• ")
                    for child in li.children:
                        _walk_inline(child, p)
            return


        if name == "table":
            _add_table_from_tag(node)
            return

        # контейнеры — раскрываем внутрь
        for child in node.children:
            _handle_block(child)

    for child in body.children:
        _handle_block(child)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf





# =========================
# FIX: сохраняем отступы/пробелы для .sql экспорта + корректный текст для буфера
# =========================
def _html_to_plain_preserving_layout(html: str, indent_spaces: int = 4) -> str:
    if not html:
        return ""

    soup = BeautifulSoup(html, "html.parser")

    for br in soup.find_all("br"):
        br.replace_with("\n")

    body = soup.body or soup

    block_tags = {"p", "div", "li", "pre", "h1", "h2", "h3", "h4", "h5", "h6"}
    out_lines: list[str] = []

    blocks_all = body.find_all(list(block_tags))
    # берем "листовые" блоки: те, у которых нет вложенных блоков
    blocks = [t for t in blocks_all if not _has_desc_block(t, block_tags)]
    if blocks:
        for el in blocks:
            txt = _text_of_leaf(el, _norm_newlines_global)

            if txt.replace("\n", "").strip() == "":
                out_lines.append("")
                continue

            lvl = _indent_level_with_root(el, body)
            prefix = (" " * (lvl * indent_spaces)) if lvl > 0 else ""

            for line in txt.split("\n"):
                if line == "":
                    out_lines.append("")
                else:
                    out_lines.append(prefix + line)
    else:
        txt = _text_of_leaf(body, _norm_newlines_global)
        out_lines.extend(txt.split("\n"))

    result = "\n".join(out_lines)
    result = _norm_newlines_global(result)
    result = re.sub(r"\n{3,}", "\n\n", result).rstrip()
    return result


def _html_to_plain_for_email(html: str, indent_spaces: int = 4) -> str:
    """
    Plain-text для eMail:
    - нормализует "служебные" переносы OneNote/Word, чтобы не было "по слову на строку"
    - сохраняет ссылки из HTML как: "Текст ссылки - URL"
    - <br> -> перенос строки, <pre> -> сохраняем переносы
    """
    if not html:
        return ""

    soup = BeautifulSoup(html, "html.parser")

    for bad in soup.find_all(["script", "style"]):
        bad.decompose()

    # безопасный маркер переноса (не управляющий символ)
    br_token = "\uE000"
    for br in soup.find_all("br"):
        br.replace_with(br_token)

    body = soup.body or soup
    block_tags = {"p", "li", "pre", "h1", "h2", "h3", "h4", "h5", "h6"}

    def _normalize_inline_text(s: str) -> str:
        """
        Нормализация текста внутри обычных блоков:
        - br_token пока оставляем (потом превратим в \n)
        - любые \n/\t из исходного HTML считаем пробелами
        """
        s = _norm_newlines_global(s).replace("\xa0", " ")
        s = s.replace(br_token, "\n")  # реальные переносы от <br>
        # все прочие переносы/табуляции схлопываем
        s = s.replace("\n", "\n")  # уже "настоящие" переносы
        # но часто OneNote приносит \n как текстовые переносы - они уже схлопнутся ниже
        # (между словами будет пробел)
        lines = s.split("\n")
        cleaned_lines = []
        for line in lines:
            line = line.replace("\t", " ")
            line = re.sub(r"[ \t\f\v]+", " ", line).strip()
            cleaned_lines.append(line)
        # сохраняем пустые строки как пустые
        return "\n".join(cleaned_lines)

    def _replace_links(root: Tag) -> None:
        """
        <a href="URL">Text</a> -> Text - URL
        Делает это ДО get_text(), чтобы URL не потерялся.
        """
        for a in list(root.find_all("a", href=True)):
            href = (a.get("href") or "").strip()
            if not href:
                continue

            # Берём видимый текст ссылки (с учётом вложенных span/strong и т.д.)
            text = a.get_text(" ", strip=True)
            text = re.sub(r"[ \t\f\v]+", " ", (text or "").replace("\xa0", " ")).strip()

            # Если текста нет — используем URL как текст
            if not text:
                repl = href
            else:
                # если текст уже равен URL — не дублируем
                if text.strip().lower() == href.strip().lower():
                    repl = href
                else:
                    repl = f"{text} - {href}"

            a.replace_with(NavigableString(repl))

    def _text_of(tag: Tag) -> str:
        # pre: сохраняем переносы строк
        if tag.name and tag.name.lower() == "pre":
            # ссылки внутри pre тоже можно сохранить (редко, но бывает)
            tag_copy = BeautifulSoup(str(tag), "html.parser")
            pre_tag = tag_copy.find("pre") or tag_copy
            _replace_links(pre_tag)

            raw = pre_tag.get_text()
            raw = raw.replace("\xa0", " ")
            raw = raw.replace(br_token, "\n")
            raw = _norm_newlines_global(raw)
            return "\n".join(line.rstrip() for line in raw.split("\n")).rstrip()

        # обычный блок: подменяем ссылки на "Text - URL"
        tag_copy = BeautifulSoup(str(tag), "html.parser")
        root = tag_copy.find(tag.name) or tag_copy
        _replace_links(root)

        raw = root.get_text(separator=" ", strip=False)
        raw = raw.replace("\xa0", " ")
        raw = raw.replace(br_token, "\n")
        raw = _norm_newlines_global(raw)

        # Важно: служебные переносы/табуляции -> пробелы, кроме тех, что пришли от <br>
        # Для этого схлопнем whitespace внутри каждой строки отдельно.
        raw = _normalize_inline_text(raw)

        # подчистим пробелы вокруг переносов
        raw = re.sub(r" *\n *", "\n", raw)
        return raw.strip()

    blocks_all = body.find_all(list(block_tags))
    blocks = [t for t in blocks_all if not _has_desc_block(t, block_tags)]

    out_lines: list[str] = []
    for el in blocks:
        txt = _text_of(el)
        if txt is None:
            continue

        txt = _norm_newlines_global(txt)
        # если блок реально пустой
        if txt.strip() == "":
            out_lines.append("")
            continue

        lvl = _indent_level_with_root(el, body)
        prefix = (" " * (lvl * indent_spaces)) if lvl > 0 else ""

        if el.name and el.name.lower() == "li":
            for line in txt.split("\n"):
                line = line.strip()
                if not line:
                    continue
                out_lines.append(prefix + "- " + line)
            continue

        for line in txt.split("\n"):
            if line.strip() == "":
                out_lines.append("")
            else:
                out_lines.append(prefix + line.rstrip())

    result = "\n".join(out_lines)
    result = _norm_newlines_global(result)
    result = re.sub(r"\n{3,}", "\n\n", result).rstrip()
    return result


def _sql_text_from_html(html: str, title: str) -> str:
    plain = _html_to_plain_preserving_layout(html or "")
    header = f"-- {title or ''}".rstrip()
    return header + ("\n\n" + plain if plain else "\n")


def export_html_to_sql_bytes(html: str, title: str, encoding: str = "utf-8") -> bytes:
    full_text = _sql_text_from_html(html or "", title or "")
    return full_text.encode(encoding, errors="replace")


def render_copy_sql_button(sql_text: str, btn_key: str) -> None:
    render_copy_text_button(
        text=sql_text or "",
        btn_key=btn_key,
        button_label="📋 Копировать",
        ok_message="SQL скопирован в буфер обмена.",
    )


def render_copy_rich_button(
    html: str,
    plain: str,
    btn_key: str,
    button_label: str = "📋 Копировать",
    ok_message: str = "Скопировано",
    button_style: str | None = None,
    align: str = "flex-end",
    height: int = 42,
) -> None:
    """
    Копирует в буфер обмена как при выделении содержимого preview:
    - text/html (для Word) + text/plain (fallback)
    """
    # ВАЖНО: html должен быть уже sanitized
    html_payload = json.dumps(html or "", ensure_ascii=False)
    plain_payload = json.dumps(plain or "", ensure_ascii=False)
    safe_id = re.sub(r"[^a-zA-Z0-9_\-]", "_", btn_key)

    style = (
        button_style
        or """
            cursor:pointer;
            padding:0.35rem 0.75rem;
            border:1px solid rgba(49,51,63,.2);
            border-radius:0.5rem;
            background: white;
            font-size:14px;
            line-height:1.2;
        """
    )

    components.html(
        f"""
        <div style="display:flex; justify-content:{align}; align-items:center; width:100%; margin:0; padding:0;">
          <button id="copy_rich_btn_{safe_id}" style="{style}">{button_label}</button>
          <span id="copy_rich_msg_{safe_id}" style="margin-left:10px; font-size:13px;"></span>
        </div>

        <script>
        (function() {{
          const html = {html_payload};
          const plain = {plain_payload};

          const btn = document.getElementById("copy_rich_btn_{safe_id}");
          const msg = document.getElementById("copy_rich_msg_{safe_id}");

          function ok(message) {{
            if (!msg) return;
            msg.textContent = message || "Скопировано";
            msg.style.color = "green";
            setTimeout(() => {{ msg.textContent = ""; }}, 2500);
          }}

          function fail(message) {{
            if (!msg) return;
            msg.textContent = message || "Не удалось скопировать";
            msg.style.color = "crimson";
            setTimeout(() => {{ msg.textContent = ""; }}, 4000);
          }}

          async function copyRich() {{
            // 1) Современный путь: ClipboardItem (HTML + Text)
            try {{
              if (navigator.clipboard && window.ClipboardItem) {{
                const htmlBlob = new Blob([html], {{ type: "text/html" }});
                const textBlob = new Blob([plain], {{ type: "text/plain" }});
                const item = new ClipboardItem({{
                  "text/html": htmlBlob,
                  "text/plain": textBlob
                }});
                await navigator.clipboard.write([item]);
                ok({json.dumps(ok_message, ensure_ascii=False)});
                return;
              }}
            }} catch(e) {{
              // идём в fallback
            }}

            // 2) Fallback: "как выделили в preview" -> contentEditable + execCommand('copy')
            // Word обычно отлично подхватывает HTML из такого копирования
            try {{
              const host = document.createElement("div");
              host.setAttribute("contenteditable", "true");
              host.style.position = "fixed";
              host.style.left = "-9999px";
              host.style.top = "0";
              host.style.opacity = "0";
              host.style.pointerEvents = "none";

              // Добавим минимальную HTML-обёртку (Word иногда лучше распознаёт)
              host.innerHTML = "<div>" + html + "</div>";

              document.body.appendChild(host);

              const sel = window.getSelection();
              const range = document.createRange();
              range.selectNodeContents(host);
              sel.removeAllRanges();
              sel.addRange(range);

              const res = document.execCommand("copy");
              sel.removeAllRanges();
              document.body.removeChild(host);

              if (res) {{
                ok({json.dumps(ok_message, ensure_ascii=False)});
                return;
              }}
              fail("Браузер запретил копирование.");
            }} catch(e2) {{
              fail("Браузер запретил копирование.");
            }}
          }}

          if (btn) {{
            btn.addEventListener("click", function(ev) {{
              ev.preventDefault();
              ev.stopPropagation();
              copyRich();
            }});
          }}
        }})();
        </script>
        """,
        height=height,
        scrolling=False,
    )



def render_copy_text_button(
    text: str,
    btn_key: str,
    button_label: str = "📋 Копировать",
    ok_message: str = "Скопировано",
    button_style: str | None = None,
    align: str = "flex-end",
    height: int = 42,
) -> None:
    payload = json.dumps(text or "", ensure_ascii=False)
    safe_id = re.sub(r"[^a-zA-Z0-9_\-]", "_", btn_key)
    style = (
        button_style
        or """
            cursor:pointer;
            padding:0.35rem 0.75rem;
            border:1px solid rgba(49,51,63,.2);
            border-radius:0.5rem;
            background: white;
            font-size:14px;
            line-height:1.2;
        """
    )

    components.html(
        f"""
        <div style="display:flex; justify-content:{align}; align-items:center; width:100%; margin:0; padding:0;">
          <button id="copy_btn_{safe_id}" style="{style}">
            {button_label}
          </button>
          <span id="copy_msg_{safe_id}" style="margin-left:10px; font-size:13px;"></span>
        </div>

        <script>
        (function() {{
          const text = {payload};
          const btn = document.getElementById("copy_btn_{safe_id}");
          const msg = document.getElementById("copy_msg_{safe_id}");

          function ok(message) {{
            if (!msg) return;
            msg.textContent = message || "Скопировано";
            msg.style.color = "green";
            setTimeout(() => {{ msg.textContent = ""; }}, 2500);
          }}

          function fail(message) {{
            if (!msg) return;
            msg.textContent = message || "Не удалось скопировать";
            msg.style.color = "crimson";
            setTimeout(() => {{ msg.textContent = ""; }}, 4000);
          }}

          async function copy() {{
            try {{
              if (navigator.clipboard && navigator.clipboard.writeText) {{
                await navigator.clipboard.writeText(text);
                ok({json.dumps(ok_message, ensure_ascii=False)});
                return;
              }}
            }} catch (e) {{}}

            try {{
              const ta = document.createElement("textarea");
              ta.value = text;
              ta.setAttribute("readonly", "");
              ta.style.position = "fixed";
              ta.style.left = "-9999px";
              ta.style.top = "0";
              document.body.appendChild(ta);
              ta.select();
              const res = document.execCommand("copy");
              document.body.removeChild(ta);
              if (res) {{
                ok({json.dumps(ok_message, ensure_ascii=False)});
              }} else {{
                fail("Браузер запретил копирование.");
              }}
            }} catch (e2) {{
              fail("Браузер запретил копирование.");
            }}
          }}

          if (btn) {{
            btn.addEventListener("click", function(ev) {{
              ev.preventDefault();
              ev.stopPropagation();
              copy();
            }});
          }}
        }})();
        </script>
        """,
        height=height,
        scrolling=False,
    )


def build_page_deeplink(page_id: int) -> str:
    base = (getattr(config, "app_base_url", "") or "").strip()
    if not base:
        return ""

    parsed = urllib.parse.urlparse(base)
    existing = dict(urllib.parse.parse_qsl(parsed.query, keep_blank_values=True))
    existing["page_id"] = str(int(page_id))
    new_query = urllib.parse.urlencode(existing)
    return urllib.parse.urlunparse(parsed._replace(query=new_query))


def get_page_location(page_id: int) -> dict[str, int | str | None] | None:
    df = run_fetch_df(
        f"""
        SELECT
            p.id AS page_id,
            p.section_id AS section_id,
            s.notebook_id AS notebook_id,
            n.department_id AS department_id
        FROM {PAGES_TABLE} p
        JOIN {SECTIONS_TABLE} s ON p.section_id = s.id
        JOIN {NOTEBOOKS_TABLE} n ON s.notebook_id = n.id
        WHERE p.id = {int(page_id)}
        """
    )
    if df is None or df.empty:
        return None
    row = df.iloc[0]
    return {
        "page_id": int(row.get("page_id")),
        "section_id": int(row.get("section_id")),
        "notebook_id": int(row.get("notebook_id")),
        "department_id": None if row.get("department_id") is None else str(row.get("department_id")),
    }


def _close_edit_dialog_state():
    # закрыть редактор и убрать любые "авто-открывалки"
    st.session_state["edit_dialog_page_id"] = None
    st.session_state.pop("force_edit_page_id", None)
    st.session_state.pop("force_edit_page_id_once", None)
    st.session_state["dialog_open"] = False


@st.dialog("Редактирование страницы", width="large")
def edit_page_dialog(page_id_local: int, title: str, html_body: str, tag: str, status: str, default_edit_mode: bool = False):
    # помечаем, что активен диалог (чтобы не допустить вложенные modals)
    st.session_state["dialog_open"] = True
    st.markdown(
        """
        <style>
        div[data-testid="stDialog"] div[role="dialog"]{
            width: 96vw !important;
            max-width: 1400px !important;
        }
        div[data-testid="stDialog"] div[role="dialog"] > div{
            max-height: 90vh !important;
            overflow: auto !important;
        }
        div[data-testid="stDialog"] .ql-container{
            height: 60vh !important;
        }
        div[data-testid="stDialog"] .ql-editor{
            min-height: 60vh !important;
        }

        /* (опционально) скрыть крестик закрытия, чтобы не оставлять "висячие" состояния */
        div[data-testid="stDialog"] button[aria-label="Close"]{
            display: none !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    col_l, col_r = st.columns([2, 3])
    with col_l:
        new_title = st.text_input("Название страницы", value=title, key=f"dlg_title_{page_id_local}")


    status_options = [c for c, _ in STATUS_COLOR_OPTIONS]
    status_labels = {c: label for c, label in STATUS_COLOR_OPTIONS}
    current_status = _normalize_status(status)

    with col_r:
        tag_col, status_col, help_col = st.columns([4.8,1.5, 0.7])

        with tag_col:
            new_tag = st.text_input(
                "Теги",
                value=tag or "",
                key=f"dlg_tag_{page_id_local}",
                placeholder="Введите тег(и) через запятую без символа #",
            )

        with status_col:
            try:
                status_index = status_options.index(current_status)
            except ValueError:
                status_index = 0
            new_status = st.selectbox(
                "Статус",
                options=status_options,
                index=status_index,
                format_func=lambda v: status_labels.get(v, v),
                key=f"dlg_status_{page_id_local}",
            )

        with help_col:
            st.markdown("###")  # выравнивание кнопки по высоте
            with st.popover("❓", use_container_width=True):
                st.markdown(
                    """
    ### 📝 Справка по редактору (Quill)

    **Форматирование**
    - **B** — жирный, *I* — курсив, <u>U</u> — подчёркивание, ~~S~~ — зачёркивание  
    - x₂ — нижний индекс, x² — верхний индекс  

    **Цвет**
    - A — цвет текста  
    - A (с фоном) — цвет подсветки (фон)

    **Списки и выравнивание**
    - Маркированный / нумерованный список  
    - Выравнивание: слева / центр / справа / по ширине

    **Заголовки**
    - H1 / H2 — заголовки  
    - Normal — обычный текст

    **Дополнительно**
    - fx — формулы  
    - “ ” — цитата  
    - </> — код / блок кода  
    - Tx — очистить форматирование

    **Вставка**
    - 🔗 — ссылка  
    - 🖼 — изображение  

    ℹ️ Вставка из Word/Excel может приносить таблицы. HTML дополнительно очищается от опасных тегов.
    """
                )


    edit_col, col_2 = st.columns([1, 2])

    with edit_col:
        edit_mode = st.checkbox(
            "Редактировать содержимое",
            value=default_edit_mode,
            key=f"dlg_edit_mode_{page_id_local}",
        )

    editable_html = html_body or ""

    if edit_mode:
        quill_value = st_quill(
            value=editable_html,
            html=True,
            placeholder="Начните ввод текста...",
            key=f"dlg_quill_{page_id_local}",
        )
        quill_html = editable_html if quill_value is None else (quill_value or "")
    else:
        safe_view_html = sanitize_html_safe(editable_html)
        components.html(safe_view_html, height=420, scrolling=True)
        quill_html = editable_html

    c1, c2 = st.columns([1, 1])
    with c1:
        if st.button("Сохранить", key=f"dlg_save_{page_id_local}", use_container_width=True):
            update_page(page_id_local, new_title, quill_html, new_tag, new_status)

            _notebook_id,_section_id = get_nb_section_id(page_id_local)
            add_event_log(
                topic="PAGE",
                subtopic="UPGRADE",
                notebook_id=_notebook_id,
                section_id=_section_id,
                page_id=page_id_local,
                event="new_title = " + _escape(new_title) + "; new_tag = " + new_tag + "; new_status = " + new_status,
                body_html=quill_html,
            )
            st.success("Страница обновлена")

            # ✅ закрыть редактор и убрать любые авто-открывалки
            _close_edit_dialog_state()

            st.session_state["current_page_id"] = page_id_local
            st.session_state["force_page_id"] = page_id_local
            st.rerun()

    with c2:
        if st.button("Отмена", key=f"dlg_cancel_{page_id_local}", use_container_width=True):
            _close_edit_dialog_state()
            st.rerun()


# =========================
# ✅ STATUS column: renderer + dblclick cycle
# =========================
STATUS_VALUE_FORMATTER = JsCode(
    """
function(params) {
  return "●"; // рисуем кружок символом, без HTML/DOM
}
"""
)

STATUS_CELL_STYLE = JsCode(
    """
function(params) {
  const c = (params.value || "#FFFFFF").toString().trim();
  return {
    display: "flex",
    justifyContent: "center",
    alignItems: "center",
    // размер кружка = размер шрифта символа "●"
    fontSize: "27px",
    lineHeight: "27px",
    padding: "0",
    color: c,
    // лёгкая обводка/контраст на белом можно имитировать тенью:
    textShadow: "0 0 1px rgba(0,0,0,1)"
  };
}
"""
)



STATUS_DBLCLICK_HANDLER = JsCode(
    """
function(e) {
  if (!e || !e.colDef || e.colDef.field !== 'status') return;

const colors = ['#FFFFFF', '#FFF59D', '#90CAF9', '#A5D6A7', '#EF9A9A', '#CE93D8']; // белый, желтый, зеленый, красный, синий
  let cur = (e.data.status || '#FFFFFF').toString().trim().toUpperCase();

  const idx = colors.indexOf(cur);
  const next = colors[(idx >= 0 ? idx + 1 : 0) % colors.length];

  e.node.setDataValue('status', next);
  e.api.refreshCells({ rowNodes: [e.node], columns: ['status'], force: true });
}
"""
)



HELP_SEARCH_MD = r"""
## 🔎 Расширенный поиск

Поиск работает по **названию** и **содержимому** страницы, а если начать запрос с `#` то поиск работает по **тегам**.

### Основные правила
- **Пробел = AND** (пробел между словами в поиске - оба слова должны присутствовать в тексте/заголовке/теге)
- **`|` = OR** (символ | между словами в поиске - одно из слов должно присутствовать)
- **Скобки** `(...)` группируют выражения, т.е. можно сочитать условия AND, OR, NOT ...
- **`+`** — обязательное слово (аналогично просто слову, но явнее)
- **`-`** — исключить слово/фразу (NOT)
- **Кавычки** `"..."` — точная фраза
- **`*`** — любой набор символов (wildcard)

### Приоритеты (важно!)
`AND` сильнее чем `OR`, как в Google.

Пример:
`a | b c`  трактуется как `a OR (b AND c)`

### Примеры
- `(ноутбук | компьютер) + dell - asus`  
  → (ноутбук ИЛИ компьютер) И обязательно dell И НЕТ asus

- `"искусственный интеллект" | "машинное обучение"`  
  → одна из точных фраз

- `лучший*в*мире + ресторан - дорогой`  
  → wildcard + AND + NOT

### Поиск по тегам
- Чтобы искать **только по тегам**, начни запрос с `#`  
  Пример: `#1969 | 258`
"""


@st.dialog("Справка: Расширенный поиск", width="large")
def help_search_dialog() -> None:
    st.markdown(HELP_SEARCH_MD)


def main():
    st.set_page_config(
        layout="wide",
        page_title="ДФИП_Notes",
        page_icon="assets/favicon-stack-128.png",
        initial_sidebar_state="expanded",
    )

    cleanup_docs_dir()

    st.markdown(
        """
        <style>
        .block-container { padding-top: 0.4rem !important; }
        header[data-testid="stHeader"] { padding-top: 0; padding-bottom: 0; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <style>
        [data-testid="stSidebar"] { min-width: 350px; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    ensure_db_credentials()
    owned_notebooks_df = pd.DataFrame()

    # --- deep-link: открыть страницу по `?page_id=<id>` ---
    try:
        raw_page_id = st.query_params.get("page_id")
        if isinstance(raw_page_id, list):
            raw_page_id = raw_page_id[0] if raw_page_id else None
        qp_page_id = int(str(raw_page_id)) if raw_page_id not in (None, "", []) else None
    except Exception:
        qp_page_id = None

    if qp_page_id is not None:
        last_handled = st.session_state.get("deeplink_handled_page_id")
        if last_handled != qp_page_id:
            loc = get_page_location(qp_page_id)
            if loc:
                dept_id = loc.get("department_id")
                if dept_id:
                    st.session_state["force_department_id"] = dept_id
                st.session_state["force_notebook_id"] = int(loc["notebook_id"])
                st.session_state["force_section_id"] = int(loc["section_id"])
                st.session_state["current_page_id"] = int(loc["page_id"])
                st.session_state["force_page_id"] = int(loc["page_id"])
                st.session_state["page_search"] = ""
                st.session_state["deeplink_handled_page_id"] = int(loc["page_id"])
            else:
                st.warning(f"Страница с ID={qp_page_id} не найдена или недоступна.")
                st.session_state["deeplink_handled_page_id"] = qp_page_id

        # чтобы параметр не "фиксировал" навигацию при дальнейшем использовании приложения
        try:
            st.query_params.pop("page_id", None)
        except Exception:
            pass

    users_df = list_users()
    user_records = list(users_df.itertuples(index=False))
    user_map = {row.login: row.full_name for row in user_records}
    user_dept_map = {row.login: row.department_id for row in user_records}
    registered_users = {row.login for row in user_records}
    login_options = [row.login for row in user_records]

    stored_login = st.session_state.get("current_user_login")
    db_user = st.session_state.get("db_credentials", {}).get("user")
    preferred_login: str | None = stored_login or db_user or (login_options[0] if login_options else None)

    selected_login: str | None = preferred_login
    if selected_login:
        st.session_state["current_user_login"] = selected_login
    else:
        st.sidebar.info("Нет доступных пользователей.")
        return

    owned_notebooks_df = get_owned_notebooks(selected_login)

    departments_df = get_departments()
    department_map = {row.department_id: row.name_department for row in departments_df.itertuples(index=False)}

    user_dep_id = user_dept_map.get(selected_login)

    # --- Левый сайдбар: подразделение ---
    selected_department_id = "00"
    dept_records = list(departments_df.itertuples(index=False)) if not departments_df.empty else []
    is_user_department_selected = False
    forced_department_id = st.session_state.pop("force_department_id", None)

    if dept_records:
        if forced_department_id:
            forced_row = next((r for r in dept_records if str(r.department_id) == str(forced_department_id)), None)
            if forced_row is not None:
                st.session_state["department_selector"] = forced_row

        if "department_selector" not in st.session_state and not forced_department_id:
            has_owned_books = not owned_notebooks_df.empty
            default_dep_id = "99" if has_owned_books else "00"
            default_row = next((r for r in dept_records if str(r.department_id) == default_dep_id), None)
            if default_row is None:
                default_row = dept_records[0]
            st.session_state["department_selector"] = default_row

        selected_department = st.sidebar.selectbox(
            "Поиск по подразделению",
            dept_records,
            format_func=lambda row: row.name_department,
            key="department_selector",
        )
        selected_department_id = str(selected_department.department_id)
        user_dep_id = user_dept_map.get(selected_login)
        is_user_department_selected = bool(user_dep_id) and (str(selected_department_id) == str(user_dep_id))
    else:
        selected_department_id = "00"

    st.session_state["current_department_id"] = selected_department_id

    # --- поиск страниц ---
    def _clear_page_search():
        st.session_state["page_search"] = ""

    # сбрасываем поиск до отрисовки инпута (если была операция создания/импорта)
    if st.session_state.pop("reset_page_search_once", False):
        st.session_state["page_search"] = ""

    # уникальный префикс (можно оставить константой)
    SEARCH_UI_PREFIX = "pages_search_ui"

    search_col, clear_col = st.sidebar.columns([13, 3])

    with search_col:
        search_raw = st.text_input(
            label="",
            key="page_search",
                    placeholder="Начните ввод текста...",
            label_visibility="collapsed",
        )

    with clear_col:
        st.button(
            "🔄",
            key=f"{SEARCH_UI_PREFIX}_clear_btn",
            help="Очистить фильтр",
            on_click=_clear_page_search,
            use_container_width=True,
        )





    search_raw = (search_raw or "").strip()
    search_tags_only = search_raw.startswith("#")
    search_text = search_raw[1:].strip() if search_tags_only else search_raw

    current_user_can_create_notebook = selected_login in registered_users

    # --- список книг ---
    notebooks_df = _cached_df(
        "cache_notebooks",
        params=(selected_login, user_dep_id),
        fetcher=lambda: get_notebooks(selected_login, user_dep_id),
    )
    filtered_notebooks_df = notebooks_df.copy()
    # ✅ Все книги, доступные пользователю по правилам get_notebooks (owners/00/99/подразделения)
    visible_notebook_ids = notebooks_df["id"].astype(int).tolist() if not notebooks_df.empty else []
    owned_ids = set(owned_notebooks_df["id"].astype(int).tolist()) if not owned_notebooks_df.empty else set()
    current_department_id: str = st.session_state.get("current_department_id", "00")



    if current_department_id != "00" and not filtered_notebooks_df.empty:
        dep_col = filtered_notebooks_df["department_id"].fillna("00").astype(str)
        owned_mask = filtered_notebooks_df["id"].astype(int).isin(owned_ids)

        if str(current_department_id) == "99":
            mask = owned_mask
        else:
            prefix = str(current_department_id).strip()
            prefix_like = prefix + "."
            dept_mask = (dep_col == "00") | (dep_col == prefix) | dep_col.str.startswith(prefix_like)
            mask = dept_mask | owned_mask

        filtered_notebooks_df = filtered_notebooks_df[mask]

    selected_notebook_id: int | None = None
    selected_section_id: int | None = None
    selected_notebook_row: pd.Series | None = None
    can_edit_notebook: bool = False

    notebook_records = list(filtered_notebooks_df.itertuples(index=False))
    top_col1, top_col2, top_col3 = st.columns([5, 1, 5])

    # --- диалог "Новая книга" ---
    if current_user_can_create_notebook:

        @st.dialog("Новая книга", width="small")
        def new_notebook_dialog():
            new_nb_name = st.text_input("Название новой книги", key="new_notebook_name_modal")
            create_clicked = st.button("Создать книгу", key="create_notebook_btn_modal")
            if create_clicked:
                user_department_id = user_dept_map.get(selected_login)
                new_nb_id = create_notebook(new_nb_name, selected_login, user_department_id)
                if user_department_id:
                    st.session_state["force_department_id"] = user_department_id
                st.session_state["force_notebook_id"] = new_nb_id
                add_event_log(topic="NOTEBOOK", subtopic="CREATE",notebook_id=new_nb_id, section_id=0, page_id=0, event="notebook_name = " + new_nb_name + "; user_department_id = " + user_department_id, body_html="")
                st.success("Книга создана")
                _bump_data_version()
                st.rerun()

    # --- верхняя панель: книга + кнопки (ВЕРНУЛИ 🔐) ---
    with top_col1:
        select_col, plus_col, info_col = st.columns([14, 2, 2])

        with select_col:
            if notebook_records:
                forced_nb_id = st.session_state.pop("force_notebook_id", None)
                if forced_nb_id is not None:
                    target_row = next((r for r in notebook_records if int(r.id) == int(forced_nb_id)), None)
                    if target_row is not None:
                        st.session_state["notebook_selector"] = target_row

                def _nb_label(row):
                    if is_user_department_selected:
                        return f"{row.name}"
                    dept_id = getattr(row, "department_id", None)
                    dept_name = department_map.get(dept_id, "") if dept_id else ""
                    dept_txt = f" [{dept_name}]" if dept_name else ""
                    return f"{row.name}{dept_txt}"

                st.markdown("###### ")
                st.markdown("###### Записная книга")

                selected_notebook = st.selectbox(
                    label="",
                    options=notebook_records,
                    format_func=_nb_label,
                    key="notebook_selector",
                    label_visibility="collapsed",
                )

                selected_notebook_id = int(selected_notebook.id)
                selected_notebook_row = filtered_notebooks_df[filtered_notebooks_df["id"] == selected_notebook_id].iloc[0]
                can_edit_notebook = bool(is_notebook_owner(selected_notebook_id, selected_login))
            else:
                st.info("Нет доступных книг")




    # ✅ диалог "Права доступа на книгу" (вернули)
    if selected_notebook_id is not None and can_edit_notebook:

        @st.dialog("Права доступа на книгу/Переименование", width="small")
        def notebook_access_dialog():
            owners_df = get_notebook_owners(selected_notebook_id)
            owners_text = (
                ", ".join(f"{row.full_name or row.login} ({row.login})" for row in owners_df.itertuples(index=False))
                or "Нет владельцев"
            )

            dept_id = selected_notebook_row.get("department_id", None) if selected_notebook_row is not None else None
            dept_id = "00" if dept_id is None or str(dept_id).strip() == "" else str(dept_id)
            dept_value = department_map.get(dept_id, "не задано")

            st.caption(f"Подразделение (видимость): {dept_value}")
            st.caption(f"Владельцы: {owners_text}")

            owner_logins = set(owners_df["login"].tolist()) if not owners_df.empty else set()

            # ⬇️ текущее имя книги (для подсказки/плейсхолдера)
            current_book_name = ""
            try:
                current_book_name = str(selected_notebook_row.get("name", "") or "").strip()
            except Exception:
                current_book_name = ""

            with st.form(f"access_form_{selected_notebook_id}"):
                dept_records_all = list(departments_df.itertuples(index=False)) if not departments_df.empty else []
                current_nb_dept = str(dept_id or "00")

                current_dept_row = next(
                    (r for r in dept_records_all if str(r.department_id) == current_nb_dept),
                    None,
                )
                if current_dept_row is None and dept_records_all:
                    current_dept_row = dept_records_all[0]

                selected_dept_row = st.selectbox(
                    "Область видимости книги (подразделение)",
                    options=dept_records_all,
                    index=dept_records_all.index(current_dept_row) if current_dept_row in dept_records_all else 0,
                    format_func=lambda r: f"{r.name_department}",
                    key=f"nb_department_{selected_notebook_id}",
                )
                new_department_id = str(selected_dept_row.department_id) if selected_dept_row else "00"

                selectable_users = [login for login in login_options if login not in owner_logins]
                new_owner_login = st.selectbox(
                    "Добавить владельца книги",
                    options=[""] + selectable_users,
                    format_func=lambda login: "—" if login == "" else f"{user_map.get(login, login)} ({login})",
                    key=f"add_owner_{selected_notebook_id}",
                )

                removable_owners = [login for login in owner_logins if login != selected_login]
                remove_owner_login = st.selectbox(
                    "Удалить владельца книги",
                    options=[""] + removable_owners,
                    format_func=lambda login: "—" if login == "" else f"{user_map.get(login, login)} ({login})",
                    key=f"remove_owner_{selected_notebook_id}",
                )

                # ✅ НОВОЕ ПОЛЕ: переименование книги (внизу над кнопкой)
                # st.markdown("---")
                new_book_name = st.text_input(
                    "Переименовать книгу",
                    value="",
                    placeholder=(f"Текущее название: {current_book_name}" if current_book_name else "Введите новое название книги"),
                    key=f"rename_notebook_{selected_notebook_id}",
                )

                # ✅ кнопка переименована
                submitted = st.form_submit_button("Сохранить")

                if submitted:
                    # 1) видимость
                    set_notebook_department(selected_notebook_id, new_department_id)

                    # 2) владельцы
                    if new_owner_login:
                        add_notebook_owner(selected_notebook_id, new_owner_login)
                        add_event_log(topic="NOTEBOOK", subtopic="ADD_OWNER",notebook_id=selected_notebook_id, section_id=0, page_id=0, event="new_owner = " + new_owner_login, body_html="")
                    if remove_owner_login and remove_owner_login != selected_login:
                        remove_notebook_owner(selected_notebook_id, remove_owner_login)
                        add_event_log(topic="NOTEBOOK", subtopic="REMOVE_OWNER",notebook_id=selected_notebook_id, section_id=0, page_id=0, event="remove_owner = " + remove_owner_login, body_html="")
                    # 3) переименование (только если введено)
                    nb_new = (new_book_name or "").strip()
                    if nb_new and nb_new != current_book_name:
                        run_execute(
                            f"""
                            UPDATE {NOTEBOOKS_TABLE}
                            SET name = '{_escape(nb_new)}', updated_at = NOW()
                            WHERE id = {int(selected_notebook_id)}
                            """
                        )
                        add_event_log(topic="NOTEBOOK", subtopic="RENAME",notebook_id=selected_notebook_id, section_id=0, page_id=0, event="new_notebook_name = " + _escape(nb_new) + "; old_notebook_name = " + _escape(current_book_name), body_html="")
                    st.success("Сохранено")
                    st.rerun()

    if current_user_can_create_notebook:
        with plus_col:
            st.markdown("###### ")
            st.markdown("###### ")
            if st.button("➕", key="open_new_notebook_dialog", help="Создать новую книгу", use_container_width=True):
                new_notebook_dialog()

    # ✅ КНОПКА 🔐 
    if selected_notebook_id is not None and can_edit_notebook:
        with info_col:
            st.markdown("###### ")
            st.markdown("###### ")
            if st.button("🔐", key="open_notebook_access_dialog", help="Права доступа на книгу/Переименование", use_container_width=True):
                notebook_access_dialog()






    # --- диалог "Новый раздел" ---
    if selected_notebook_id is not None and can_edit_notebook:

        @st.dialog("Новый раздел", width="small")
        def new_section_dialog():
            new_section_name = st.text_input("Название раздела", key="new_section_name_modal")
            create_clicked = st.button("Создать раздел", key="create_section_btn_modal")
            if create_clicked:
                section_new_id=create_section(selected_notebook_id, new_section_name, selected_login)
                st.success("Раздел создан")
                _bump_data_version()
                add_event_log(topic="SECTION", subtopic="CREATE",notebook_id=selected_notebook_id, section_id=section_new_id, page_id=0, event="section_name = " + _escape(new_section_name), body_html="")
                st.rerun()

        @st.dialog("Переименовать или удалить раздел", width="small")
        def section_manage_dialog(section_row):
            section_id_local = int(section_row.id)
            st.caption(f"Текущий раздел: **{section_row.name}**")
            old_section_name = section_row.name
            new_name = st.text_input(
                "Новое название раздела",
                value=section_row.name,
                key=f"rename_section_name_{section_id_local}",
            )
            pages_cnt = get_section_pages_count(section_id_local)

            col_rename, col_delete = st.columns(2)
            with col_rename:
                if st.button("Сохранить название", key=f"btn_rename_section_{section_id_local}", use_container_width=True):
                    rename_section(section_id_local, new_name)
                    _bump_data_version()
                    add_event_log(topic="SECTION", subtopic="RENAME",notebook_id=selected_notebook_id, section_id=section_id_local, page_id=0, event="new_section_name = " + _escape(new_name) + "; old_saction_name = " + old_section_name, body_html="")
                    st.success("Название раздела обновлено")
                    st.rerun()

            with col_delete:
                if pages_cnt > 0:
                    st.caption(f"В разделе есть страницы ({pages_cnt}). Удаление недоступно.")
                else:
                    if st.button(
                        "Удалить раздел",
                        key=f"btn_delete_section_{section_id_local}",
                        use_container_width=True,
                        type="secondary",
                    ):
                        delete_section(section_id_local)
                        _bump_data_version()
                        add_event_log(topic="SECTION", subtopic="DELETE",notebook_id=selected_notebook_id, section_id=section_id_local, page_id=0, event="saction_name = " + old_section_name, body_html="")
                        st.success("Раздел удалён")
                        st.rerun()

    # --- список разделов + кнопки ➕/✎ (вернули) ---
    sections_df = pd.DataFrame()
    section_records: list = []
    if selected_notebook_id is not None:
        sections_df = _cached_df(
            "cache_sections",
            params=(selected_notebook_id,),
            fetcher=lambda: get_sections(selected_notebook_id),
        )
        section_records = list(sections_df.itertuples(index=False))

    with top_col3:
        select_col2, plus_col2, manage_col2 = st.columns([14, 2, 2])

        selected_section = None
        with select_col2:
            st.markdown("###### ")
            st.markdown("###### Раздел")
            if section_records:
                forced_section_id = st.session_state.pop("force_section_id", None)
                if forced_section_id is not None:
                    target_row = next(
                        (r for r in section_records if int(getattr(r, "id", 0)) == int(forced_section_id)),
                        None,
                    )
                    if target_row is not None:
                        st.session_state["section_selector"] = target_row

                selected_section = st.selectbox(
                    label="",
                    options=section_records,
                    format_func=lambda row: row.name,
                    key="section_selector",
                    label_visibility="collapsed",
                )
                selected_section_id = int(selected_section.id)
            else:
                st.warning("В книге нет разделов.")

        if can_edit_notebook:
            with plus_col2:
                st.markdown("###### ")
                st.markdown("###### ")
                if st.button("➕", key="open_new_section_dialog", help="Создать новый раздел", use_container_width=True):
                    new_section_dialog()

            with manage_col2:
                st.markdown("###### ")
                st.markdown("###### ")
                if selected_section is not None:
                    if st.button("✎", key="open_section_manage_dialog", help="Переименовать или удалить раздел", use_container_width=True):
                        section_manage_dialog(selected_section)

    # ---------- Загрузка страниц ----------
    dept_notebook_ids = filtered_notebooks_df["id"].astype(int).tolist() if not filtered_notebooks_df.empty else []

    if search_text:
        # ✅ при поиске ищем по ВСЕМ доступным книгам пользователя
        search_notebook_id = None
        search_section_id = None
        search_allowed_ids = visible_notebook_ids
    else:
        # ✅ без поиска показываем только страницы выбранного раздела выбранной книги
        if not selected_section_id:
            pages_df = pd.DataFrame()
        search_notebook_id = selected_notebook_id
        search_section_id = selected_section_id
        search_allowed_ids = [selected_notebook_id] if selected_notebook_id else []

    if not search_text and not selected_section_id:
        pages_df = pd.DataFrame()
    else:
        pages_df = load_pages_df(
            search_notebook_id,
            search_section_id,
            search_allowed_ids,
            search_text or None,
            search_tags_only,
        )


    # ✅ защита: load_pages_df / run_fetch_df должны возвращать DataFrame
    if pages_df is None:
        pages_df = pd.DataFrame()


    if pages_df.empty:
        pages_df = pd.DataFrame(
            columns=[
                "id",
                "title",
                "status",
            ]
        )

    # ---------- Кнопки "Новая страница" и "Импорт страниц" ----------
    new_page_clicked = False

    if can_edit_notebook:

        @st.dialog("Импорт страниц", width="large")
        def import_pages_dialog():
            if not selected_section_id:
                st.info("Выберите раздел для импорта.")
                return

            uploaded = st.file_uploader(
                "Загрузите файлы для импорта",
                type=["mht", "docx", "xlsx", "xlsm", "csv", "png", "txt", "sql"],
                accept_multiple_files=True,
                key="import_files_modal",
            )

            if uploaded and st.button("Импортировать", key="import_files_btn_modal"):
                imported = 0
                last_imported_id: int | None = None
                errors: list[str] = []
                for file in uploaded:
                    try:
                        file_name = file.name or "file"
                        _, ext = os.path.splitext(file_name)
                        ext = ext.lower()
                        data = file.getvalue()

                        if ext == ".mht":
                            pages = parse_mht_to_pages(data, file_name)  # ✅ sanitized per page
                            for title, body_html in pages:
                                last_imported_id = insert_page_with_content(
                                    selected_section_id, title, body_html, selected_login
                                )
                                add_event_log(topic="PAGE", subtopic="CREATE_IMPORT_.MHT",notebook_id=selected_notebook_id, section_id=selected_section_id, page_id=last_imported_id, event="page_title = " + title, body_html=body_html)
                                imported += 1
                            continue

                        page_title = file_name
                        body_html = ""

                        if ext == ".docx":
                            body_html = _docx_bytes_to_html(data)
                        elif ext in (".xlsx", ".xlsm"):
                            body_html = _excel_bytes_to_html(data)
                        elif ext == ".csv":
                            body_html = _csv_bytes_to_html(data)
                        elif ext in (".txt", ".sql"):
                            body_html = _text_bytes_to_html(data)
                        elif ext == ".png":
                            body_html = _png_bytes_to_html(data)
                        else:
                            raise ValueError(f"Неподдерживаемый тип файла: {ext}")

                        last_imported_id = insert_page_with_content(
                            selected_section_id, page_title, body_html, selected_login
                        )
                        add_event_log(topic="PAGE", subtopic="CREATE_IMPORT_" + str(ext.upper()),notebook_id=selected_notebook_id, section_id=selected_section_id, page_id=last_imported_id, event="page_title = " + page_title, body_html=body_html)
                        imported += 1
                    except Exception as exc:
                        errors.append(f"{file.name}: {exc}")

                if imported:
                    st.success(f"Импортировано {imported} страниц")
                    if last_imported_id is not None:
                        st.session_state["current_page_id"] = last_imported_id
                        st.session_state["force_page_id"] = last_imported_id
                        st.session_state["reset_page_search_once"] = True
                    _bump_data_version()
                    st.rerun()
                if errors:
                    st.warning(";\n".join(errors))

        btn_col_new, btn_col_import = st.sidebar.columns([1, 1])
        with btn_col_new:
            new_page_clicked = st.button("Новая страница", use_container_width=True)
        with btn_col_import:
            import_clicked = st.button("Импорт страниц", use_container_width=True)

        if import_clicked:
            import_pages_dialog()

    if new_page_clicked:
        if not selected_section_id:
            st.sidebar.warning("Сначала создайте и выберите раздел.")
        else:
            new_page_id = create_page(section_id=selected_section_id, user_login=selected_login, title=None)
            st.session_state["current_page_id"] = new_page_id
            st.session_state["force_page_id"] = new_page_id
            st.session_state["reset_page_search_once"] = True
            # ✅ открыть редактор ровно один раз на следующем ререндере
            st.session_state["open_editor_once_for_page"] = new_page_id
            st.session_state["open_editor_once_edit_mode"] = True
            _bump_data_version()

            add_event_log(topic="PAGE", subtopic="CREATE_HAND",notebook_id=selected_notebook_id, section_id=selected_section_id, page_id=new_page_id, event="create: page_id = " + str(new_page_id), body_html="")

            st.rerun()

    # ---------- Список страниц ----------
    # ✅ добавили статус в отображаемый df
    df_display = pages_df[["id", "title", "status"]].copy().reset_index(drop=True)

    if "id" in df_display.columns:
        # гарантируем числовые id, чтобы корректно проставлять выбранные строки
        df_display["id"] = pd.to_numeric(df_display["id"], errors="coerce")

    # нормализация на всякий случай (если NULL)
    if "status" in df_display.columns:
        df_display["status"] = df_display["status"].fillna("#FFFFFF").astype(str)

    # снимок для сравнения изменений (чтобы поймать, какие status реально поменялись)
    if "pages_grid_prev_df" not in st.session_state:
        st.session_state["pages_grid_prev_df"] = df_display.copy()

    # если сменился контекст (книга/раздел/поиск), обновляем "предыдущее" состояние,
    # чтобы не ловить ложные изменения статусов при навигации
    current_pages_signature = (selected_notebook_id, selected_section_id, search_text, search_tags_only)
    prev_signature = st.session_state.get("pages_grid_signature")
    if prev_signature != current_pages_signature:
        st.session_state["pages_grid_signature"] = current_pages_signature
        st.session_state["pages_grid_prev_df"] = df_display.copy()

    gb = GridOptionsBuilder.from_dataframe(df_display)

    # selection
    gb.configure_selection("single", use_checkbox=False)

    # ВАЖНО: id скрыт и не участвует в подгонке ширины
    gb.configure_column("id", header_name="ID", hide=True, suppressSizeToFit=True)

    # ✅ колонка кружка status (СПРАВА от "Страница")
    gb.configure_column(
        "title",
        header_name="Страница",
        flex=1,
        minWidth=140,
        resizable=True,
    )
    gb.configure_column(
        "status",
        header_name="",
        width=52,
        minWidth=40,
        maxWidth=60,
        flex=0,
        sortable=False,
        filter=False,
        editable=False,
        resizable=False,
        valueFormatter=STATUS_VALUE_FORMATTER,
        cellStyle=STATUS_CELL_STYLE,
        suppressSizeToFit=False,
    )

    # (опционально) defaultColDef — на всякий случай
    gb.configure_default_column(resizable=True)

    force_page_id = st.session_state.pop("force_page_id", None)
    if force_page_id is not None and not df_display.empty:
        try:
            row_index = int(df_display.index[df_display["id"] == force_page_id][0])
            gb.configure_selection("single", pre_selected_rows=[row_index])
        except Exception:
            pass

    # JS: растянуть колонку на ширину грида (и при первом рендере, и при resize окна)
    on_grid_ready = JsCode(
        """
        function(params) {
            try { params.api.sizeColumnsToFit(); } catch(e) {}

            setTimeout(function(){ try { params.api.sizeColumnsToFit(); } catch(e) {} }, 50);
            setTimeout(function(){ try { params.api.sizeColumnsToFit(); } catch(e) {} }, 200);

            if (!window.__notes_pages_resize_bound) {
                window.__notes_pages_resize_bound = true;
                window.addEventListener('resize', function() {
                    try { params.api.sizeColumnsToFit(); } catch(e) {}
                });
            }
        }
        """
    )

    on_first_data_rendered = JsCode(
        """
        function(params) {
            try { params.api.sizeColumnsToFit(); } catch(e) {}
        }
        """
    )

    gb.configure_grid_options(
        onGridReady=on_grid_ready,
        onFirstDataRendered=on_first_data_rendered,
        onGridSizeChanged=JsCode("function(params){ try { params.api.sizeColumnsToFit(); } catch(e) {} }"),
        # ✅ dblclick по кружку только владельцам
        onCellDoubleClicked=(STATUS_DBLCLICK_HANDLER if can_edit_notebook else None),
    )

    list_container = st.sidebar.container()
    with list_container:
        grid_response_pages = AgGrid(
            df_display,
            gridOptions=gb.build(),
            enable_enterprise_modules=False,
            # ✅ важно: чтобы изменения из JS (setDataValue) возвращались в python
            # update_mode=GridUpdateMode.MODEL_CHANGED | GridUpdateMode.SELECTION_CHANGED,
            update_mode=GridUpdateMode.MODEL_CHANGED,
            data_return_mode=DataReturnMode.AS_INPUT,
            height=650,
            fit_columns_on_grid_load=True,
            allow_unsafe_jscode=True,
        )


    # --- безопасно достаём data из grid_response_pages (может быть list или DataFrame) ---
    raw_data = grid_response_pages.get("data", [])
    if isinstance(raw_data, pd.DataFrame):
        df_after = raw_data.copy()
    elif isinstance(raw_data, list):
        df_after = pd.DataFrame(raw_data)
    else:
        df_after = pd.DataFrame()

    df_before = st.session_state.get("pages_grid_prev_df", df_display.copy())


    if not df_after.empty and "id" in df_after.columns and "status" in df_after.columns:
        df_after["status"] = df_after["status"].fillna("#FFFFFF").astype(str)
        df_before_local = df_before.copy()

        if not isinstance(df_before, pd.DataFrame):
            df_before = df_display.copy()


        merged = df_after[["id", "status"]].merge(
            df_before_local[["id", "status"]],
            on="id",
            how="left",
            suffixes=("_new", "_old"),
        )

        changed = merged[merged["status_new"] != merged["status_old"]]
        if not changed.empty:
            if can_edit_notebook:
                for row in changed.itertuples(index=False):
                    page_id_upd = int(row.id)
                    status_old = str(row.status_old or "#FFFFFF").strip()
                    new_status = str(row.status_new or "#FFFFFF").strip()

                    # защита: разрешаем только заданные цвета, иначе ставим белый
                    new_status = new_status.upper()
                    allowed = {"#FFFFFF", "#FFF59D", "#90CAF9", "#A5D6A7", "#EF9A9A", "#CE93D8"}
                    if new_status not in allowed:
                        new_status = "#FFFFFF"

                    # Уберем изменение статуса при создании страницы, когда предыдущего статуса нет Nane  and status_old is not None
                    if status_old != "nan":
                        run_execute(
                            f"""
                            UPDATE {PAGES_TABLE}
                            SET status = '{_escape(new_status)}'
                            WHERE id = {int(page_id_upd)}
                                AND COALESCE(status,'') <> '{_escape(new_status)}'
                            """
                        )

                        add_event_log(topic="PAGE", subtopic="CHANGE_STATUS",notebook_id=selected_notebook_id, section_id=selected_section_id, page_id=page_id_upd, event="new_status = " + _escape(new_status) + "; old_status = " + status_old, body_html="")
                # обновляем "предыдущее" состояние, чтобы не обновлять повторно
                st.session_state["pages_grid_prev_df"] = df_after.copy()

                # перечитать/перерисовать, чтобы pages_df внизу был синхронен
                _bump_data_version()
                st.rerun()
            else:
                # не владелец — откатываем снимок и предупреждаем
                st.session_state["pages_grid_prev_df"] = df_display.copy()
                st.warning("Менять статус страниц может только владелец книги.")
                st.rerun()
    else:
        st.session_state["pages_grid_prev_df"] = df_after.copy() if not df_after.empty else df_display.copy()

    # --- выбор страницы из грида (ОБЯЗАТЕЛЬНО объявляем page_id заранее) ---
    page_id: int | None = None

    selected_rows = grid_response_pages.get("selected_rows", [])
    if isinstance(selected_rows, pd.DataFrame):
        selected_rows = selected_rows.to_dict("records")

    if selected_rows:
        row = selected_rows[0]
        try:
            page_id = int(row["id"])
            st.session_state["current_page_id"] = page_id
        except Exception:
            page_id = None
    else:
        stored_page_id = st.session_state.get("current_page_id")
        if stored_page_id is not None and not pages_df.empty:
            if (pages_df["id"].astype(int) == int(stored_page_id)).any():
                page_id = int(stored_page_id)

    # st.sidebar.caption(f"selected_rows: {len(grid_response_pages.get('selected_rows') or [])}")
    st.sidebar.caption(f"selected_rows:{page_id=}")
                       

    # ---------- Просмотр / редактирование выбранной страницы ----------
    if page_id is not None:
        page_detail_df = load_page_detail_df(page_id)
        if page_detail_df is None:
            page_detail_df = pd.DataFrame()
        if page_detail_df.empty:
            st.warning("Не удалось загрузить данные страницы.")
            st.stop()

        current_page = page_detail_df.iloc[0]
        current_title = current_page.get("title", "")
        current_html = current_page.get("body_html") or ""
        current_tag = current_page.get("tag") or ""

        dept_id_for_page = current_page.get("notebook_department_id")
        dept_name_for_page = department_map.get(dept_id_for_page, "") if dept_id_for_page else ""
        dept_prefix = f"[{dept_name_for_page}] " if dept_name_for_page else ""

        safe_title = current_title or f"Страница_{page_id}"
        page_path = f"{current_page['notebook_name']} > {current_page['section_name']} > {current_page['title']}"
        info_left, info_right = st.columns([22, 19])
        with info_left:
            st.caption(
                f"{dept_prefix}  {current_page['notebook_name']}  >  "
                f"{current_page['section_name']}  >  {current_page['title']}"
            )
            if current_tag:
                st.caption(f"Tag: {current_tag}")

        with info_right:
            sql_text = _sql_text_from_html(current_html or "", safe_title)
            page_link = build_page_deeplink(page_id)

            col_copy, col_link = st.columns([3, 3])
            with col_copy:
                # копируем именно то, что видите в preview (HTML для Word + plain fallback)
                safe_preview_html_for_copy = sanitize_html_safe(current_html or "")
                plain_for_copy = _html_to_plain_preserving_layout(current_html or "")

                render_copy_rich_button(
                    html=safe_preview_html_for_copy,
                    plain=plain_for_copy,
                    btn_key=f"copy_preview_rich_{page_id}",
                    button_label="📋 Копировать страницу",
                    ok_message="Скопировано.",
                    button_style="""
                        cursor:pointer;
                        width:100%;
                        padding:0.35rem 0.75rem;
                        border:1px solid rgba(49,51,63,.2);
                        border-radius:0.5rem;
                        background: white;
                        font-size:14px;
                        line-height:1.2;
                        white-space:nowrap;
                    """,
                    align="flex-end",
                    height=42,
                )


            with col_link:
                if page_link:
                    render_copy_text_button(
                        text=page_link,
                        btn_key=f"copy_page_link_{page_id}",
                        button_label="🔗 Копировать ссылку на страницу",
                        ok_message="Скопирована",
                        button_style="""
                            cursor:pointer;
                            width:100%;
                            padding:0.35rem 0.75rem;
                            border:1px solid rgba(49,51,63,.2);
                            border-radius:0.5rem;
                            background: white;
                            font-size:14px;
                            line-height:1.2;
                            white-space:normal;
                        """,
                        align="flex-end",
                        height=42,
                    )
                else:
                    st.caption("Заполните `config.app_base_url`.")

        safe_preview_html = sanitize_html_safe(current_html or "")
        highlight_query = (search_text or "").strip()
        highlight_block = ""
        if highlight_query and not pages_df.empty:
            highlight_block = f"""
            <style>
            mark.search-hit {{
                background-color: #fff59d;
                padding: 0 2px;
            }}
            </style>
            <script>
            (() => {{
                const rawQuery = {json.dumps(highlight_query)};
                const query = (rawQuery || "").trim();
                if (!query) return;

                const root = document.querySelector(".preview-body");
                if (!root) return;

                // Убираем предыдущие подсветки перед новым поиском
                root.querySelectorAll("mark.search-hit").forEach((el) => {{
                    const textNode = document.createTextNode(el.textContent);
                    el.replaceWith(textNode);
                }});

                const queryLower = query.toLowerCase();
                const queryLength = query.length;
                const walker = document.createTreeWalker(root, NodeFilter.SHOW_TEXT, {{
                    acceptNode: (node) => {{
                        const parent = node?.parentNode;
                        if (!parent) return NodeFilter.FILTER_REJECT;
                        const tag = parent.nodeName;
                        if (tag === "SCRIPT" || tag === "STYLE") return NodeFilter.FILTER_REJECT;
                        if (!node.textContent || !node.textContent.trim()) return NodeFilter.FILTER_SKIP;
                        return NodeFilter.FILTER_ACCEPT;
                    }},
                }});

                let firstMark = null;
                while (true) {{
                    const textNode = walker.nextNode();
                    if (!textNode) break;

                    const original = textNode.textContent || "";
                    const lower = original.toLowerCase();
                    let idx = lower.indexOf(queryLower);
                    if (idx === -1) continue;

                    let lastIndex = 0;
                    const frag = document.createDocumentFragment();
                    while (idx !== -1) {{
                        if (idx > lastIndex) {{
                            frag.appendChild(document.createTextNode(original.slice(lastIndex, idx)));
                        }}
                        const mark = document.createElement("mark");
                        mark.className = "search-hit";
                        mark.textContent = original.slice(idx, idx + queryLength);
                        frag.appendChild(mark);
                        if (!firstMark) firstMark = mark;
                        lastIndex = idx + queryLength;
                        idx = lower.indexOf(queryLower, lastIndex);
                    }}
                    if (lastIndex < original.length) {{
                        frag.appendChild(document.createTextNode(original.slice(lastIndex)));
                    }}
                    textNode.replaceWith(frag);
                }}

                if (firstMark && typeof firstMark.scrollIntoView === "function") {{
                    firstMark.scrollIntoView({{ behavior: "smooth", block: "center" }});
                }}
            }})();
            </script>
            """

        preview_html = f"""
        <style>
        .page-preview-wrapper {{
            border: 1px solid #d0d4da;
            border-radius: 6px;
            padding: 16px 18px;
            background-color: #ffffff;
            box-shadow: 0 1px 2px rgba(0,0,0,0.04);
            min-height: 560px;
            box-sizing: border-box;
        }}
        .preview-body *,
        .preview-body p,
        .preview-body li {{
            line-height: 1.15 !important;
        }}
        .preview-body p {{
            margin: 0.2em 0 !important;
        }}
        </style>

        <div class="page-preview-wrapper">
            <div class="preview-body">
                {safe_preview_html or "<p><em>Нет содержимого</em></p>"}
            </div>
        </div>
        {highlight_block}
        """
        components.html(preview_html, height=580, scrolling=True)

        # ✅ one-shot авто-открытие редактора (после создания/копирования), НЕ будет открываться после открытия вложений
        st.session_state.setdefault("open_editor_once_for_page", None)
        open_once = st.session_state.pop("open_editor_once_for_page", None)
        default_edit_mode = bool(st.session_state.pop("open_editor_once_edit_mode", False))
        if can_edit_notebook and open_once == page_id:
            edit_page_dialog(page_id, current_title, current_html, current_tag, current_page.get("status", ""), default_edit_mode=default_edit_mode)

        # --- кнопка редактирования + экспорт + вложения + перемещение ---
        col1, col2, col3, col4 = st.columns([1.8, 1.2, 3, 3])

        # ---------------- Редактировать ----------------
        with col1:
            if can_edit_notebook:
                if st.button("Редактировать страницу", key=f"open_edit_dialog_{page_id}", use_container_width=True):
                    edit_page_dialog(page_id, current_title, current_html, current_tag, current_page.get("status", ""), default_edit_mode=False)
            else:
                st.caption("Просмотр (редактирование недоступно)")


        # ---------------- Экспорт ----------------
        with col2:
            exp_export_nonce_key = f"exp_export_nonce_{page_id}"
            st.session_state.setdefault(exp_export_nonce_key, 0)

            # состояние docx на страницу
            docx_state_key = f"docx_state_{page_id}"  # dict | None
            st.session_state.setdefault(docx_state_key, {"data": None, "file_name": None, "mime": None, "error": None})

            def _collapse_export():
                st.session_state[exp_export_nonce_key] += 1

            exp_export_label = "Экспорт" + ("\u200b" * int(st.session_state[exp_export_nonce_key]))

            with st.expander(exp_export_label, expanded=False):
                safe_title2 = current_title or f"Страница_{page_id}"

                # ---------- ОДНА КНОПКА: .docx ----------
                # при нажатии: генерируем bytes, сохраняем в state и автокликаем download
                clicked_docx = st.button(".docx", key=f"btn_docx_{page_id}", use_container_width=True)

                if clicked_docx:
                    try:
                        buf = export_html_to_docx_bytes(current_html or "", safe_title2)
                        data_bytes = buf.getvalue() if hasattr(buf, "getvalue") else bytes(buf)

                        st.session_state[docx_state_key] = {
                            "data": data_bytes,
                            "file_name": _safe_filename(safe_title2, "docx"),
                            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            "error": None,
                        }
                    except Exception as exc:
                        st.session_state[docx_state_key] = {
                            "data": None,
                            "file_name": _safe_filename(safe_title2, "docx"),
                            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            "error": str(exc),
                        }

                docx_state = st.session_state.get(docx_state_key) or {}
                if docx_state.get("error"):
                    st.error(f"Ошибка формирования DOCX: {docx_state['error']}")

                # Если есть готовые bytes — рисуем download_button и автокликаем его (один клик пользователя)
                if docx_state.get("data"):
                    dl_key = f"dl_docx_hidden_{page_id}"

                    # сам download_button можно сделать невидимым, но он должен быть в DOM
                    st.download_button(
                        label="Скачать DOCX",
                        data=docx_state["data"],
                        file_name=docx_state.get("file_name") or _safe_filename(safe_title2, "docx"),
                        mime=docx_state.get("mime")
                        or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=dl_key,
                        on_click=_collapse_export,
                        use_container_width=True,
                    )

                    # авто-клик по кнопке download_button
                    # (ищем DOM-кнопку по data-testid + тексту; работает стабильно в Streamlit)
                    components.html(
                        f"""
                        <script>
                        (function() {{
                          try {{
                            const root = window.parent.document;
                            const buttons = root.querySelectorAll('button');
                            let target = null;
                            buttons.forEach(b => {{
                              if ((b.innerText || '').trim() === 'Скачать DOCX') target = b;
                            }});
                            if (target) {{
                              setTimeout(() => target.click(), 30);
                            }}
                          }} catch(e) {{}}
                        }})();
                        </script>
                        """,
                        height=0,
                    )

                    # чтобы не автокликало при следующем rerun (и не скачивало повторно)
                    # сбрасываем data сразу после постановки автоклика
                    st.session_state[docx_state_key]["data"] = None



                # ---------- SQL (как было) ----------
                st.download_button(
                    ".sql (utf-8)",
                    data=export_html_to_sql_bytes(current_html or "", safe_title2, encoding="utf-8"),
                    file_name=_safe_filename(safe_title2, "sql"),
                    mime="text/plain; charset=utf-8",
                    on_click=_collapse_export,
                    key=f"dl_sql_u8_{page_id}",
                )

                st.download_button(
                    ".sql (cp1251)",
                    data=export_html_to_sql_bytes(current_html or "", safe_title2, encoding="cp1251"),
                    file_name=_safe_filename(safe_title2, "sql"),
                    mime="text/plain; charset=windows-1251",
                    on_click=_collapse_export,
                    key=f"dl_sql_1251_{page_id}",
                )

                if st.button("eMail", key=f"open_email_dialog_{page_id}", use_container_width=True):
                    st.session_state["email_dialog_page_id"] = page_id

                if st.session_state.get("email_dialog_page_id") == page_id:
                    email_message_dialog(page_id, safe_title2, current_html or "", selected_login, page_path)
                    _collapse_export()



        # ---------------- Вложения ----------------
        with col3:
            exp_nonce_key = f"exp_files_nonce_{page_id}"
            up_nonce_files_key = f"uploader_nonce_files_{page_id}"
            up_nonce_links_key = f"uploader_nonce_links_{page_id}"
            question_number_key = f"attachment_question_number_{page_id}"

            st.session_state.setdefault(exp_nonce_key, 0)
            st.session_state.setdefault(up_nonce_files_key, 0)
            st.session_state.setdefault(up_nonce_links_key, 0)
            st.session_state.setdefault(question_number_key, "")

            exp_label = "Прикрепть файлы или ссылки" + ("\u200b" * int(st.session_state[exp_nonce_key]))

            with st.expander(exp_label, expanded=False):
                if can_edit_notebook:
                    uploader_key = f"files_uploader_{page_id}_{st.session_state[up_nonce_files_key]}"
                    uploaded_files = st.file_uploader(
                        "Загрузите файлы",
                        type=None,
                        accept_multiple_files=True,
                        key=uploader_key,
                    )

                    st.text_input("Номер вопроса", key=question_number_key)

                    save_files_clicked = st.button(
                        "Прикрепить файлы",
                        key=f"save_files_btn_{page_id}",
                    )

                    if save_files_clicked:
                        if not uploaded_files:
                            st.warning("Сначала выберите файлы для загрузки.")
                        else:
                            question_number_val = st.session_state.get(question_number_key, "")
                            try:
                                for uf in uploaded_files:
                                    save_file_attachment(page_id, uf, selected_login, question_number_val)

                                st.success("Файлы прикреплены")
                                st.session_state[up_nonce_files_key] += 1
                                st.session_state[exp_nonce_key] += 1
                                st.rerun()
                            except Exception as e:
                                st.error(f"Ошибка при сохранении файлов: {e}")

                    st.markdown("---")

                    link_title_key = f"page_link_title_{page_id}"
                    link_url_key = f"page_link_url_{page_id}"

                    link_title = st.text_input("Название для ссылки", key=link_title_key)
                    link_url = st.text_input("URL", key=link_url_key)

                    if st.button("Сохранить ссылку", key=f"btn_save_link_{page_id}", use_container_width=True):
                        question_number_val = st.session_state.get(question_number_key, "")
                        try:
                            save_link_attachment(page_id, link_url, link_title, selected_login, question_number_val)
                            st.success("Ссылка сохранена")
                            st.session_state[exp_nonce_key] += 1
                            st.rerun()
                        except ValueError as exc:
                            st.warning(str(exc))
                        except Exception as exc:
                            st.error(f"Ошибка при сохранении ссылки: {exc}")
                else:
                    st.caption("У вас нет прав на добавление файлов и ссылок.")




        # ---------------- Переместить/скопировать ----------------
        with col4:
            exp_move_nonce_key = f"exp_move_nonce_{page_id}"
            st.session_state.setdefault(exp_move_nonce_key, 0)

            def _collapse_move():
                st.session_state[exp_move_nonce_key] += 1

            exp_move_label = "Переместить или скопировать" + ("\u200b" * int(st.session_state[exp_move_nonce_key]))

            with st.expander(exp_move_label, expanded=False):
                st.write(f"Текущая страница: **{current_title or f'ID {page_id}'}**")

                if owned_notebooks_df.empty:
                    st.info(
                        "У вас нет записных книг, в которых вы являетесь владельцем. "
                        "Копирование и перемещение недоступны."
                    )
                else:
                    dest_nb_records = list(owned_notebooks_df.itertuples(index=False))

                    cur_nb_id = int(current_page["notebook_id"])
                    default_nb = next((r for r in dest_nb_records if int(r.id) == cur_nb_id), dest_nb_records[0])
                    nb_index = dest_nb_records.index(default_nb)

                    dest_notebook = st.selectbox(
                        "Записная книга (для копирования/перемещения)",
                        dest_nb_records,
                        format_func=lambda r: r.name,
                        index=nb_index,
                        key=f"move_nb_{page_id}",
                    )

                    dest_sections_df = get_sections(int(dest_notebook.id))
                    dest_sec_records = list(dest_sections_df.itertuples(index=False))

                    if not dest_sec_records:
                        st.warning("В выбранной книге нет разделов. Сначала создайте раздел.")
                    else:
                        cur_sec_id = int(current_page["section_id"])
                        default_sec = next((r for r in dest_sec_records if int(r.id) == cur_sec_id), dest_sec_records[0])
                        sec_index = dest_sec_records.index(default_sec)

                        dest_section = st.selectbox(
                            "Раздел",
                            dest_sec_records,
                            format_func=lambda r: r.name,
                            index=sec_index,
                            key=f"move_sec_{page_id}",
                        )

                        col_move, col_copy, col_cancel = st.columns(3)

                        with col_move:
                            if can_edit_notebook:
                                move_clicked = st.button(
                                    "Переместить",
                                    type="primary",
                                    key=f"btn_move_{page_id}",
                                    use_container_width=True,
                                )
                            else:
                                move_clicked = False
                                st.caption("Перемещение доступно только владельцу книги.")

                        with col_copy:
                            copy_clicked = st.button(
                                "Копировать",
                                key=f"btn_copy_{page_id}",
                                use_container_width=True,
                            )

                        with col_cancel:
                            cancel_clicked = st.button(
                                "Отмена",
                                key=f"btn_cancel_{page_id}",
                                use_container_width=True,
                            )

                        if move_clicked and can_edit_notebook:
                            run_execute(
                                f"""
                                UPDATE {PAGES_TABLE}
                                SET section_id = {int(dest_section.id)}, updated_at = NOW()
                                WHERE id = {int(page_id)}
                                """
                            )
                            st.success("Страница перемещена")
                            st.session_state["current_page_id"] = page_id
                            st.session_state["force_page_id"] = page_id
                            _collapse_move()

                            add_event_log(topic="PAGE", subtopic="MOVE",notebook_id=int(dest_notebook.id), section_id=int(dest_section.id), page_id=page_id, event="move: new_section_id = " + str(dest_section.id) + "; old_section_id = " + str(selected_section_id) + "; old_notebook_id = " + str(selected_notebook_id), body_html="")                           
                            st.rerun()

                        if copy_clicked:
                            new_page_id = run_scalar(
                                f"""
                                INSERT INTO {PAGES_TABLE}
                                    (section_id, title, tag, body_html, created_by)
                                VALUES
                                    ({int(dest_section.id)},
                                     '{_escape(current_title)}',
                                     '{_escape(current_tag)}',
                                     '{_escape(current_html)}',
                                     '{_escape(selected_login)}')
                                RETURNING id
                                """
                            )
                            add_event_log(topic="PAGE", subtopic="COPY",notebook_id=int(dest_notebook.id), section_id=int(dest_section.id), page_id=new_page_id, event="copy page: new_page_id = " + str(new_page_id) + "; from_page_id = " + str(current_page["id"]), body_html=current_page["body_html"]) 

                            if new_page_id is None:
                                st.error("Ошибка при копировании страницы (не получен новый id).")
                            else:
                                new_page_id = int(new_page_id)
                                st.success("Страница скопирована")
                                st.session_state["current_page_id"] = new_page_id
                                st.session_state["force_page_id"] = new_page_id
                                st.session_state["open_editor_once_for_page"] = new_page_id
                                st.session_state["open_editor_once_edit_mode"] = False
                                _collapse_move()
                                st.rerun()

                        if cancel_clicked:
                            _collapse_move()
                            st.rerun()



        # --- таблица вложений ---
        attachments_df = get_page_attachments(page_id)
        if not attachments_df.empty:
            att_display = attachments_df.copy()
            att_display["Размер"] = att_display["file_size"].apply(_format_file_size)
            att_display["Создано"] = pd.to_datetime(att_display["created_at"], errors="coerce").dt.strftime("%Y-%m-%d %H:%M")
            att_display["Тип"] = att_display["attachment_type"].map({"file": "Файл", "link": "Ссылка"}).fillna(att_display["attachment_type"])
            att_display["Название"] = att_display["file_name"]
            att_display["Автор"] = att_display["created_by"]
            att_display["URL"] = att_display["url"].fillna("")
            att_display["Номер вопроса"] = att_display["question_number"].apply(lambda v: str(v).strip() if pd.notna(v) else "")
            grid_df = att_display[["id", "Номер вопроса", "Название", "Тип", "Размер", "Создано", "Автор", "URL"]]

            links_data = []
            for row in attachments_df.itertuples(index=False):
                if row.attachment_type == "file":
                    payload = get_attachment_file(row.id)
                    if payload:
                        file_bytes, file_name, mime_type = payload
                        b64 = base64.b64encode(file_bytes).decode("ascii")
                        links_data.append({"id": int(row.id), "name": str(file_name), "href": f"data:{mime_type};base64,{b64}"})

            if links_data:
                js_payload = json.dumps(links_data, ensure_ascii=False)
                components.html(
                    f"""
                    <style>
                    html, body {{
                        margin: 0 !important;
                        padding: 0 !important;
                        height: 0 !important;
                        overflow: hidden !important;
                    }}
                    </style>
                    <script>
                    (function() {{
                        var data = {js_payload};
                        var doc = window.parent.document;
                        var container = doc.getElementById('attachments_hidden_links');
                        if (!container) {{
                            container = doc.createElement('div');
                            container.id = 'attachments_hidden_links';
                            container.style.display = 'none';
                            doc.body.appendChild(container);
                        }}
                        data.forEach(function(item) {{
                            var id = item.id;
                            var a = doc.getElementById('att_dl_' + id);
                            if (!a) {{
                                a = doc.createElement('a');
                                a.id = 'att_dl_' + id;
                                a.download = item.name;
                                a.href = item.href;
                                container.appendChild(a);
                            }} else {{
                                a.download = item.name;
                                a.href = item.href;
                            }}
                        }});
                    }})();
                    </script>
                    """,
                    height=1,
                    scrolling=False,
                )

            gb_att = GridOptionsBuilder.from_dataframe(grid_df)
            gb_att.configure_selection("single", use_checkbox=False)
            gb_att.configure_column("id", hide=True)
            gb_att.configure_column("URL", hide=True)
            gb_att.configure_column("Номер вопроса", width=100)
            gb_att.configure_column("Название", width=500)
            gb_att.configure_column("Размер", width=90)
            gb_att.configure_column("Тип", width=80)

            row_doubleclick_js = JsCode(
                """
                function (e) {
                    var d = e.data || {};
                    var id = d.id;
                    var type = d["Тип"];
                    var url = d["URL"] || "";

                    if (!id) { return; }

                    if (type === "Файл") {
                        var a = window.parent.document.getElementById("att_dl_" + id);
                        if (a) { a.click(); }
                    } else if (type === "Ссылка" && url) {
                        window.parent.open(url, "_blank");
                    }
                }
                """
            )
            gb_att.configure_grid_options(onRowDoubleClicked=row_doubleclick_js)

            grid_response = AgGrid(
                grid_df,
                gridOptions=gb_att.build(),
                enable_enterprise_modules=False,
                update_on=["selectionChanged"],
                height=220,
                fit_columns_on_grid_load=True,
                allow_unsafe_jscode=True,
            )

            selected_rows = grid_response.get("selected_rows", [])
            if isinstance(selected_rows, pd.DataFrame):
                selected_rows = selected_rows.to_dict("records")
            selected_att = selected_rows[0] if selected_rows else None

            if selected_att:
                att_id = int(selected_att["id"])
                att_type = selected_att.get("Тип")
                url_val = selected_att.get("URL") or ""

                if att_type == "Файл":
                    if can_edit_notebook:
                        del_col, _ = st.columns([1,  4])
                        with del_col:
                            if st.button("Удалить вложение", key=f"delete_attachment_{att_id}", use_container_width=True):
                                delete_attachment(att_id)
                                if st.session_state.get("download_att_id") == att_id:
                                    st.session_state["download_att_id"] = None
                                    st.session_state["download_payload"] = None
                                    st.session_state["download_error"] = None

                                add_event_log(topic="ATTACHMENT", subtopic="DELETE",notebook_id=selected_notebook_id, section_id=selected_section_id, page_id=page_id, event="delete: attachment_id = " + str(att_id) + ": attachment_type = " + selected_att.get("Тип") + ": name = " + selected_att.get("Название"), body_html="")
                                st.success("Вложение удалено")
                                st.rerun()


                elif att_type == "Ссылка":
                    if not url_val:
                        st.warning("Ссылка не указана.")
                    else:
                        del_col, _ = st.columns([1, 4])
                        if can_edit_notebook:
                            with del_col:
                                if st.button("Удалить ссылку", key=f"delete_link_{att_id}", use_container_width=True):
                                    delete_attachment(att_id)
                                    st.success("Ссылка удалена")
                                    st.rerun()


        # --- Удаление страницы (с проверкой вложений) + справка (в одну строку) ---
        col_a, col_b, col_c, col_d = st.columns([1.8, 1.2, 3, 3])

        with col_d:
            exp_help_nonce_key = f"exp_help_nonce_{page_id}"
            st.session_state.setdefault(exp_help_nonce_key, 0)

            def _open_help(topic: str) -> None:
                st.session_state[f"help_topic_{page_id}"] = topic
                st.session_state[exp_help_nonce_key] += 1

            exp_help_label = "Справка ?" + ("\u200b" * int(st.session_state[exp_help_nonce_key]))

            with st.expander(exp_help_label, expanded=False):
                st.button(
                    "🔎 Расширенный поиск",
                    key=f"help_search_open_{page_id}",
                    help="Правила AND/OR/NOT, теги, wildcard и приоритеты",
                    on_click=_open_help,
                    args=("search",),
                    use_container_width=True,
                )

            help_topic = st.session_state.pop(f"help_topic_{page_id}", None)
            if help_topic == "search":
                help_search_dialog()



        if can_edit_notebook:
            attachments_df_for_delete = attachments_df
            has_attachments = not attachments_df_for_delete.empty

            if has_attachments:
                st.warning("Удаление страницы запрещено: сначала удалите все прикреплённые файлы и ссылки.")

            with col_a:
                confirm_delete = st.checkbox(
                    "Подтвердить удаление",
                    key=f"confirm_delete_{page_id}",
                    disabled=has_attachments,
                )

            with col_b:
                delete_clicked = st.button(
                    "Удалить страницу",
                    key=f"delete_{page_id}",
                    type="secondary",
                    use_container_width=True,
                    disabled=has_attachments,
                )

            if delete_clicked:
                # если вложения есть — кнопка должна быть disabled,
                # но оставим защиту на всякий случай
                if has_attachments:
                    st.error("Нельзя удалить страницу: сначала удалите все вложения.")
                    st.stop()

                if not confirm_delete:
                    st.warning("Поставьте галочку для подтверждения.")
                    st.stop()

                _body_html = run_scalar(
                                            f"""
                                            SELECT body_html
                                            FROM {PAGES_TABLE}
                                            WHERE id = {int(page_id)}
                                            LIMIT 1
                                            """
                                        )
                delete_page(page_id)
                add_event_log(topic="PAGE", subtopic="DELETE",notebook_id=selected_notebook_id, section_id=selected_section_id, page_id=page_id, event="delete: page_id = " + str(page_id), body_html=_body_html)
                st.success("Страница удалена")

                if st.session_state.get("edit_dialog_page_id") == page_id:
                    st.session_state["edit_dialog_page_id"] = None

                st.session_state["current_page_id"] = None
                st.rerun()
        else:
            st.info("У вас права только на просмотр этой записной книжки.")



    st.logo("assets/logo.png", size="medium")


if __name__ == "__main__":
    main()
